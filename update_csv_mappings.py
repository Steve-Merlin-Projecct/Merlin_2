#!/usr/bin/env python3
"""
Update CSV Mapping Files to Match Templates

This script updates all CSV mapping files to accurately reflect the variables
actually used in their corresponding .docx templates.

Created: 2025-10-24
"""

import os
import csv
import re
from typing import Dict, List, Set
from docx import Document


class CSVMappingUpdater:
    """Updates CSV mapping files to match template variables"""

    # Extended standard variable definitions
    VARIABLE_DEFINITIONS = {
        # Personal Information
        'user_first_name': ('User Profile', 'First name', 'user_profile'),
        'user_last_name': ('User Profile', 'Last name', 'user_profile'),
        'user_email': ('User Profile', 'Email address', 'user_profile'),
        'user_phone': ('User Profile', 'Phone number', 'user_profile'),
        'user_linkedin': ('User Profile', 'LinkedIn URL', 'user_profile'),
        'user_city_prov': ('User Profile', 'City and province', 'user_profile'),
        'user_portfolio': ('User Profile', 'Portfolio URL', 'user_profile'),
        'user_github': ('User Profile', 'GitHub URL', 'user_profile'),
        'professional_title': ('User Profile', 'Professional title', 'user_profile'),

        # Professional Summary
        'professional_summary': ('Profile', 'Professional summary paragraph', 'user_profile'),
        'executive_summary': ('Profile', 'Executive summary paragraph', 'user_profile'),
        'executive_title': ('Profile', 'Executive title', 'user_profile'),
        'tech_specialty': ('Profile', 'Technology specialty', 'user_profile'),

        # Skills
        'technical_summary': ('Skills', 'Technical skills list', 'user_skills'),
        'methodology_summary': ('Skills', 'Methodologies list', 'user_skills'),
        'domain_summary': ('Skills', 'Domain expertise list', 'user_skills'),

        # Leadership Competencies
        'leadership_competency_1': ('Skills', 'Leadership competency 1', 'user_skills'),
        'leadership_competency_2': ('Skills', 'Leadership competency 2', 'user_skills'),
        'leadership_competency_3': ('Skills', 'Leadership competency 3', 'user_skills'),
        'leadership_competency_4': ('Skills', 'Leadership competency 4', 'user_skills'),

        # Skills-Based Achievements
        'achievement_technical_1': ('Skills', 'Technical achievement 1', 'sentence_bank_resume'),
        'achievement_technical_2': ('Skills', 'Technical achievement 2', 'sentence_bank_resume'),
        'achievement_leadership_1': ('Skills', 'Leadership achievement 1', 'sentence_bank_resume'),
        'achievement_leadership_2': ('Skills', 'Leadership achievement 2', 'sentence_bank_resume'),

        # Education
        'edu_1_name': ('Education', 'Educational institution name', 'user_education'),
        'edu_1_degree': ('Education', 'Degree type', 'user_education'),
        'edu_1_concentration': ('Education', 'Field of study', 'user_education'),
        'edu_1_specialization': ('Education', 'Specialization area', 'user_education'),
        'edu_1_grad_date': ('Education', 'Graduation date', 'user_education'),
        'edu_1_location': ('Education', 'Institution location', 'user_education'),

        # Work Experience - Job 1
        'work_experience_1_position': ('Work Experience', 'Job position/title', 'user_work_experience'),
        'work_experience_1_name': ('Work Experience', 'Company name', 'user_work_experience'),
        'work_experience_1_location': ('Work Experience', 'Job location', 'user_work_experience'),
        'work_experience_1_dates': ('Work Experience', 'Employment dates', 'user_work_experience'),
        'work_experience_1_context': ('Work Experience', 'Company context', 'user_work_experience'),
        'work_experience_1_tech_stack': ('Work Experience', 'Technology stack', 'user_work_experience'),
        'work_experience_1_skill1': ('Work Experience', 'Achievement 1', 'sentence_bank_resume'),
        'work_experience_1_skill2': ('Work Experience', 'Achievement 2', 'sentence_bank_resume'),
        'work_experience_1_skill3': ('Work Experience', 'Achievement 3', 'sentence_bank_resume'),
        'work_experience_1_skill4': ('Work Experience', 'Achievement 4', 'sentence_bank_resume'),
        'work_experience_1_skill5': ('Work Experience', 'Achievement 5', 'sentence_bank_resume'),
        'work_experience_1_skill6': ('Work Experience', 'Achievement 6', 'sentence_bank_resume'),

        # Work Experience - Job 2
        'work_experience_2_position': ('Work Experience', 'Job 2 position/title', 'user_work_experience'),
        'work_experience_2_name': ('Work Experience', 'Job 2 company name', 'user_work_experience'),
        'work_experience_2_location': ('Work Experience', 'Job 2 location', 'user_work_experience'),
        'work_experience_2_dates': ('Work Experience', 'Job 2 employment dates', 'user_work_experience'),
        'work_experience_2_skill1': ('Work Experience', 'Job 2 achievement 1', 'sentence_bank_resume'),
        'work_experience_2_skill2': ('Work Experience', 'Job 2 achievement 2', 'sentence_bank_resume'),
        'work_experience_2_skill3': ('Work Experience', 'Job 2 achievement 3', 'sentence_bank_resume'),

        # Volunteer Experience
        'volunteer_1_position': ('Volunteer', 'Volunteer position', 'user_volunteer'),
        'volunteer_1_name': ('Volunteer', 'Volunteer organization', 'user_volunteer'),
        'volunteer_1_location': ('Volunteer', 'Volunteer location', 'user_volunteer'),
        'volunteer_1_dates': ('Volunteer', 'Volunteer dates', 'user_volunteer'),
        'volunteer_1_description': ('Volunteer', 'Volunteer description', 'user_volunteer'),

        # Certifications
        'certifications_list': ('Certifications', 'Professional certifications', 'user_certifications'),

        # Cover Letter - Company Info
        'company_name': ('Job Data', 'Target company name', 'job_data'),
        'company_address': ('Job Data', 'Company address', 'job_data'),
        'company_city_prov': ('Job Data', 'Company city and province', 'job_data'),
        'hiring_manager_name': ('Job Data', 'Hiring manager full name', 'job_data'),
        'hiring_manager_first_name': ('Job Data', 'Hiring manager first name', 'job_data'),
        'hiring_manager_title': ('Job Data', 'Hiring manager title', 'job_data'),

        # Cover Letter - Content (General)
        'cover_letter_opening': ('Content', 'Opening paragraph', 'sentence_bank_cover_letter'),
        'cover_letter_skills_alignment': ('Content', 'Skills alignment paragraph', 'sentence_bank_cover_letter'),
        'cover_letter_achievement': ('Content', 'Achievement paragraph', 'sentence_bank_cover_letter'),
        'cover_letter_closing': ('Content', 'Closing paragraph', 'sentence_bank_cover_letter'),

        # Cover Letter - Content (Modern)
        'cover_letter_hook': ('Content', 'Opening hook', 'sentence_bank_cover_letter'),
        'cover_letter_value_prop': ('Content', 'Value proposition', 'sentence_bank_cover_letter'),
        'cover_letter_company_interest': ('Content', 'Company interest statement', 'sentence_bank_cover_letter'),

        # Cover Letter - Content (Email)
        'cover_letter_email_hook': ('Content', 'Email opening hook', 'sentence_bank_cover_letter'),
        'cover_letter_value_prop_brief': ('Content', 'Brief value proposition', 'sentence_bank_cover_letter'),
        'cover_letter_achievement_brief': ('Content', 'Brief achievement', 'sentence_bank_cover_letter'),
        'cover_letter_cta': ('Content', 'Call to action', 'sentence_bank_cover_letter'),

        # Cover Letter - Content (Career Change)
        'cover_letter_transition_hook': ('Content', 'Career transition hook', 'sentence_bank_cover_letter'),
        'cover_letter_transferable_skills': ('Content', 'Transferable skills paragraph', 'sentence_bank_cover_letter'),
        'cover_letter_enthusiasm': ('Content', 'Enthusiasm statement', 'sentence_bank_cover_letter'),
        'cover_letter_preparation': ('Content', 'Preparation statement', 'sentence_bank_cover_letter'),

        # Cover Letter - Content (T-Format)
        'cover_letter_intro': ('Content', 'Introduction paragraph', 'sentence_bank_cover_letter'),
        'cover_letter_summary': ('Content', 'Summary paragraph', 'sentence_bank_cover_letter'),
        'job_requirement_1': ('Content', 'Job requirement 1', 'job_data'),
        'job_requirement_2': ('Content', 'Job requirement 2', 'job_data'),
        'job_requirement_3': ('Content', 'Job requirement 3', 'job_data'),
        'job_requirement_4': ('Content', 'Job requirement 4', 'job_data'),
        'matching_qualification_1': ('Content', 'Matching qualification 1', 'sentence_bank_cover_letter'),
        'matching_qualification_2': ('Content', 'Matching qualification 2', 'sentence_bank_cover_letter'),
        'matching_qualification_3': ('Content', 'Matching qualification 3', 'sentence_bank_cover_letter'),
        'matching_qualification_4': ('Content', 'Matching qualification 4', 'sentence_bank_cover_letter'),

        # Metadata
        'current_date': ('Metadata', 'Document generation date', 'system_generated'),
    }

    def __init__(self, worktree_root: str):
        """
        Initialize CSV mapping updater

        Args:
            worktree_root: Path to worktree root
        """
        self.worktree_root = worktree_root
        self.template_dir = os.path.join(worktree_root, 'content_template_library')
        self.variable_pattern = re.compile(r'<<([^>]+)>>')
        self.updates_made = []

    def extract_template_variables(self, docx_path: str) -> Set[str]:
        """
        Extract all variables from a .docx template

        Args:
            docx_path: Path to .docx template file

        Returns:
            Set of variable names
        """
        doc = Document(docx_path)
        variables = set()

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            found = self.variable_pattern.findall(paragraph.text)
            variables.update(found)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        found = self.variable_pattern.findall(paragraph.text)
                        variables.update(found)

        return variables

    def update_csv_mapping(self, docx_path: str, csv_path: str):
        """
        Update CSV mapping file to match template variables

        Args:
            docx_path: Path to .docx template
            csv_path: Path to CSV mapping file
        """
        # Extract variables from template
        template_vars = self.extract_template_variables(docx_path)

        # Read existing CSV mapping
        existing_vars = set()
        if os.path.exists(csv_path):
            with open(csv_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    if row.get('Is_Variable') == 'TRUE' and row.get('Variable_name'):
                        var_names = row['Variable_name'].strip().split()
                        existing_vars.update(var_names)

        # Determine what needs to be added
        vars_to_add = template_vars - existing_vars
        vars_to_keep = template_vars & existing_vars

        if not vars_to_add:
            print(f"   ✓ {os.path.basename(csv_path)} - Already up to date")
            return

        # Create new CSV content
        csv_rows = []

        # Add header
        csv_rows.append({
            'Original_Text': 'Original_Text',
            'Is_Variable': 'Is_Variable',
            'Is_Static': 'Is_Static',
            'Is_discarded': 'Is_discarded',
            'Variable_name': 'Variable_name',
            'Variable_intention': 'Variable_intention',
            'Variable_category': 'Variable_category',
            'Static_Text': 'Static_Text',
            'Content_source': 'Content_source',
        })

        # Add all template variables
        for var in sorted(template_vars):
            if var in self.VARIABLE_DEFINITIONS:
                category, intention, source = self.VARIABLE_DEFINITIONS[var]
            else:
                # Default values for unknown variables
                category = 'Unknown'
                intention = f'Variable: {var}'
                source = 'unknown'

            # Create placeholder text based on variable name
            placeholder_text = var.replace('_', ' ').title()

            csv_rows.append({
                'Original_Text': placeholder_text,
                'Is_Variable': 'TRUE',
                'Is_Static': 'FALSE',
                'Is_discarded': 'FALSE',
                'Variable_name': var,
                'Variable_intention': intention,
                'Variable_category': category,
                'Static_Text': '',
                'Content_source': source,
            })

        # Write updated CSV
        with open(csv_path, 'w', encoding='utf-8', newline='') as f:
            fieldnames = ['Original_Text', 'Is_Variable', 'Is_Static', 'Is_discarded',
                         'Variable_name', 'Variable_intention', 'Variable_category',
                         'Static_Text', 'Content_source']
            writer = csv.DictWriter(f, fieldnames=fieldnames)

            for row in csv_rows:
                writer.writerow(row)

        self.updates_made.append({
            'csv_file': os.path.relpath(csv_path, self.worktree_root),
            'added_vars': sorted(vars_to_add),
            'kept_vars': len(vars_to_keep),
            'total_vars': len(template_vars),
        })

        print(f"   ✓ {os.path.basename(csv_path)} - Updated ({len(vars_to_add)} added, {len(template_vars)} total)")

    def update_all_mappings(self):
        """Update all CSV mapping files to match their templates"""
        print("=" * 80)
        print("UPDATING CSV MAPPING FILES")
        print("=" * 80)

        # Process all templates
        for root, dirs, files in os.walk(self.template_dir):
            for file in files:
                if file.endswith('.docx') and not file.startswith('~'):
                    docx_path = os.path.join(root, file)
                    csv_path = os.path.splitext(docx_path)[0] + '_mapping.csv'

                    try:
                        self.update_csv_mapping(docx_path, csv_path)
                    except Exception as e:
                        print(f"   ✗ {file} - ERROR: {e}")

        print("\n" + "=" * 80)
        print("UPDATE COMPLETE")
        print("=" * 80)
        print(f"Files updated: {len(self.updates_made)}")

        if self.updates_made:
            print("\nDetails:")
            for update in self.updates_made:
                print(f"\n  {update['csv_file']}")
                print(f"    Added {len(update['added_vars'])} variables")
                print(f"    Total {update['total_vars']} variables")
                if update['added_vars']:
                    print(f"    New variables: {', '.join(update['added_vars'][:5])}")
                    if len(update['added_vars']) > 5:
                        print(f"                   ... and {len(update['added_vars']) - 5} more")


def main():
    """Main execution function"""
    worktree_root = '/workspace/.trees/word-document-cover-letter-and-resume-template-var'

    updater = CSVMappingUpdater(worktree_root)
    updater.update_all_mappings()


if __name__ == '__main__':
    main()
