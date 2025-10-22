#!/usr/bin/env python3
"""
Production Variable Insertion Script for Manually Converted Templates

This script populates converted Word templates with actual data from JSON input,
preserving all formatting while replacing variable placeholders.

Author: Template Conversion System
Date: October 2024
Version: 1.0
"""

import json
import os
import sys
from datetime import datetime
from docx import Document
from typing import Dict, List, Optional, Tuple
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class TemplateVariableInserter:
    """
    Handles variable insertion into converted Word templates

    This class:
    - Loads converted templates with variable placeholders
    - Accepts JSON data for variable values
    - Replaces variables while preserving formatting
    - Generates final documents ready for use
    """

    def __init__(self):
        """Initialize the variable inserter"""
        self.templates_dir = "/workspace/content_template_library/manual_converted"
        self.output_dir = "/workspace/content_template_library/generated"
        self.variable_pattern = "<<{}>>"

        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)

        # Template definitions
        self.templates = {
            "restaurant_manager": {
                "file": "restaurant_manager_template_converted.docx",
                "variables": self._get_restaurant_manager_variables()
            },
            "accountant": {
                "file": "accountant_template_converted.docx",
                "variables": self._get_accountant_variables()
            },
            "uiux_designer": {
                "file": "uiux_designer_template_converted.docx",
                "variables": self._get_uiux_designer_variables()
            }
        }

    def _get_restaurant_manager_variables(self) -> List[str]:
        """Get list of variables for restaurant manager template"""
        return [
            "first_name", "last_name", "street_address", "city", "state",
            "zip_code", "phone", "email", "linkedin", "professional_summary",
            "position_title", "company_name", "start_date", "end_date",
            "skill_1", "skill_2", "skill_3", "skill_4", "skill_5", "skill_6",
            "education_institution", "degree", "graduation_date", "interests"
        ]

    def _get_accountant_variables(self) -> List[str]:
        """Get list of variables for accountant template"""
        return [
            "first_name", "last_name", "street_address", "city", "state",
            "zip_code", "phone", "email", "linkedin", "professional_summary",
            "position_title", "company_name", "start_date", "end_date",
            "degree", "minor", "education_institution", "graduation_date",
            "skill_1", "skill_2", "skill_3", "skill_4", "skill_5", "skill_6"
        ]

    def _get_uiux_designer_variables(self) -> List[str]:
        """Get list of variables for UI/UX designer template"""
        return [
            "first_name", "last_name", "email", "portfolio_website", "phone",
            "city", "job_title", "professional_bio", "education_institution",
            "degree", "major", "graduation_year", "skill_1", "skill_2",
            "skill_3", "skill_4", "position_1", "company_1", "start_date_1",
            "end_date_1", "job_description_1_1", "job_description_1_2",
            "job_description_1_3", "job_description_1_4"
        ]

    def populate_template(self, template_name: str, data: Dict[str, str],
                         output_filename: Optional[str] = None) -> str:
        """
        Populate a template with variable data

        Args:
            template_name: Name of template to use
            data: Dictionary of variable names to values
            output_filename: Optional custom output filename

        Returns:
            Path to generated document
        """
        if template_name not in self.templates:
            raise ValueError(f"Unknown template: {template_name}")

        template_info = self.templates[template_name]
        template_path = os.path.join(self.templates_dir, template_info["file"])

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")

        logger.info(f"Loading template: {template_path}")
        doc = Document(template_path)

        # Track replacements
        replacements_made = {}

        # Replace variables in paragraphs
        for paragraph in doc.paragraphs:
            for var_name in template_info["variables"]:
                if var_name in data:
                    placeholder = self.variable_pattern.format(var_name)
                    if placeholder in paragraph.text:
                        self._replace_in_paragraph(paragraph, placeholder, data[var_name])
                        replacements_made[var_name] = data[var_name]

        # Replace variables in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for var_name in template_info["variables"]:
                        if var_name in data:
                            placeholder = self.variable_pattern.format(var_name)
                            if placeholder in cell.text:
                                self._replace_in_cell(cell, placeholder, data[var_name])
                                replacements_made[var_name] = data[var_name]

        # Generate output filename
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{template_name}_{timestamp}.docx"

        output_path = os.path.join(self.output_dir, output_filename)

        # Save the populated document
        doc.save(output_path)
        logger.info(f"Document saved: {output_path}")
        logger.info(f"Variables replaced: {len(replacements_made)}")

        # Log any missing variables
        missing_vars = set(template_info["variables"]) - set(replacements_made.keys())
        if missing_vars:
            logger.warning(f"Variables not provided: {missing_vars}")

        return output_path

    def _replace_in_paragraph(self, paragraph, placeholder: str, value: str):
        """Replace placeholder in paragraph while preserving formatting"""
        if placeholder in paragraph.text:
            # Work with runs to preserve formatting
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    def _replace_in_cell(self, cell, placeholder: str, value: str):
        """Replace placeholder in table cell while preserving formatting"""
        for paragraph in cell.paragraphs:
            self._replace_in_paragraph(paragraph, placeholder, value)

    def populate_from_json(self, template_name: str, json_path: str,
                          output_filename: Optional[str] = None) -> str:
        """
        Populate template from JSON file

        Args:
            template_name: Name of template to use
            json_path: Path to JSON file with data
            output_filename: Optional custom output filename

        Returns:
            Path to generated document
        """
        with open(json_path, 'r') as f:
            json_data = json.load(f)

        # Extract variables based on template type
        data = self._extract_variables_from_json(template_name, json_data)

        return self.populate_template(template_name, data, output_filename)

    def _extract_variables_from_json(self, template_name: str, json_data: Dict) -> Dict[str, str]:
        """
        Extract variables from comprehensive JSON structure

        Maps the complex JSON structure to simple variable names
        """
        data = {}

        # Check if this is the comprehensive defaults format
        if "resume_data" in json_data:
            resume_data = json_data["resume_data"]

            # Personal information
            if "personal" in resume_data:
                personal = resume_data["personal"]
                data["first_name"] = personal.get("first_name", "")
                data["last_name"] = personal.get("last_name", "")
                data["email"] = personal.get("email", "")
                data["phone"] = personal.get("phone", "")
                data["city"] = personal.get("city", "")
                data["state"] = personal.get("province", "")  # Map province to state
                data["zip_code"] = personal.get("postal_code", "")
                data["street_address"] = personal.get("address", "123 Main Street")
                data["linkedin"] = personal.get("linkedin", "")

            # Professional summary
            data["professional_summary"] = resume_data.get("professional_summary", "")
            data["professional_bio"] = resume_data.get("professional_summary", "")

            # Experience
            if "experience" in resume_data and len(resume_data["experience"]) > 0:
                exp = resume_data["experience"][0]
                data["position_title"] = exp.get("position", "")
                data["company_name"] = exp.get("company", "")
                data["start_date"] = exp.get("start_date", "")
                data["end_date"] = exp.get("end_date", "")

                # For UI/UX template
                data["position_1"] = exp.get("position", "")
                data["company_1"] = exp.get("company", "")
                data["start_date_1"] = exp.get("start_date", "")
                data["end_date_1"] = exp.get("end_date", "")

                # Job descriptions
                bullets = exp.get("bullets", [])
                for i, bullet in enumerate(bullets[:4], 1):
                    data[f"job_description_1_{i}"] = bullet

            # Education
            if "education" in resume_data and len(resume_data["education"]) > 0:
                edu = resume_data["education"][0]
                data["education_institution"] = edu.get("institution", "")
                data["degree"] = edu.get("degree", "")
                data["major"] = edu.get("concentration", "")
                data["graduation_date"] = edu.get("graduation_date", "")
                data["graduation_year"] = edu.get("graduation_date", "")
                data["minor"] = ""  # Not in standard data

            # Skills
            if "skills" in resume_data:
                skills = resume_data["skills"]
                # Combine different skill categories
                all_skills = []
                if isinstance(skills, dict):
                    for category, skill_list in skills.items():
                        if isinstance(skill_list, str):
                            all_skills.extend([s.strip() for s in skill_list.split(",")])

                # Assign to numbered skill variables
                for i, skill in enumerate(all_skills[:6], 1):
                    data[f"skill_{i}"] = skill

            # Interests
            if "skills" in resume_data and "interests" in resume_data["skills"]:
                data["interests"] = resume_data["skills"]["interests"]

        # Add defaults for any critical missing fields
        if template_name == "uiux_designer":
            data["job_title"] = data.get("position_title", "UI/UX Designer")
            data["portfolio_website"] = "www.portfolio.com"

        return data

    def generate_sample_data(self, template_name: str) -> Dict[str, str]:
        """
        Generate sample data for testing

        Args:
            template_name: Template to generate sample data for

        Returns:
            Dictionary of sample data
        """
        base_data = {
            "first_name": "John",
            "last_name": "Doe",
            "street_address": "123 Main Street",
            "city": "New York",
            "state": "NY",
            "zip_code": "10001",
            "phone": "(555) 123-4567",
            "email": "john.doe@email.com",
            "linkedin": "linkedin.com/in/johndoe",
            "professional_summary": "Experienced professional with proven track record",
            "professional_bio": "Passionate professional dedicated to excellence",
            "position_title": "Senior Professional",
            "company_name": "Example Company",
            "start_date": "Jan 2020",
            "end_date": "Present",
            "education_institution": "University Name",
            "degree": "Bachelor of Science",
            "major": "Computer Science",
            "minor": "Business",
            "graduation_date": "May 2019",
            "graduation_year": "2019",
            "interests": "Technology, Innovation, Leadership"
        }

        # Add skills
        for i in range(1, 7):
            base_data[f"skill_{i}"] = f"Professional Skill {i}"

        # Add UI/UX specific fields
        if template_name == "uiux_designer":
            base_data["job_title"] = "Senior UI/UX Designer"
            base_data["portfolio_website"] = "www.johndoe-portfolio.com"
            base_data["position_1"] = "Senior UI/UX Designer"
            base_data["company_1"] = "Design Agency"
            base_data["start_date_1"] = "Jan 2020"
            base_data["end_date_1"] = "Present"
            base_data["job_description_1_1"] = "Led design initiatives for major clients"
            base_data["job_description_1_2"] = "Conducted user research and testing"
            base_data["job_description_1_3"] = "Created wireframes and prototypes"
            base_data["job_description_1_4"] = "Managed team of junior designers"

        return base_data


def main():
    """Main function demonstrating template population"""

    inserter = TemplateVariableInserter()

    print("=" * 60)
    print("TEMPLATE VARIABLE INSERTION SYSTEM")
    print("=" * 60)
    print()

    # Available templates
    print("Available Templates:")
    for name, info in inserter.templates.items():
        print(f"  - {name}: {info['file']}")
        print(f"    Variables: {len(info['variables'])}")
    print()

    # Test with sample data
    print("Testing with sample data...")

    for template_name in inserter.templates.keys():
        print(f"\nProcessing: {template_name}")

        try:
            # Generate sample data
            sample_data = inserter.generate_sample_data(template_name)

            # Populate template
            output_path = inserter.populate_template(
                template_name,
                sample_data,
                f"{template_name}_sample_output.docx"
            )

            print(f"  ✓ Success: {output_path}")

        except Exception as e:
            print(f"  ✗ Error: {str(e)}")

    # Test with Steve Glen data if available
    json_path = "/workspace/content_template_library/steve_glen_comprehensive_defaults.json"
    if os.path.exists(json_path):
        print("\n" + "=" * 60)
        print("Testing with Steve Glen data...")

        for template_name in inserter.templates.keys():
            print(f"\nProcessing: {template_name}")

            try:
                output_path = inserter.populate_from_json(
                    template_name,
                    json_path,
                    f"{template_name}_steve_glen.docx"
                )
                print(f"  ✓ Success: {output_path}")

            except Exception as e:
                print(f"  ✗ Error: {str(e)}")

    print("\n" + "=" * 60)
    print("Variable insertion complete!")
    print(f"Generated documents saved to: {inserter.output_dir}")


if __name__ == "__main__":
    main()