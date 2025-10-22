#!/usr/bin/env python3
"""
Template Converter Script
Converts resume templates to variable-based templates following the conversion guide.
"""

import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from typing import List, Dict, Set
from collections import defaultdict


def analyze_template(doc_path: str) -> Dict:
    """
    Analyze a Word document template to understand its structure.

    Returns:
        Dictionary with template analysis including:
        - Total paragraphs
        - Text content
        - Tables found
        - Potential sections
        - Hardcoded content that needs replacement
    """
    print(f"\n{'='*80}")
    print(f"ANALYZING: {doc_path}")
    print(f"{'='*80}\n")

    doc = Document(doc_path)
    analysis = {
        'path': doc_path,
        'total_paragraphs': 0,
        'paragraphs_with_text': 0,
        'tables': 0,
        'sections': [],
        'content': [],
        'potential_names': [],
        'potential_companies': [],
        'potential_dates': [],
        'potential_degrees': []
    }

    # Analyze paragraphs
    for i, para in enumerate(doc.paragraphs):
        analysis['total_paragraphs'] += 1
        if para.text.strip():
            analysis['paragraphs_with_text'] += 1
            analysis['content'].append({
                'index': i,
                'text': para.text,
                'style': para.style.name,
                'alignment': para.alignment
            })

    # Analyze tables
    analysis['tables'] = len(doc.tables)

    # Extract text from tables
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    print(f"  Row {row_idx}, Cell {cell_idx}: {cell.text[:100]}")
                    analysis['content'].append({
                        'index': f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}",
                        'text': cell.text,
                        'style': 'table_cell',
                        'alignment': None
                    })

    # Print summary
    print(f"\nTemplate Summary:")
    print(f"  Total paragraphs: {analysis['total_paragraphs']}")
    print(f"  Paragraphs with text: {analysis['paragraphs_with_text']}")
    print(f"  Tables: {analysis['tables']}")

    print(f"\nFirst 10 paragraphs with content:")
    count = 0
    for item in analysis['content']:
        if isinstance(item['index'], int) and count < 10:
            print(f"  [{item['index']}] {item['style']}: {item['text'][:80]}")
            count += 1

    return analysis


def extract_all_text(doc_path: str) -> List[str]:
    """Extract all text from document including paragraphs and tables."""
    doc = Document(doc_path)
    all_text = []

    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    all_text.append(cell.text)

    return all_text


def print_full_content(doc_path: str):
    """Print all content from the document for manual review."""
    print(f"\n{'='*80}")
    print(f"FULL CONTENT: {doc_path}")
    print(f"{'='*80}\n")

    all_text = extract_all_text(doc_path)

    for i, text in enumerate(all_text):
        print(f"[{i:3d}] {text}")

    print(f"\nTotal text elements: {len(all_text)}")


if __name__ == "__main__":
    # Analyze both templates
    template_4_path = "/workspace/.trees/convert-raw-template-files-into-ready-for-producti/content_template_library/manual_converted/template_4_source.docx"
    template_5_path = "/workspace/.trees/convert-raw-template-files-into-ready-for-producti/content_template_library/manual_converted/template_5_source.docx"

    # Analyze and print full content for Template 4
    analysis_4 = analyze_template(template_4_path)
    print_full_content(template_4_path)

    print("\n" + "="*80 + "\n")

    # Analyze and print full content for Template 5
    analysis_5 = analyze_template(template_5_path)
    print_full_content(template_5_path)
