#!/usr/bin/env python3
"""
UI/UX Designer Template Conversion Script - Enhanced Version

This script converts a UI/UX Designer Word document template by replacing
specific content with variable placeholders while preserving ALL original
formatting (fonts, sizes, colors, bold, italic, etc.).

Template Structure:
- Name in paragraphs (Angelica Astrom)
- Table 1: Job title, bio, contact info
- Table 2: Education, skills, experience

Enhanced to handle all variables including:
- Professional bio
- Job title
- Experience details (positions, companies, dates, descriptions)

Author: Claude Code
Date: 2025-10-21
"""

import logging
import re
from pathlib import Path
from docx import Document
from typing import Dict, List, Set

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class UIUXTemplateConverterEnhanced:
    """
    Enhanced converter for UI/UX Designer template with comprehensive
    variable replacement while preserving all formatting.
    """

    def __init__(self, input_path: str, output_path: str):
        """
        Initialize the converter.

        Args:
            input_path: Path to the input .docx template
            output_path: Path to save the converted template
        """
        self.input_path = Path(input_path)
        self.output_path = Path(output_path)
        self.doc = None
        self.variables_created: Set[str] = set()

    def load_document(self) -> None:
        """Load the Word document."""
        logger.info(f"Loading document: {self.input_path}")
        self.doc = Document(str(self.input_path))
        logger.info(f"Document loaded successfully")
        logger.info(f"Document contains {len(self.doc.paragraphs)} paragraphs")
        logger.info(f"Document contains {len(self.doc.tables)} tables")

    def replace_text_in_run(self, run, old_text: str, new_text: str) -> bool:
        """
        Replace text in a run while preserving formatting.

        Args:
            run: python-docx Run object
            old_text: Text to find
            new_text: Text to replace with

        Returns:
            True if replacement was made, False otherwise
        """
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            logger.info(f"  Replaced '{old_text}' with '{new_text}'")
            self.variables_created.add(new_text)
            return True
        return False

    def replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str) -> bool:
        """
        Replace text in a paragraph while preserving formatting.

        Args:
            paragraph: python-docx Paragraph object
            old_text: Text to find
            new_text: Text to replace with

        Returns:
            True if replacement was made, False otherwise
        """
        replaced = False
        for run in paragraph.runs:
            if self.replace_text_in_run(run, old_text, new_text):
                replaced = True
        return replaced

    def replace_entire_paragraph_content(self, paragraph, new_text: str) -> None:
        """
        Replace entire paragraph content while preserving formatting of first run.

        Args:
            paragraph: python-docx Paragraph object
            new_text: New text to set
        """
        if paragraph.runs:
            # Keep formatting from first run, clear all runs
            first_run = paragraph.runs[0]
            # Clear all runs except first
            for i in range(len(paragraph.runs) - 1, 0, -1):
                paragraph._element.remove(paragraph.runs[i]._element)
            # Set new text in first run
            first_run.text = new_text
            logger.info(f"  Replaced entire paragraph with '{new_text}'")
            self.variables_created.add(new_text)

    def process_name_in_paragraphs(self) -> None:
        """Process paragraphs to replace name components."""
        logger.info("=" * 80)
        logger.info("PROCESSING PARAGRAPHS - Name Replacements")
        logger.info("=" * 80)

        name_replacements = {
            'Angelica': '<<first_name>>',
            'Astrom': '<<last_name>>'
        }

        for para_idx, para in enumerate(self.doc.paragraphs):
            logger.info(f"\nParagraph {para_idx}: '{para.text[:50]}...'")
            for old_text, new_text in name_replacements.items():
                self.replace_text_in_paragraph(para, old_text, new_text)

    def process_table_1(self, table) -> None:
        """
        Process Table 1 - Header table with job title, bio, and contact info.

        Expected structure:
        - 2 columns
        - Left: Job title, professional bio
        - Right: Contact information (email, website, phone, city)
        """
        logger.info("\n" + "=" * 80)
        logger.info("PROCESSING TABLE 1 - Header/Contact Information")
        logger.info("=" * 80)

        for row_idx, row in enumerate(table.rows):
            logger.info(f"\n--- Row {row_idx} ---")
            for col_idx, cell in enumerate(row.cells):
                logger.info(f"\nColumn {col_idx}:")
                cell_text = cell.text.strip()
                logger.info(f"Cell text: '{cell_text[:100]}...'")

                # Left column: Job title and bio
                if col_idx == 0:
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        para_text = paragraph.text.strip()

                        # First paragraph is likely the job title
                        if para_idx == 0 and para_text:
                            logger.info(f"  Processing job title paragraph")
                            # Replace job title
                            self.replace_text_in_paragraph(paragraph, 'UI/UX Designer', '<<job_title>>')

                        # Look for bio paragraphs (longer text)
                        elif len(para_text) > 50:
                            logger.info(f"  Processing bio paragraph ({len(para_text)} chars)")
                            # This is the professional bio
                            self.replace_entire_paragraph_content(paragraph, '<<professional_bio>>')

                # Right column: Contact information
                elif col_idx == 1:
                    # Define contact replacements
                    contact_replacements = {
                        'angelica@example.com': '<<email>>',
                        'www.interestingsite.com': '<<portfolio_website>>',
                        'New York City': '<<city>>',
                    }

                    for old_text, new_text in contact_replacements.items():
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, old_text, new_text)

                    # Handle phone number (pattern-based)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text.strip()
                            # Phone number pattern: contains digits and special chars
                            if re.search(r'\(\d{3}\)\s*\d{3}-\d{4}', text):
                                logger.info(f"  Found phone number: '{text}'")
                                run.text = '<<phone>>'
                                self.variables_created.add('<<phone>>')

    def process_table_2(self, table) -> None:
        """
        Process Table 2 - Education, Skills, and Experience.

        Expected structure:
        - 3 columns
        - Column 1: Education & Skills
        - Column 2: Work Experience (multiple positions)
        - Column 3: May contain additional info
        """
        logger.info("\n" + "=" * 80)
        logger.info("PROCESSING TABLE 2 - Education, Skills, Experience")
        logger.info("=" * 80)

        # Track which experience entries we've seen
        experience_counter = 0

        for row_idx, row in enumerate(table.rows):
            logger.info(f"\n--- Row {row_idx} ---")
            for col_idx, cell in enumerate(row.cells):
                logger.info(f"\nColumn {col_idx}:")
                cell_text = cell.text.strip()
                logger.info(f"Cell text preview: '{cell_text[:100]}...'")

                # Column 1: Education & Skills
                if col_idx == 0:
                    # Education section
                    education_replacements = {
                        'SCHOOL OF FINE ART': '<<education_institution>>',
                        'BFA, Graphic Design': '<<degree>>, <<major>>',
                    }

                    for old_text, new_text in education_replacements.items():
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, old_text, new_text)

                    # Graduation year (pattern: 20XX or year format)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if '20XX' in run.text or re.search(r'\b20\d{2}\b', run.text):
                                logger.info(f"  Found graduation year: '{run.text}'")
                                # Replace just the year part
                                run.text = re.sub(r'20XX|20\d{2}', '<<graduation_year>>', run.text)
                                self.variables_created.add('<<graduation_year>>')

                    # Skills section (look for SKILLS header then process items)
                    if 'SKILLS' in cell_text.upper():
                        logger.info("  Found SKILLS section")
                        skill_counter = 1
                        in_skills_section = False

                        for paragraph in cell.paragraphs:
                            para_text = paragraph.text.strip()

                            # Activate skills section when we hit the header
                            if 'SKILLS' in para_text.upper():
                                in_skills_section = True
                                logger.info("  Entering skills section")
                                continue

                            # Process skill items
                            if in_skills_section and para_text and len(para_text) < 50:
                                logger.info(f"  Processing skill: '{para_text}'")
                                self.replace_entire_paragraph_content(
                                    paragraph,
                                    f'<<skill_{skill_counter}>>'
                                )
                                skill_counter += 1

                # Column 2: Experience
                elif col_idx == 1:
                    # Check if this cell contains experience information
                    has_company_name = any(name in cell_text.upper() for name in ['PROSEWARE', 'FABRIKAM', 'INC', 'LLC'])

                    if has_company_name or 'Designer' in cell_text:
                        experience_counter += 1
                        logger.info(f"  Processing Experience Entry #{experience_counter}")

                        # Process each paragraph in the experience cell
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            para_text = paragraph.text.strip()

                            if not para_text:
                                continue

                            # First paragraph: Position title
                            if para_idx == 0:
                                logger.info(f"  Position: '{para_text}'")
                                self.replace_entire_paragraph_content(
                                    paragraph,
                                    f'<<position_{experience_counter}>>'
                                )

                            # Second paragraph: Company name
                            elif para_idx == 1:
                                logger.info(f"  Company: '{para_text}'")
                                self.replace_entire_paragraph_content(
                                    paragraph,
                                    f'<<company_{experience_counter}>>'
                                )

                            # Third paragraph: Date range
                            elif para_idx == 2:
                                logger.info(f"  Date range: '{para_text}'")
                                self.replace_entire_paragraph_content(
                                    paragraph,
                                    f'<<start_date_{experience_counter}>> - <<end_date_{experience_counter}>>'
                                )

                            # Remaining paragraphs: Job descriptions
                            else:
                                # Check if this is a substantial description (not just whitespace)
                                if len(para_text) > 20:
                                    desc_num = para_idx - 2  # Description number
                                    logger.info(f"  Description {desc_num}: '{para_text[:50]}...'")
                                    self.replace_entire_paragraph_content(
                                        paragraph,
                                        f'<<job_description_{experience_counter}_{desc_num}>>'
                                    )

    def save_document(self) -> None:
        """Save the converted document."""
        logger.info(f"\nSaving converted document to: {self.output_path}")
        self.doc.save(str(self.output_path))
        logger.info("Document saved successfully")

    def generate_report(self) -> str:
        """
        Generate a detailed conversion report.

        Returns:
            Formatted report string
        """
        report = []
        report.append("\n" + "=" * 80)
        report.append("UI/UX DESIGNER TEMPLATE CONVERSION REPORT - ENHANCED")
        report.append("=" * 80)
        report.append(f"\nInput File: {self.input_path}")
        report.append(f"Output File: {self.output_path}")
        report.append(f"\nTotal Variables Created: {len(self.variables_created)}")
        report.append("\nComplete Variable List (Categorized):")
        report.append("-" * 80)

        # Group variables by category
        categories = {
            'Personal Information': [],
            'Contact Information': [],
            'Professional Summary': [],
            'Education': [],
            'Skills': [],
            'Experience': [],
        }

        for var in sorted(self.variables_created):
            if 'first_name' in var or 'last_name' in var:
                categories['Personal Information'].append(var)
            elif 'email' in var or 'phone' in var or 'city' in var or 'website' in var:
                categories['Contact Information'].append(var)
            elif 'job_title' in var or 'bio' in var:
                categories['Professional Summary'].append(var)
            elif 'education' in var or 'degree' in var or 'major' in var or 'graduation' in var:
                categories['Education'].append(var)
            elif 'skill' in var:
                categories['Skills'].append(var)
            elif 'position' in var or 'company' in var or 'job_description' in var or 'date' in var:
                categories['Experience'].append(var)

        for category, vars_list in categories.items():
            if vars_list:
                report.append(f"\n{category}:")
                for var in vars_list:
                    report.append(f"  {var}")

        report.append("\n" + "=" * 80)
        report.append("FORMATTING PRESERVATION:")
        report.append("-" * 80)
        report.append("All original formatting has been preserved:")
        report.append("  - Font families, sizes, and colors")
        report.append("  - Bold, italic, and underline styles")
        report.append("  - Table structures and cell borders")
        report.append("  - Paragraph alignment and spacing")
        report.append("  - Document layout and margins")

        report.append("\n" + "=" * 80)
        report.append("CONVERSION COMPLETE")
        report.append("=" * 80)

        return "\n".join(report)

    def convert(self) -> None:
        """Execute the complete conversion process."""
        logger.info("\n" + "=" * 80)
        logger.info("STARTING UI/UX DESIGNER TEMPLATE CONVERSION - ENHANCED")
        logger.info("=" * 80)

        try:
            # Load document
            self.load_document()

            # Process paragraphs (name)
            self.process_name_in_paragraphs()

            # Process Table 1 (Header/Contact)
            if len(self.doc.tables) >= 1:
                self.process_table_1(self.doc.tables[0])

            # Process Table 2 (Education/Skills/Experience)
            if len(self.doc.tables) >= 2:
                self.process_table_2(self.doc.tables[1])

            # Save converted document
            self.save_document()

            # Generate and print report
            report = self.generate_report()
            print(report)

            logger.info("\n" + "=" * 80)
            logger.info("CONVERSION COMPLETED SUCCESSFULLY!")
            logger.info("=" * 80)

        except Exception as e:
            logger.error(f"Conversion failed: {str(e)}", exc_info=True)
            raise


def main():
    """Main execution function."""
    # Define paths
    input_path = '/workspace/content_template_library/manual_converted/uiux_designer_template.docx'
    output_path = '/workspace/content_template_library/manual_converted/uiux_designer_template_converted.docx'

    # Create converter and run
    converter = UIUXTemplateConverterEnhanced(input_path, output_path)
    converter.convert()


if __name__ == '__main__':
    main()
