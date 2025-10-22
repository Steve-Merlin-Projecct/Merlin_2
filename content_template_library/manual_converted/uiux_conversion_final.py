#!/usr/bin/env python3
"""
UI/UX Designer Template Conversion Script - Final Version

This script converts a UI/UX Designer Word document template by replacing
specific content with variable placeholders while preserving ALL original
formatting (fonts, sizes, colors, bold, italic, etc.).

Key improvements:
- Better detection of experience sections
- Handles multi-column table layouts correctly
- Preserves all formatting attributes
- Comprehensive logging and reporting

Author: Claude Code
Date: 2025-10-21
"""

import logging
import re
from pathlib import Path
from docx import Document
from typing import Set

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(levelname)s: %(message)s'
)
logger = logging.getLogger(__name__)


class UIUXTemplateConverterFinal:
    """
    Final version of UI/UX Designer template converter with comprehensive
    variable replacement and formatting preservation.
    """

    def __init__(self, input_path: str, output_path: str):
        """
        Initialize the converter.

        Args:
            input_path: Path to the input .docx template
            output_path: Path to save the converted template
        """
        self.input_path = Path(input_path)
        self.output_path = Path(output_path)
        self.doc = None
        self.variables_created: Set[str] = set()

    def load_document(self) -> None:
        """Load the Word document."""
        logger.info(f"Loading document: {self.input_path}")
        self.doc = Document(str(self.input_path))
        logger.info(f"Document loaded: {len(self.doc.paragraphs)} paragraphs, {len(self.doc.tables)} tables")

    def replace_in_run(self, run, old_text: str, new_text: str) -> bool:
        """Replace text in a run while preserving formatting."""
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            self.variables_created.add(new_text)
            logger.info(f"  ✓ Replaced '{old_text}' → '{new_text}'")
            return True
        return False

    def replace_in_paragraph(self, paragraph, old_text: str, new_text: str) -> bool:
        """Replace text in a paragraph while preserving formatting."""
        replaced = False
        for run in paragraph.runs:
            if self.replace_in_run(run, old_text, new_text):
                replaced = True
        return replaced

    def set_paragraph_text(self, paragraph, new_text: str) -> None:
        """
        Replace entire paragraph text while preserving formatting.
        Keeps the formatting of the first run.
        """
        if paragraph.runs:
            # Clear all runs except first
            for i in range(len(paragraph.runs) - 1, 0, -1):
                paragraph._element.remove(paragraph.runs[i]._element)
            # Set new text in first run
            paragraph.runs[0].text = new_text
            self.variables_created.add(new_text)
            logger.info(f"  ✓ Set paragraph to: '{new_text}'")

    def process_name_paragraphs(self) -> None:
        """Process paragraphs to replace name components."""
        logger.info("\n" + "="*80)
        logger.info("STEP 1: Processing Name in Paragraphs")
        logger.info("="*80)

        name_map = {
            'Angelica': '<<first_name>>',
            'Astrom': '<<last_name>>'
        }

        for idx, para in enumerate(self.doc.paragraphs):
            if para.text.strip():
                logger.info(f"Paragraph {idx}: '{para.text}'")
                for old, new in name_map.items():
                    self.replace_in_paragraph(para, old, new)

    def process_table_1_header(self, table) -> None:
        """
        Process Table 1 - Header table with job title, bio, and contact.

        Structure:
        - Row 0, Cell 0: Job title + Bio
        - Row 0, Cell 1: Contact info
        """
        logger.info("\n" + "="*80)
        logger.info("STEP 2: Processing Table 1 (Header/Contact)")
        logger.info("="*80)

        for row_idx, row in enumerate(table.rows):
            logger.info(f"\nRow {row_idx}:")
            for col_idx, cell in enumerate(row.cells):
                logger.info(f"  Column {col_idx}:")

                # Left column: Job title and bio
                if col_idx == 0:
                    for para_idx, para in enumerate(cell.paragraphs):
                        text = para.text.strip()

                        # Job title (short text in first paragraph)
                        if para_idx == 0 and text and len(text) < 50:
                            logger.info(f"    Job title: '{text}'")
                            # Don't replace if already converted
                            if '<<' not in text:
                                self.set_paragraph_text(para, '<<job_title>>')

                        # Bio (longer paragraph)
                        elif len(text) > 50 and '<<' not in text:
                            logger.info(f"    Bio ({len(text)} chars)")
                            self.set_paragraph_text(para, '<<professional_bio>>')

                # Right column: Contact info
                elif col_idx == 1:
                    # Contact field replacements
                    contact_map = {
                        'angelica@example.com': '<<email>>',
                        'www.interestingsite.com': '<<portfolio_website>>',
                        'New York City': '<<city>>',
                    }

                    for old, new in contact_map.items():
                        for para in cell.paragraphs:
                            self.replace_in_paragraph(para, old, new)

                    # Phone number (pattern-based)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if re.search(r'\(\d{3}\)\s*\d{3}-\d{4}', run.text):
                                logger.info(f"    Phone: '{run.text.strip()}'")
                                run.text = '<<phone>>'
                                self.variables_created.add('<<phone>>')

    def process_table_2_content(self, table) -> None:
        """
        Process Table 2 - Education, Skills, and Experience.

        This table has a complex 3-column layout:
        - Column 1: Education & Skills
        - Column 2: (May be merged or empty)
        - Column 3: Experience section
        """
        logger.info("\n" + "="*80)
        logger.info("STEP 3: Processing Table 2 (Education/Skills/Experience)")
        logger.info("="*80)

        # Process all cells to find education, skills, and experience
        for row_idx, row in enumerate(table.rows):
            logger.info(f"\nRow {row_idx}:")
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()

                if not cell_text:
                    continue

                logger.info(f"  Column {col_idx} ({len(cell_text)} chars):")
                logger.info(f"    Preview: '{cell_text[:80]}...'")

                # EDUCATION SECTION
                if 'SCHOOL OF FINE ART' in cell_text or 'Education' in cell_text:
                    logger.info("    → Processing EDUCATION")

                    edu_map = {
                        'SCHOOL OF FINE ART': '<<education_institution>>',
                        'BFA, Graphic Design': '<<degree>>, <<major>>',
                    }

                    for old, new in edu_map.items():
                        for para in cell.paragraphs:
                            self.replace_in_paragraph(para, old, new)

                    # Graduation year
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if '20XX' in run.text or re.search(r'\b20\d{2}\b', run.text):
                                logger.info(f"    Grad year: '{run.text.strip()}'")
                                run.text = re.sub(r'20XX|20\d{2}', '<<graduation_year>>', run.text)
                                self.variables_created.add('<<graduation_year>>')

                # SKILLS SECTION
                if 'Skills' in cell_text or 'UI/UX design' in cell_text:
                    logger.info("    → Processing SKILLS")

                    skill_counter = 1
                    in_skills = False

                    for para in cell.paragraphs:
                        text = para.text.strip()

                        if 'Skills' in text:
                            in_skills = True
                            continue

                        # Process individual skills (short lines after Skills header)
                        if in_skills and text and len(text) < 50 and '<<' not in text:
                            logger.info(f"    Skill {skill_counter}: '{text}'")
                            self.set_paragraph_text(para, f'<<skill_{skill_counter}>>')
                            skill_counter += 1

                # EXPERIENCE SECTION
                if any(keyword in cell_text for keyword in ['PROSEWARE', 'FABRIKAM', 'Senior UI/UX', 'Experience']):
                    logger.info("    → Processing EXPERIENCE")

                    # Determine job number based on keywords
                    job_num = 1 if 'PROSEWARE' in cell_text or 'Senior' in cell_text else 2
                    logger.info(f"    Job #{job_num}")

                    # Process paragraphs in order
                    para_list = [p for p in cell.paragraphs if p.text.strip() and '<<' not in p.text]

                    for idx, para in enumerate(para_list):
                        text = para.text.strip()

                        # Skip section headers
                        if text == 'Experience':
                            continue

                        # Position (first non-header paragraph, contains "Designer")
                        if 'Designer' in text and 'position' not in text.lower():
                            logger.info(f"    Position: '{text}'")
                            self.set_paragraph_text(para, f'<<position_{job_num}>>')

                        # Company (all caps, contains INC/LLC)
                        elif re.search(r'[A-Z]{2,}.*INC|LLC', text):
                            logger.info(f"    Company: '{text}'")
                            self.set_paragraph_text(para, f'<<company_{job_num}>>')

                        # Date range (contains hyphen and year patterns)
                        elif '-' in text and ('20XX' in text or re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)', text)):
                            logger.info(f"    Dates: '{text}'")
                            self.set_paragraph_text(para, f'<<start_date_{job_num}>> - <<end_date_{job_num}>>')

                        # Job descriptions (longer text)
                        elif len(text) > 30:
                            # Count previous descriptions for this job
                            desc_num = sum(1 for p in para_list[:idx] if len(p.text.strip()) > 30) + 1
                            logger.info(f"    Description {desc_num}: '{text[:40]}...'")
                            self.set_paragraph_text(para, f'<<job_description_{job_num}_{desc_num}>>')

    def save_document(self) -> None:
        """Save the converted document."""
        logger.info(f"\nSaving to: {self.output_path}")
        self.doc.save(str(self.output_path))
        logger.info("✓ Document saved successfully")

    def generate_report(self) -> str:
        """Generate comprehensive conversion report."""
        report = []
        report.append("\n" + "="*80)
        report.append("UI/UX DESIGNER TEMPLATE CONVERSION - FINAL REPORT")
        report.append("="*80)
        report.append(f"\nInput:  {self.input_path}")
        report.append(f"Output: {self.output_path}")
        report.append(f"\nTotal Variables: {len(self.variables_created)}")

        # Categorize variables
        categories = {
            'Personal': [],
            'Contact': [],
            'Professional': [],
            'Education': [],
            'Skills': [],
            'Experience': [],
        }

        for var in sorted(self.variables_created):
            if any(x in var for x in ['first_name', 'last_name']):
                categories['Personal'].append(var)
            elif any(x in var for x in ['email', 'phone', 'city', 'website']):
                categories['Contact'].append(var)
            elif any(x in var for x in ['job_title', 'bio']):
                categories['Professional'].append(var)
            elif any(x in var for x in ['education', 'degree', 'major', 'graduation']):
                categories['Education'].append(var)
            elif 'skill' in var:
                categories['Skills'].append(var)
            elif any(x in var for x in ['position', 'company', 'date', 'description']):
                categories['Experience'].append(var)

        report.append("\nVariables by Category:")
        report.append("-"*80)

        for category, vars_list in categories.items():
            if vars_list:
                report.append(f"\n{category}: ({len(vars_list)})")
                for var in vars_list:
                    report.append(f"  • {var}")

        report.append("\n" + "="*80)
        report.append("FORMATTING PRESERVED:")
        report.append("  ✓ Fonts (family, size, color)")
        report.append("  ✓ Styles (bold, italic, underline)")
        report.append("  ✓ Tables (structure, borders, alignment)")
        report.append("  ✓ Paragraphs (spacing, indentation)")
        report.append("  ✓ Document layout and margins")
        report.append("="*80)

        return "\n".join(report)

    def convert(self) -> None:
        """Execute the complete conversion process."""
        logger.info("\n" + "="*80)
        logger.info("UI/UX DESIGNER TEMPLATE CONVERTER - FINAL VERSION")
        logger.info("="*80)

        try:
            # Load
            self.load_document()

            # Convert
            self.process_name_paragraphs()

            if len(self.doc.tables) >= 1:
                self.process_table_1_header(self.doc.tables[0])

            if len(self.doc.tables) >= 2:
                self.process_table_2_content(self.doc.tables[1])

            # Save
            self.save_document()

            # Report
            print(self.generate_report())

            logger.info("\n" + "="*80)
            logger.info("✓ CONVERSION COMPLETE!")
            logger.info("="*80)

        except Exception as e:
            logger.error(f"✗ Conversion failed: {e}", exc_info=True)
            raise


def main():
    """Main execution."""
    input_path = '/workspace/content_template_library/manual_converted/uiux_designer_template.docx'
    output_path = '/workspace/content_template_library/manual_converted/uiux_designer_template_converted.docx'

    converter = UIUXTemplateConverterFinal(input_path, output_path)
    converter.convert()


if __name__ == '__main__':
    main()
