#!/usr/bin/env python3
"""
UI/UX Designer Template Conversion Script

This script converts a UI/UX Designer Word document template by replacing
specific content with variable placeholders while preserving ALL original
formatting (fonts, sizes, colors, bold, italic, etc.).

Template Structure:
- Name in paragraphs (Angelica Astrom)
- Table 1: Job title, bio, contact info
- Table 2: Education, skills, experience

Author: Claude Code
Date: 2025-10-21
"""

import logging
from pathlib import Path
from docx import Document
from typing import Dict, List, Tuple

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class UIUXTemplateConverter:
    """
    Converts UI/UX Designer template by replacing content with variables
    while preserving all formatting.
    """

    def __init__(self, input_path: str, output_path: str):
        """
        Initialize the converter.

        Args:
            input_path: Path to the input .docx template
            output_path: Path to save the converted template
        """
        self.input_path = Path(input_path)
        self.output_path = Path(output_path)
        self.doc = None
        self.variables_created = []

        # Define all replacement mappings
        self.replacements = self._define_replacements()

    def _define_replacements(self) -> Dict[str, str]:
        """
        Define all text-to-variable replacement mappings.

        Returns:
            Dictionary mapping original text to variable placeholders
        """
        return {
            # Name components (in paragraphs)
            'Angelica': '<<first_name>>',
            'Astrom': '<<last_name>>',

            # Job title
            'UI/UX Designer': '<<job_title>>',

            # Contact information
            'angelica@example.com': '<<email>>',
            'www.interestingsite.com': '<<portfolio_website>>',
            'New York City': '<<city>>',

            # Education
            'SCHOOL OF FINE ART': '<<education_institution>>',
            'BFA, Graphic Design': '<<degree>>, <<major>>',

            # Note: Phone numbers, dates, and multi-line content need special handling
        }

    def load_document(self) -> None:
        """Load the Word document."""
        logger.info(f"Loading document: {self.input_path}")
        self.doc = Document(str(self.input_path))
        logger.info(f"Document loaded successfully")
        logger.info(f"Document contains {len(self.doc.paragraphs)} paragraphs")
        logger.info(f"Document contains {len(self.doc.tables)} tables")

    def replace_in_run(self, run, old_text: str, new_text: str) -> bool:
        """
        Replace text in a run while preserving formatting.

        Args:
            run: python-docx Run object
            old_text: Text to find
            new_text: Text to replace with

        Returns:
            True if replacement was made, False otherwise
        """
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
        return False

    def replace_in_paragraph(self, paragraph, old_text: str, new_text: str) -> bool:
        """
        Replace text in a paragraph while preserving formatting.

        Args:
            paragraph: python-docx Paragraph object
            old_text: Text to find
            new_text: Text to replace with

        Returns:
            True if replacement was made, False otherwise
        """
        replaced = False
        for run in paragraph.runs:
            if self.replace_in_run(run, old_text, new_text):
                replaced = True
        return replaced

    def process_paragraphs(self) -> None:
        """Process all paragraphs and replace name components."""
        logger.info("Processing paragraphs for name replacements...")

        name_replacements = {
            'Angelica': '<<first_name>>',
            'Astrom': '<<last_name>>'
        }

        replacements_made = 0
        for para in self.doc.paragraphs:
            for old_text, new_text in name_replacements.items():
                if self.replace_in_paragraph(para, old_text, new_text):
                    logger.info(f"Replaced '{old_text}' with '{new_text}' in paragraph")
                    replacements_made += 1
                    if new_text not in self.variables_created:
                        self.variables_created.append(new_text)

        logger.info(f"Made {replacements_made} replacements in paragraphs")

    def replace_in_cell(self, cell, old_text: str, new_text: str) -> bool:
        """
        Replace text in a table cell while preserving formatting.

        Args:
            cell: python-docx Cell object
            old_text: Text to find
            new_text: Text to replace with

        Returns:
            True if replacement was made, False otherwise
        """
        replaced = False
        for paragraph in cell.paragraphs:
            if self.replace_in_paragraph(paragraph, old_text, new_text):
                replaced = True
        return replaced

    def process_table_1_contact_info(self, table) -> None:
        """
        Process Table 1 (Contact Information Table).

        Table 1 structure:
        - 2 columns
        - Left: Job title, bio
        - Right: Contact details
        """
        logger.info("Processing Table 1 (Contact Information)...")

        # Replacements for Table 1
        table1_replacements = {
            'UI/UX Designer': '<<job_title>>',
            'angelica@example.com': '<<email>>',
            'www.interestingsite.com': '<<portfolio_website>>',
            'New York City': '<<city>>',
        }

        replacements_made = 0
        for row in table.rows:
            for cell in row.cells:
                for old_text, new_text in table1_replacements.items():
                    if self.replace_in_cell(cell, old_text, new_text):
                        logger.info(f"Replaced '{old_text}' with '{new_text}' in Table 1")
                        replacements_made += 1
                        if new_text not in self.variables_created:
                            self.variables_created.append(new_text)

                # Handle professional bio (multi-paragraph)
                cell_text = cell.text
                if 'Lorem ipsum' in cell_text or 'professional' in cell_text.lower():
                    # This is likely the bio section
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip() and len(paragraph.text) > 50:
                            # Replace the entire bio paragraph
                            for run in paragraph.runs:
                                if run.text.strip():
                                    run.text = '<<professional_bio>>'
                                    logger.info("Replaced bio text with <<professional_bio>>")
                                    if '<<professional_bio>>' not in self.variables_created:
                                        self.variables_created.append('<<professional_bio>>')
                                    break
                            break

                # Handle phone number (pattern: digits with formatting)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text.strip()
                        # Look for phone number patterns
                        if any(char.isdigit() for char in text) and len(text) < 20:
                            # Check if it looks like a phone number
                            digit_count = sum(1 for c in text if c.isdigit())
                            if digit_count >= 7:  # At least 7 digits
                                run.text = '<<phone>>'
                                logger.info(f"Replaced phone number '{text}' with <<phone>>")
                                if '<<phone>>' not in self.variables_created:
                                    self.variables_created.append('<<phone>>')

        logger.info(f"Made {replacements_made} replacements in Table 1")

    def process_table_2_experience(self, table) -> None:
        """
        Process Table 2 (Education, Skills, Experience).

        Table 2 structure:
        - 3 columns
        - Column 1: Education & Skills
        - Column 2: Experience (multiple jobs)
        - Column 3: Additional content
        """
        logger.info("Processing Table 2 (Education, Skills, Experience)...")

        # Education replacements
        education_replacements = {
            'SCHOOL OF FINE ART': '<<education_institution>>',
            'BFA, Graphic Design': '<<degree>>, <<major>>',
        }

        # Experience replacements (pattern-based for positions 1 and 2)
        experience_patterns = [
            ('Senior UI/UX Designer', '<<position_1>>'),
            ('PROSEWARE, INC.', '<<company_1>>'),
            ('UI/UX Designer', '<<position_2>>'),
            ('FABRIKAM, INC.', '<<company_2>>'),
        ]

        replacements_made = 0

        # Process each row and cell
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text

                # Column 1: Education & Skills
                if col_idx == 0:
                    # Education replacements
                    for old_text, new_text in education_replacements.items():
                        if self.replace_in_cell(cell, old_text, new_text):
                            logger.info(f"Replaced '{old_text}' with '{new_text}' in Table 2")
                            replacements_made += 1
                            if new_text not in self.variables_created:
                                self.variables_created.append(new_text)

                    # Replace graduation year (20XX pattern)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if '20XX' in run.text or ('20' in run.text and 'XX' in run.text):
                                run.text = run.text.replace('20XX', '<<graduation_year>>')
                                logger.info("Replaced graduation year with <<graduation_year>>")
                                if '<<graduation_year>>' not in self.variables_created:
                                    self.variables_created.append('<<graduation_year>>')

                    # Handle skills (bulleted or listed items)
                    if 'SKILLS' in cell_text.upper():
                        skill_counter = 1
                        for paragraph in cell.paragraphs:
                            text = paragraph.text.strip()
                            # Skip headers and empty lines
                            if text and 'SKILLS' not in text.upper() and len(text) < 30:
                                for run in paragraph.runs:
                                    if run.text.strip() and len(run.text.strip()) > 2:
                                        run.text = f'<<skill_{skill_counter}>>'
                                        logger.info(f"Replaced skill with <<skill_{skill_counter}>>")
                                        if f'<<skill_{skill_counter}>>' not in self.variables_created:
                                            self.variables_created.append(f'<<skill_{skill_counter}>>')
                                        skill_counter += 1
                                        break

                # Column 2: Experience
                elif col_idx == 1:
                    # Determine which job this is (1 or 2)
                    job_num = 1
                    if 'Senior' not in cell_text and 'PROSEWARE' not in cell_text:
                        job_num = 2

                    # Replace position and company
                    for old_text, new_text in experience_patterns:
                        if self.replace_in_cell(cell, old_text, new_text):
                            logger.info(f"Replaced '{old_text}' with '{new_text}' in Table 2")
                            replacements_made += 1
                            if new_text not in self.variables_created:
                                self.variables_created.append(new_text)

                    # Replace date ranges (Jan 20XX - Dec 20XX)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text
                            if '20XX' in text and '-' in text:
                                # This is a date range
                                run.text = f'<<start_date_{job_num}>> - <<end_date_{job_num}>>'
                                logger.info(f"Replaced date range with <<start_date_{job_num}>> - <<end_date_{job_num}>>")
                                if f'<<start_date_{job_num}>>' not in self.variables_created:
                                    self.variables_created.append(f'<<start_date_{job_num}>>')
                                    self.variables_created.append(f'<<end_date_{job_num}>>')

                    # Replace job descriptions (bullet points)
                    desc_counter = 1
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()
                        # Look for bullet points or description text
                        if text and len(text) > 30:  # Likely a description
                            for run in paragraph.runs:
                                if run.text.strip() and len(run.text.strip()) > 20:
                                    run.text = f'<<job_description_{job_num}_{desc_counter}>>'
                                    logger.info(f"Replaced job description with <<job_description_{job_num}_{desc_counter}>>")
                                    if f'<<job_description_{job_num}_{desc_counter}>>' not in self.variables_created:
                                        self.variables_created.append(f'<<job_description_{job_num}_{desc_counter}>>')
                                    desc_counter += 1
                                    break

        logger.info(f"Made {replacements_made} replacements in Table 2")

    def process_tables(self) -> None:
        """Process all tables in the document."""
        logger.info("Processing tables...")

        if len(self.doc.tables) >= 1:
            logger.info("Processing Table 1...")
            self.process_table_1_contact_info(self.doc.tables[0])

        if len(self.doc.tables) >= 2:
            logger.info("Processing Table 2...")
            self.process_table_2_experience(self.doc.tables[1])

        logger.info("Table processing complete")

    def save_document(self) -> None:
        """Save the converted document."""
        logger.info(f"Saving converted document to: {self.output_path}")
        self.doc.save(str(self.output_path))
        logger.info("Document saved successfully")

    def generate_report(self) -> str:
        """
        Generate a detailed conversion report.

        Returns:
            Formatted report string
        """
        report = []
        report.append("=" * 80)
        report.append("UI/UX DESIGNER TEMPLATE CONVERSION REPORT")
        report.append("=" * 80)
        report.append(f"\nInput File: {self.input_path}")
        report.append(f"Output File: {self.output_path}")
        report.append(f"\nTotal Variables Created: {len(self.variables_created)}")
        report.append("\nComplete Variable List:")
        report.append("-" * 80)

        # Group variables by category
        categories = {
            'Personal Information': [],
            'Contact Information': [],
            'Education': [],
            'Skills': [],
            'Experience': [],
            'Other': []
        }

        for var in sorted(self.variables_created):
            if 'first_name' in var or 'last_name' in var:
                categories['Personal Information'].append(var)
            elif 'email' in var or 'phone' in var or 'city' in var or 'website' in var:
                categories['Contact Information'].append(var)
            elif 'education' in var or 'degree' in var or 'major' in var or 'graduation' in var:
                categories['Education'].append(var)
            elif 'skill' in var:
                categories['Skills'].append(var)
            elif 'position' in var or 'company' in var or 'job_description' in var or 'date' in var:
                categories['Experience'].append(var)
            else:
                categories['Other'].append(var)

        for category, vars_list in categories.items():
            if vars_list:
                report.append(f"\n{category}:")
                for var in vars_list:
                    report.append(f"  {var}")

        report.append("\n" + "=" * 80)
        report.append("CONVERSION COMPLETE")
        report.append("=" * 80)

        return "\n".join(report)

    def convert(self) -> None:
        """Execute the complete conversion process."""
        logger.info("Starting UI/UX Designer template conversion...")

        try:
            # Load document
            self.load_document()

            # Process paragraphs (name)
            self.process_paragraphs()

            # Process tables
            self.process_tables()

            # Save converted document
            self.save_document()

            # Generate and print report
            report = self.generate_report()
            print("\n" + report)

            logger.info("Conversion completed successfully!")

        except Exception as e:
            logger.error(f"Conversion failed: {str(e)}", exc_info=True)
            raise


def main():
    """Main execution function."""
    # Define paths
    input_path = '/workspace/content_template_library/manual_converted/uiux_designer_template.docx'
    output_path = '/workspace/content_template_library/manual_converted/uiux_designer_template_converted.docx'

    # Create converter and run
    converter = UIUXTemplateConverter(input_path, output_path)
    converter.convert()


if __name__ == '__main__':
    main()
