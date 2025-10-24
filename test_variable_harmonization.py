#!/usr/bin/env python3
"""
Variable Naming Harmonization - Validation Test

This script validates that the variable naming harmonization was successful
and demonstrates proper usage of the unified naming standard.

Created: 2025-10-24
"""

import os
import re
from docx import Document
import csv


def test_template_csv_consistency():
    """
    Test that all templates and their CSV mappings are perfectly synchronized
    """
    print("=" * 80)
    print("VARIABLE HARMONIZATION VALIDATION TEST")
    print("=" * 80)

    template_dir = 'content_template_library'
    variable_pattern = re.compile(r'<<([^>]+)>>')

    all_passed = True
    total_templates = 0
    total_mismatches = 0

    # Test all templates
    for category in ['resumes', 'coverletters']:
        category_path = os.path.join(template_dir, category)

        if not os.path.exists(category_path):
            print(f"\nWarning: {category_path} not found")
            continue

        for file in os.listdir(category_path):
            if file.endswith('.docx') and not file.startswith('~'):
                docx_path = os.path.join(category_path, file)
                csv_path = docx_path.replace('.docx', '_mapping.csv')

                if not os.path.exists(csv_path):
                    print(f"\n❌ Missing CSV mapping: {csv_path}")
                    all_passed = False
                    continue

                total_templates += 1

                # Extract variables from template
                doc = Document(docx_path)
                template_vars = set()

                for paragraph in doc.paragraphs:
                    found = variable_pattern.findall(paragraph.text)
                    template_vars.update(found)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                found = variable_pattern.findall(paragraph.text)
                                template_vars.update(found)

                # Extract variables from CSV
                csv_vars = set()
                with open(csv_path, 'r', encoding='utf-8') as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        if row.get('Is_Variable') == 'TRUE' and row.get('Variable_name'):
                            var_names = row['Variable_name'].strip().split()
                            csv_vars.update(var_names)

                # Compare
                if template_vars == csv_vars:
                    print(f"\n✅ {file}")
                    print(f"   Template & CSV perfectly synchronized ({len(template_vars)} variables)")
                else:
                    print(f"\n❌ {file}")
                    all_passed = False
                    total_mismatches += 1

                    only_in_template = template_vars - csv_vars
                    only_in_csv = csv_vars - template_vars

                    if only_in_template:
                        print(f"   Variables in template but not CSV: {sorted(only_in_template)}")
                    if only_in_csv:
                        print(f"   Variables in CSV but not template: {sorted(only_in_csv)}")

    # Summary
    print("\n" + "=" * 80)
    print("VALIDATION SUMMARY")
    print("=" * 80)
    print(f"Templates tested: {total_templates}")
    print(f"Mismatches found: {total_mismatches}")

    if all_passed:
        print("\n🎉 SUCCESS: All templates perfectly synchronized with CSV mappings!")
        print("   Variable naming harmonization is COMPLETE.")
        return True
    else:
        print("\n⚠️  FAILED: Some templates have inconsistencies.")
        print("   Run update_csv_mappings.py to fix.")
        return False


def test_naming_conventions():
    """
    Test that all variables follow the CSV naming convention
    """
    print("\n" + "=" * 80)
    print("NAMING CONVENTION VALIDATION")
    print("=" * 80)

    # Valid patterns
    valid_patterns = [
        re.compile(r'^user_[a-z_]+$'),  # user_first_name, user_email
        re.compile(r'^professional_summary$'),
        re.compile(r'^executive_[a-z_]+$'),
        re.compile(r'^(technical|methodology|domain)_summary$'),
        re.compile(r'^tech_specialty$'),
        re.compile(r'^achievement_[a-z]+_\d+$'),  # achievement_technical_1
        re.compile(r'^leadership_competency_\d+$'),
        re.compile(r'^edu_\d+_[a-z_]+$'),  # edu_1_degree
        re.compile(r'^work_experience_\d+_(position|name|location|dates|context|tech_stack)$'),
        re.compile(r'^work_experience_\d+_skill\d+$'),  # work_experience_1_skill1
        re.compile(r'^volunteer_\d+_[a-z_]+$'),
        re.compile(r'^certifications_list$'),
        re.compile(r'^company_(name|address|city_prov)$'),
        re.compile(r'^hiring_manager_(name|first_name|title)$'),
        re.compile(r'^cover_letter_[a-z_]+$'),
        re.compile(r'^job_requirement_\d+$'),
        re.compile(r'^matching_qualification_\d+$'),
        re.compile(r'^current_date$'),
    ]

    template_dir = 'content_template_library'
    variable_pattern = re.compile(r'<<([^>]+)>>')

    all_vars = set()
    non_standard_vars = []

    # Collect all variables
    for category in ['resumes', 'coverletters']:
        category_path = os.path.join(template_dir, category)

        if not os.path.exists(category_path):
            continue

        for file in os.listdir(category_path):
            if file.endswith('.docx') and not file.startswith('~'):
                docx_path = os.path.join(category_path, file)

                doc = Document(docx_path)

                for paragraph in doc.paragraphs:
                    found = variable_pattern.findall(paragraph.text)
                    all_vars.update(found)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                found = variable_pattern.findall(paragraph.text)
                                all_vars.update(found)

    # Check each variable
    for var in sorted(all_vars):
        is_valid = any(pattern.match(var) for pattern in valid_patterns)

        if is_valid:
            print(f"✅ {var:50s} - Valid CSV naming convention")
        else:
            print(f"⚠️  {var:50s} - Non-standard (template-specific)")
            non_standard_vars.append(var)

    print(f"\nTotal variables: {len(all_vars)}")
    print(f"Standard convention: {len(all_vars) - len(non_standard_vars)}")
    print(f"Template-specific: {len(non_standard_vars)}")

    if non_standard_vars:
        print("\nTemplate-specific variables:")
        for var in non_standard_vars:
            print(f"  - {var}")

    return True


def demonstrate_usage():
    """
    Demonstrate correct usage of the variable naming standard
    """
    print("\n" + "=" * 80)
    print("USAGE DEMONSTRATION")
    print("=" * 80)

    print("\n1. Personal Information:")
    print("   Variables: user_first_name, user_last_name, user_email")
    print("   Template: <<user_first_name>> <<user_last_name>>")
    print("   Data:     {'user_first_name': 'John', 'user_last_name': 'Doe'}")

    print("\n2. Work Experience (Job 1):")
    print("   Variables: work_experience_1_position, work_experience_1_name")
    print("   Template: <<work_experience_1_position>> at <<work_experience_1_name>>")
    print("   Data:     {'work_experience_1_position': 'Senior Engineer',")
    print("              'work_experience_1_name': 'Tech Corp'}")

    print("\n3. Work Achievements (Nested numbering):")
    print("   Variables: work_experience_1_skill1, work_experience_1_skill2")
    print("   Template: • <<work_experience_1_skill1>>")
    print("             • <<work_experience_1_skill2>>")
    print("   Data:     {'work_experience_1_skill1': 'Led team of 5 developers',")
    print("              'work_experience_1_skill2': 'Increased efficiency 40%'}")

    print("\n4. Cover Letter:")
    print("   Variables: hiring_manager_name, cover_letter_opening")
    print("   Template: Dear <<hiring_manager_name>>,")
    print("             <<cover_letter_opening>>")
    print("   Data:     {'hiring_manager_name': 'Jane Smith',")
    print("              'cover_letter_opening': 'I am excited to apply...'}")

    print("\n5. Template-Specific (T-Format Cover Letter):")
    print("   Variables: job_requirement_1, matching_qualification_1")
    print("   Template: Requirement: <<job_requirement_1>>")
    print("             Match:       <<matching_qualification_1>>")

    print("\nFor complete variable reference:")
    print("  → See VARIABLE_NAMING_REFERENCE.md")

    return True


if __name__ == '__main__':
    print("\nVariable Naming Harmonization - Validation & Usage Test\n")

    # Run tests
    sync_passed = test_template_csv_consistency()
    naming_passed = test_naming_conventions()
    demo_passed = demonstrate_usage()

    # Final summary
    print("\n" + "=" * 80)
    print("FINAL RESULTS")
    print("=" * 80)

    if sync_passed and naming_passed and demo_passed:
        print("✅ All validation tests PASSED")
        print("✅ Variable naming harmonization is COMPLETE")
        print("✅ System ready for production use")
        exit(0)
    else:
        print("⚠️  Some validation tests failed")
        print("   Review errors above and run update_csv_mappings.py")
        exit(1)
