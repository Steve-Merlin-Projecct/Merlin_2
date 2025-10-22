#!/usr/bin/env python3
"""
Test Steve Glen data insertion into Restaurant Manager template
"""

import json
import sys
import re
sys.path.insert(0, './scripts')

from simplified_inserter import SimplifiedInserter
from docx import Document

def map_steve_glen_to_template(json_path):
    """Map Steve Glen's data to template variables"""

    with open(json_path, 'r') as f:
        data = json.load(f)

    resume = data['resume_data']
    personal = resume['personal']

    # Create mapping
    mapped_data = {}

    # CONTACT - Direct mapping
    mapped_data['first_name'] = personal.get('first_name', '')
    mapped_data['last_name'] = personal.get('last_name', '')
    mapped_data['email'] = personal.get('email', '')
    mapped_data['phone'] = personal.get('phone', '')
    mapped_data['city'] = personal.get('city', '')
    mapped_data['state'] = personal.get('province', '')  # province -> state
    mapped_data['street_address'] = personal.get('address', '')
    mapped_data['zip_code'] = personal.get('postal_code', '')
    mapped_data['linkedin_url'] = personal.get('linkedin', '')

    # CAREER OVERVIEW - Need to split professional summary
    # Steve has one paragraph, template needs 5 statements
    prof_summary = resume.get('professional_summary', '')
    # For now, just put the whole thing in career_overview_1
    mapped_data['career_overview_1'] = prof_summary
    mapped_data['career_overview_2'] = ''  # Missing
    mapped_data['career_overview_3'] = ''  # Missing
    mapped_data['career_overview_4'] = ''  # Missing
    mapped_data['career_overview_5'] = ''  # Missing

    # JOB EXPERIENCE - Job 1
    if resume.get('experience') and len(resume['experience']) > 0:
        job1 = resume['experience'][0]
        mapped_data['position_1'] = job1.get('position', '')
        # Note: company_1 not in template variables!
        # Note: start_date_1 not in template variables!
        mapped_data['end_date_1'] = job1.get('end_date', '')

        # Experience bullets
        bullets = job1.get('bullets', [])
        mapped_data['job_1_experience_1'] = bullets[0] if len(bullets) > 0 else ''
        mapped_data['job_1_experience_2'] = bullets[1] if len(bullets) > 1 else ''
        mapped_data['job_1_experience_3'] = bullets[2] if len(bullets) > 2 else ''

    # JOB 2 - Steve only has 1 job, so these will be empty
    mapped_data['position_2'] = ''
    mapped_data['company_2'] = ''
    mapped_data['start_date_2'] = ''
    mapped_data['end_date_2'] = ''
    mapped_data['job_2_experience_1'] = ''
    mapped_data['job_2_experience_3'] = ''  # Note: job_2_experience_2 not in template!
    mapped_data['job_2_experience_4'] = ''

    # EDUCATION
    if resume.get('education') and len(resume['education']) > 0:
        edu1 = resume['education'][0]
        mapped_data['institution_1'] = edu1.get('institution', '')

        # Parse location "Edmonton, AB" into city and state
        location = edu1.get('location', '')
        if ',' in location:
            parts = location.split(',')
            mapped_data['institution_city_1'] = parts[0].strip()
            mapped_data['institution_state_1'] = parts[1].strip() if len(parts) > 1 else ''
        else:
            mapped_data['institution_city_1'] = ''
            mapped_data['institution_state_1'] = ''

        # Note: degree_1 and graduation_date_1 not in template variables!

    # SKILLS - Need to extract from categorized format
    skills_data = resume.get('skills', {})

    # Steve has skills in categories, template needs 6 individual items
    # Let's extract some skills from different categories
    all_skills = []

    # Extract from digital_marketing
    if 'digital_marketing' in skills_data:
        dm_skills = skills_data['digital_marketing'].split(',')
        all_skills.extend([s.strip() for s in dm_skills[:2]])

    # Extract from technical_expertise
    if 'technical_expertise' in skills_data:
        tech_skills = skills_data['technical_expertise'].split(',')
        all_skills.extend([s.strip() for s in tech_skills[:2]])

    # Extract from business_analytics
    if 'business_analytics' in skills_data:
        ba_skills = skills_data['business_analytics'].split(',')
        all_skills.extend([s.strip() for s in ba_skills[:2]])

    # Map to skill_1 through skill_6
    for i in range(1, 7):
        mapped_data[f'skill_{i}'] = all_skills[i-1] if i-1 < len(all_skills) else ''

    # INTERESTS - Split comma-separated string
    interests_str = skills_data.get('interests', '')
    if interests_str:
        interests_list = [i.strip() for i in interests_str.split(',')]
        for i in range(1, 7):
            mapped_data[f'interest_{i}'] = interests_list[i-1] if i-1 < len(interests_list) else ''
    else:
        for i in range(1, 7):
            mapped_data[f'interest_{i}'] = ''

    # Headers (use defaults)
    mapped_data['profile_header'] = 'Profile'
    mapped_data['experience_header'] = 'Experience'
    mapped_data['education_header'] = 'Education'
    mapped_data['skills_header'] = 'Skills'
    mapped_data['interests_header'] = 'Interests'

    return mapped_data


def test_insertion():
    """Test the insertion process"""

    print("="*60)
    print("STEVE GLEN DATA INSERTION TEST")
    print("="*60)
    print()

    # Map Steve's data
    print("Step 1: Mapping Steve Glen data to template variables...")
    steve_data = map_steve_glen_to_template(
        '/workspace/content_template_library/steve_glen_comprehensive_defaults.json'
    )

    # Show what we mapped
    print(f"  Mapped {len([v for v in steve_data.values() if v])} non-empty variables")
    print()

    # Insert into template
    print("Step 2: Inserting data into template...")
    inserter = SimplifiedInserter()

    output_path = inserter.populate_template(
        template_name='restaurant_manager',
        data=steve_data,
        output_filename='steve_glen_restaurant_manager_test.docx'
    )

    print(f"  Output: {output_path}")
    print()

    # Analyze result
    print("Step 3: Analyzing output document...")
    doc = Document(output_path)

    # Check for remaining placeholders
    all_text = []
    for p in doc.paragraphs:
        all_text.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                all_text.append(c.text)

    # Find remaining variables
    remaining = set()
    for text in all_text:
        found = re.findall(r'<<([^>]+)>>', text)
        remaining.update(found)

    print(f"  Remaining placeholders: {len(remaining)}")
    if remaining:
        print("  Variables NOT replaced:")
        for var in sorted(remaining):
            print(f"    - <<{var}>>")

    print()
    print("="*60)
    print("Step 4: Summary Report")
    print("="*60)
    print()

    total_vars = 43  # Known from template analysis
    replaced = total_vars - len(remaining)

    print(f"Total template variables: {total_vars}")
    print(f"Successfully replaced: {replaced}")
    print(f"Still placeholders: {len(remaining)}")
    print(f"Success rate: {replaced/total_vars*100:.1f}%")
    print()

    # Show what DID get inserted
    print("Sample of successfully inserted data:")
    non_empty = {k: v for k, v in steve_data.items() if v and k not in remaining}
    for k, v in list(non_empty.items())[:10]:
        print(f"  {k}: {v[:50]}..." if len(v) > 50 else f"  {k}: {v}")

    return {
        'total': total_vars,
        'replaced': replaced,
        'remaining': len(remaining),
        'remaining_vars': sorted(remaining),
        'output_path': output_path
    }


if __name__ == "__main__":
    results = test_insertion()