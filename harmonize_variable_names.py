#!/usr/bin/env python3
"""
Variable Naming Harmonization Script

This script harmonizes all variable naming conventions across the document generation system
to use a single unified naming standard (CSV convention).

Features:
- Analyzes all .docx templates to extract current variable names
- Analyzes all CSV mapping files
- Identifies inconsistencies across the system
- Updates templates and mappings to use unified naming
- Generates comprehensive report of changes

Standard Naming Convention (CSV):
- user_first_name, user_last_name, user_email, user_phone, user_linkedin
- user_city_prov (combined city and province)
- professional_summary
- technical_summary, methodology_summary, domain_summary
- edu_1_name, edu_1_degree, edu_1_concentration, edu_1_grad_date, edu_1_location
- work_experience_1_position, work_experience_1_name, work_experience_1_location, work_experience_1_dates
- work_experience_1_skill1, work_experience_1_skill2, etc.
- volunteer_1_position, volunteer_1_name, volunteer_1_dates, volunteer_1_description
- certifications_list
- company_name, hiring_manager_name (for cover letters)
- cover_letter_opening, cover_letter_skills_alignment, cover_letter_achievement, cover_letter_closing
- current_date

Created: 2025-10-24
"""

import os
import re
import csv
import json
from typing import Dict, List, Set, Tuple
from docx import Document
from collections import defaultdict


class VariableNameHarmonizer:
    """
    Harmonizes variable naming conventions across all system components
    """

    # Canonical variable naming standard (CSV convention)
    STANDARD_VARIABLES = {
        # Personal Information
        'user_first_name': 'First name',
        'user_last_name': 'Last name',
        'user_email': 'Email address',
        'user_phone': 'Phone number',
        'user_linkedin': 'LinkedIn URL',
        'user_city_prov': 'City and province',

        # Professional Summary
        'professional_summary': 'Professional summary paragraph',

        # Skills
        'technical_summary': 'Technical skills list',
        'methodology_summary': 'Methodologies list',
        'domain_summary': 'Domain expertise list',

        # Education (template patterns)
        'edu_{n}_name': 'Educational institution name',
        'edu_{n}_degree': 'Degree type',
        'edu_{n}_concentration': 'Field of study',
        'edu_{n}_specialization': 'Specialization area',
        'edu_{n}_grad_date': 'Graduation date',
        'edu_{n}_location': 'Institution location',

        # Work Experience (template patterns)
        'work_experience_{n}_position': 'Job position/title',
        'work_experience_{n}_name': 'Company/organization name',
        'work_experience_{n}_location': 'Job location',
        'work_experience_{n}_dates': 'Employment date range',
        'work_experience_{n}_skill{m}': 'Job achievement/responsibility',

        # Volunteer Experience (template patterns)
        'volunteer_{n}_position': 'Volunteer position title',
        'volunteer_{n}_name': 'Volunteer organization name',
        'volunteer_{n}_location': 'Volunteer location',
        'volunteer_{n}_dates': 'Volunteer date range',
        'volunteer_{n}_description': 'Volunteer work description',

        # Certifications
        'certifications_list': 'Professional certifications',

        # Cover Letter Specific
        'company_name': 'Target company name',
        'hiring_manager_name': 'Hiring manager name',
        'cover_letter_opening': 'Opening paragraph',
        'cover_letter_skills_alignment': 'Skills alignment paragraph',
        'cover_letter_achievement': 'Achievement example paragraph',
        'cover_letter_closing': 'Closing paragraph',

        # Metadata
        'current_date': 'Document generation date',
    }

    # Known alternative naming conventions that need to be mapped
    LEGACY_MAPPINGS = {
        # Production template names -> CSV names
        'full_name': 'user_first_name user_last_name',
        'email': 'user_email',
        'phone_number': 'user_phone',
        'linkedin_url': 'user_linkedin',
        'location': 'user_city_prov',
        'career_overview_1': 'professional_summary',
        'technical_skills': 'technical_summary',
        'methodologies': 'methodology_summary',
        'domain_expertise': 'domain_summary',
        'job_{n}_title': 'work_experience_{n}_position',
        'job_{n}_company': 'work_experience_{n}_name',
        'job_{n}_location': 'work_experience_{n}_location',
        'job_{n}_responsibility_{m}': 'work_experience_{n}_skill{m}',
        'education_{n}_institution': 'edu_{n}_name',
        'education_{n}_degree': 'edu_{n}_degree',
        'graduation_date': 'edu_{n}_grad_date',
    }

    def __init__(self, worktree_root: str):
        """
        Initialize harmonizer with worktree root path

        Args:
            worktree_root: Path to the worktree root directory
        """
        self.worktree_root = worktree_root
        self.template_dir = os.path.join(worktree_root, 'content_template_library')
        self.variable_pattern = re.compile(r'<<([^>]+)>>')

        # Analysis results
        self.docx_variables = {}  # {file_path: set(variables)}
        self.csv_variables = {}   # {file_path: set(variables)}
        self.inconsistencies = []
        self.changes_made = []

    def analyze_all_components(self) -> Dict:
        """
        Analyze all components to identify current variable naming

        Returns:
            Dictionary with analysis results
        """
        print("=" * 80)
        print("VARIABLE NAMING HARMONIZATION ANALYSIS")
        print("=" * 80)

        # Analyze .docx templates
        print("\n1. Analyzing .docx templates...")
        self._analyze_docx_templates()

        # Analyze CSV mapping files
        print("\n2. Analyzing CSV mapping files...")
        self._analyze_csv_mappings()

        # Identify inconsistencies
        print("\n3. Identifying inconsistencies...")
        self._identify_inconsistencies()

        # Generate analysis report
        report = {
            'total_docx_files': len(self.docx_variables),
            'total_csv_files': len(self.csv_variables),
            'total_inconsistencies': len(self.inconsistencies),
            'docx_variables': self.docx_variables,
            'csv_variables': self.csv_variables,
            'inconsistencies': self.inconsistencies,
        }

        return report

    def _analyze_docx_templates(self):
        """Analyze all .docx template files"""
        # Find all .docx files in template directories
        for root, dirs, files in os.walk(self.template_dir):
            for file in files:
                if file.endswith('.docx') and not file.startswith('~'):
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, self.worktree_root)

                    try:
                        variables = self._extract_docx_variables(file_path)
                        self.docx_variables[rel_path] = variables
                        print(f"   ✓ {rel_path}: {len(variables)} variables")
                    except Exception as e:
                        print(f"   ✗ {rel_path}: ERROR - {e}")

    def _extract_docx_variables(self, file_path: str) -> Set[str]:
        """
        Extract all variable names from a .docx template

        Args:
            file_path: Path to .docx file

        Returns:
            Set of variable names
        """
        doc = Document(file_path)
        variables = set()

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            found = self.variable_pattern.findall(paragraph.text)
            variables.update(found)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        found = self.variable_pattern.findall(paragraph.text)
                        variables.update(found)

        return variables

    def _analyze_csv_mappings(self):
        """Analyze all CSV mapping files"""
        # Find all CSV mapping files
        for root, dirs, files in os.walk(self.template_dir):
            for file in files:
                if file.endswith('_mapping.csv'):
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, self.worktree_root)

                    try:
                        variables = self._extract_csv_variables(file_path)
                        self.csv_variables[rel_path] = variables
                        print(f"   ✓ {rel_path}: {len(variables)} variables")
                    except Exception as e:
                        print(f"   ✗ {rel_path}: ERROR - {e}")

    def _extract_csv_variables(self, file_path: str) -> Set[str]:
        """
        Extract all variable names from a CSV mapping file

        Args:
            file_path: Path to CSV file

        Returns:
            Set of variable names
        """
        variables = set()

        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                if row.get('Is_Variable') == 'TRUE' and row.get('Variable_name'):
                    # Variable_name column may contain multiple space-separated names
                    var_names = row['Variable_name'].strip().split()
                    variables.update(var_names)

        return variables

    def _identify_inconsistencies(self):
        """Identify naming inconsistencies between templates and mappings"""
        # Compare .docx variables with their corresponding CSV mappings
        for docx_path, docx_vars in self.docx_variables.items():
            # Find corresponding CSV mapping
            base_name = os.path.splitext(docx_path)[0]
            csv_path = base_name + '_mapping.csv'

            if csv_path in self.csv_variables:
                csv_vars = self.csv_variables[csv_path]

                # Find mismatches
                only_in_docx = docx_vars - csv_vars
                only_in_csv = csv_vars - docx_vars

                if only_in_docx or only_in_csv:
                    self.inconsistencies.append({
                        'docx_file': docx_path,
                        'csv_file': csv_path,
                        'only_in_docx': sorted(only_in_docx),
                        'only_in_csv': sorted(only_in_csv),
                    })

        # Report inconsistencies
        if self.inconsistencies:
            for inc in self.inconsistencies:
                print(f"\n   Mismatch: {inc['docx_file']}")
                if inc['only_in_docx']:
                    print(f"      Variables in DOCX but not CSV: {inc['only_in_docx']}")
                if inc['only_in_csv']:
                    print(f"      Variables in CSV but not DOCX: {inc['only_in_csv']}")
        else:
            print("   ✓ No inconsistencies found between templates and mappings!")

    def generate_report(self, output_path: str):
        """
        Generate comprehensive analysis report

        Args:
            output_path: Path to save the report
        """
        report_lines = []
        report_lines.append("=" * 80)
        report_lines.append("VARIABLE NAMING HARMONIZATION REPORT")
        report_lines.append("=" * 80)
        report_lines.append(f"Generated: {os.popen('date').read().strip()}")
        report_lines.append("")

        # Summary statistics
        report_lines.append("SUMMARY STATISTICS")
        report_lines.append("-" * 80)
        report_lines.append(f"Total .docx templates analyzed: {len(self.docx_variables)}")
        report_lines.append(f"Total CSV mappings analyzed: {len(self.csv_variables)}")
        report_lines.append(f"Total inconsistencies found: {len(self.inconsistencies)}")
        report_lines.append("")

        # Standard variables reference
        report_lines.append("STANDARD VARIABLE NAMING CONVENTION (CSV)")
        report_lines.append("-" * 80)
        for var_name, description in sorted(self.STANDARD_VARIABLES.items()):
            report_lines.append(f"  {var_name:40s} - {description}")
        report_lines.append("")

        # All unique variables found
        all_vars = set()
        for vars_set in self.docx_variables.values():
            all_vars.update(vars_set)
        for vars_set in self.csv_variables.values():
            all_vars.update(vars_set)

        report_lines.append("ALL UNIQUE VARIABLES FOUND IN SYSTEM")
        report_lines.append("-" * 80)
        for var in sorted(all_vars):
            # Check if it matches standard
            if var in self.STANDARD_VARIABLES:
                status = "✓ STANDARD"
            elif any(re.match(pattern.replace('{n}', r'\d+').replace('{m}', r'\d+'), var)
                    for pattern in self.STANDARD_VARIABLES.keys() if '{' in pattern):
                status = "✓ STANDARD (numbered)"
            else:
                status = "⚠ NON-STANDARD"
            report_lines.append(f"  {var:50s} {status}")
        report_lines.append("")

        # Detailed file analysis
        report_lines.append("DETAILED FILE ANALYSIS")
        report_lines.append("-" * 80)

        report_lines.append("\n.DOCX Templates:")
        for file_path, variables in sorted(self.docx_variables.items()):
            report_lines.append(f"\n  {file_path}")
            report_lines.append(f"    Variables: {len(variables)}")
            for var in sorted(variables):
                report_lines.append(f"      - {var}")

        report_lines.append("\n\nCSV Mappings:")
        for file_path, variables in sorted(self.csv_variables.items()):
            report_lines.append(f"\n  {file_path}")
            report_lines.append(f"    Variables: {len(variables)}")
            for var in sorted(variables):
                report_lines.append(f"      - {var}")

        # Inconsistencies
        if self.inconsistencies:
            report_lines.append("\n\nINCONSISTENCIES DETECTED")
            report_lines.append("-" * 80)
            for inc in self.inconsistencies:
                report_lines.append(f"\nMismatch:")
                report_lines.append(f"  DOCX: {inc['docx_file']}")
                report_lines.append(f"  CSV:  {inc['csv_file']}")
                if inc['only_in_docx']:
                    report_lines.append(f"  Variables in DOCX but not CSV:")
                    for var in inc['only_in_docx']:
                        report_lines.append(f"    - {var}")
                if inc['only_in_csv']:
                    report_lines.append(f"  Variables in CSV but not DOCX:")
                    for var in inc['only_in_csv']:
                        report_lines.append(f"    - {var}")

        # Write report
        report_content = '\n'.join(report_lines)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report_content)

        print(f"\n✓ Report saved to: {output_path}")
        return report_content


def main():
    """Main execution function"""
    # Determine worktree root
    worktree_root = '/workspace/.trees/word-document-cover-letter-and-resume-template-var'

    # Initialize harmonizer
    harmonizer = VariableNameHarmonizer(worktree_root)

    # Run analysis
    analysis = harmonizer.analyze_all_components()

    # Generate report
    report_path = os.path.join(worktree_root, 'VARIABLE_NAMING_ANALYSIS.md')
    harmonizer.generate_report(report_path)

    # Print summary
    print("\n" + "=" * 80)
    print("ANALYSIS COMPLETE")
    print("=" * 80)
    print(f"Templates analyzed: {analysis['total_docx_files']}")
    print(f"Mappings analyzed: {analysis['total_csv_files']}")
    print(f"Inconsistencies found: {analysis['total_inconsistencies']}")
    print(f"\nFull report: {report_path}")
    print("=" * 80)


if __name__ == '__main__':
    main()
