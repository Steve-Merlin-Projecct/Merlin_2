#!/usr/bin/env python3
"""
Generate Sample DOCX Documents with Authenticity Enhancement

This script creates real DOCX documents to verify that all authenticity
features work correctly in actual document generation.

Author: Automated Job Application System
Date: October 10, 2025
"""

import sys
import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from modules.content.document_generation import (
    SmartTypography,
    MetadataGenerator,
    AuthenticityValidator,
    validate_document_authenticity,
)


def create_sample_document_1():
    """Create sample resume document with authenticity features"""
    print("\n" + "=" * 70)
    print("SAMPLE 1: Professional Resume")
    print("=" * 70)

    # Create document
    doc = Document()

    # Add content with various typography needs
    doc.add_heading('Jane Smith', 0)
    doc.add_paragraph('Software Engineer -- Building the Future')

    # Contact info with smart typography
    contact = doc.add_paragraph()
    contact.add_run('Email: jane.smith@email.com | Phone: (555) 123-4567\n')
    contact.add_run('LinkedIn: linkedin.com/in/janesmith')

    # Professional summary
    doc.add_heading('Professional Summary', 1)
    summary = doc.add_paragraph(
        'I\'m a dedicated software engineer with 5-10 years of experience '
        'in Python, JavaScript, and cloud technologies. I believe "quality '
        'code speaks louder than words"... My passion is creating elegant '
        'solutions to complex problems.'
    )

    # Experience
    doc.add_heading('Experience', 1)
    doc.add_paragraph(
        'Senior Software Engineer (2020-2023)\n'
        'Tech Corp -- Leading Innovation\n'
        '- Developed microservices architecture handling 10 MB - 100 MB payloads\n'
        '- Improved system performance by 50%\n'
        '- Mentored team of 5-8 engineers'
    )

    # Education
    doc.add_heading('Education', 1)
    doc.add_paragraph('BS Computer Science, MIT\nGPA: 3.75/4.0')

    # Apply smart typography
    print("\n1. Applying Smart Typography...")
    typography = SmartTypography()

    for paragraph in doc.paragraphs:
        if paragraph.text:
            enhanced_text, stats = typography.apply_all(paragraph.text)
            # Replace paragraph text
            for run in paragraph.runs:
                run.text = ''
            paragraph.add_run(enhanced_text)

    print(f"   ✓ Typography enhanced")

    # Generate realistic metadata
    print("\n2. Generating Realistic Metadata...")
    generator = MetadataGenerator()
    metadata = generator.generate_metadata_for_doc_object(
        doc=doc,
        document_type='resume',
        author_name='Jane Smith',
        document_title='Jane Smith Resume',
        subject='Professional Software Engineer Resume',
        keywords='Software Engineer, Python, JavaScript, Resume',
    )

    created = doc.core_properties.created
    modified = doc.core_properties.modified
    print(f"   Created: {created.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"   Modified: {modified.strftime('%Y-%m-%d %H:%M:%S')}")

    # Save document
    output_path = '/workspace/tests/output/sample_resume_jane_smith.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\n3. Document Saved: {output_path}")
    print(f"   File Size: {os.path.getsize(output_path)} bytes")

    # Validate authenticity
    print("\n4. Validating Authenticity...")

    # Extract text for validation
    text_content = '\n'.join([p.text for p in doc.paragraphs])

    results = validate_document_authenticity(
        metadata=metadata,
        text_content=text_content,
        file_path=output_path,
    )

    print(f"   Score: {results['total_score']}/100 ({results['level'].upper()})")
    print(f"   Status: {'✓ PASSED' if results['passed'] else '✗ FAILED'}")

    return results


def create_sample_document_2():
    """Create sample cover letter with authenticity features"""
    print("\n\n" + "=" * 70)
    print("SAMPLE 2: Professional Cover Letter")
    print("=" * 70)

    # Create document
    doc = Document()

    # Header
    doc.add_paragraph('John Doe\n555 Tech Street\nSan Francisco, CA 94102')
    doc.add_paragraph(f'{datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('\n')

    # Recipient
    doc.add_paragraph('Hiring Manager\nTech Innovations Inc.\n123 Innovation Drive')
    doc.add_paragraph('\n')

    # Salutation
    doc.add_paragraph('Dear Hiring Manager,')

    # Body with typography features
    doc.add_paragraph(
        'I\'m writing to express my strong interest in the Senior Python Developer '
        'position (2020-2023 opening). With a proven track record of delivering '
        'high-quality software -- and a passion for continuous learning -- I believe '
        'I would be an excellent fit for your team.'
    )

    doc.add_paragraph(
        'In my previous role, I led projects that improved system efficiency by 75%, '
        'managed databases processing 100 MB - 1 GB daily, and mentored 3-5 junior '
        'developers. I believe "great code is readable code"...'
    )

    doc.add_paragraph(
        'I hold a BS in Computer Science from Stanford and have certifications in '
        'Python, AWS, and Docker. My skills align perfectly with your requirements.'
    )

    # Closing
    doc.add_paragraph('Thank you for considering my application. I look forward to discussing how my experience and passion can contribute to Tech Innovations Inc.')
    doc.add_paragraph('\nSincerely,\nJohn Doe')

    # Apply smart typography
    print("\n1. Applying Smart Typography...")
    typography = SmartTypography()

    enhanced_paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            enhanced_text, _ = typography.apply_all(paragraph.text)
            enhanced_paragraphs.append(enhanced_text)
            # Replace paragraph text
            for run in paragraph.runs:
                run.text = ''
            paragraph.add_run(enhanced_text)

    print(f"   ✓ Typography enhanced on {len(enhanced_paragraphs)} paragraphs")

    # Generate realistic metadata
    print("\n2. Generating Realistic Metadata...")
    generator = MetadataGenerator()
    metadata = generator.generate_metadata_for_doc_object(
        doc=doc,
        document_type='coverletter',
        author_name='John Doe',
        document_title='John Doe Cover Letter - Tech Innovations Inc',
        subject='Cover Letter for Senior Python Developer Position',
        keywords='Cover Letter, Python Developer, Software Engineer',
    )

    created = doc.core_properties.created
    modified = doc.core_properties.modified
    print(f"   Created: {created.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"   Modified: {modified.strftime('%Y-%m-%d %H:%M:%S')}")

    # Save document
    output_path = '/workspace/tests/output/sample_coverletter_john_doe.docx'
    doc.save(output_path)
    print(f"\n3. Document Saved: {output_path}")
    print(f"   File Size: {os.path.getsize(output_path)} bytes")

    # Validate authenticity
    print("\n4. Validating Authenticity...")
    text_content = '\n'.join([p.text for p in doc.paragraphs])

    results = validate_document_authenticity(
        metadata=metadata,
        text_content=text_content,
        file_path=output_path,
    )

    print(f"   Score: {results['total_score']}/100 ({results['level'].upper()})")
    print(f"   Status: {'✓ PASSED' if results['passed'] else '✗ FAILED'}")

    return results


def create_sample_document_3():
    """Create document testing all special characters and typography"""
    print("\n\n" + "=" * 70)
    print("SAMPLE 3: Typography Feature Test Document")
    print("=" * 70)

    # Create document
    doc = Document()

    doc.add_heading('Typography Enhancement Test Document', 0)

    # Smart Quotes
    doc.add_heading('1. Smart Quotes', 1)
    doc.add_paragraph('"Double quotes" should become curly quotes.')
    doc.add_paragraph("'Single quotes' and apostrophes like it's, don't, can't.")

    # Smart Dashes
    doc.add_heading('2. Smart Dashes', 1)
    doc.add_paragraph('Date ranges: 2020-2023, January-December')
    doc.add_paragraph('Em dashes -- for breaks in thought -- like this.')
    doc.add_paragraph('En dash for ranges: pages 10-20, years 2015-2020')

    # Ellipsis
    doc.add_heading('3. Ellipsis', 1)
    doc.add_paragraph('Three dots... should become ellipsis character.')

    # Non-breaking spaces
    doc.add_heading('4. Non-Breaking Spaces', 1)
    doc.add_paragraph('Dr. Smith and Mr. Johnson met with Prof. Anderson.')
    doc.add_paragraph('File sizes: 10 MB, 500 KB, 2 GB')
    doc.add_paragraph('Measurements: 50%, 25 kg, 100 cm')

    # Special characters
    doc.add_heading('5. Special Characters', 1)
    doc.add_paragraph('Copyright (C) 2025, Trademark (TM), Registered (R)')
    doc.add_paragraph('Temperature: 72 degrees Fahrenheit')
    doc.add_paragraph('Fractions: 1/2 cup, 1/4 teaspoon, 3/4 pound')
    doc.add_paragraph('Multiplication: 10 x 20 inches')

    # Apply smart typography
    print("\n1. Applying Smart Typography...")
    typography = SmartTypography()

    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.style.name.startswith('Normal'):
            original = paragraph.text
            enhanced_text, stats = typography.apply_all(paragraph.text)

            # Show transformation
            if original != enhanced_text:
                print(f"\n   Original:  {repr(original)}")
                print(f"   Enhanced:  {repr(enhanced_text)}")

            # Replace paragraph text
            for run in paragraph.runs:
                run.text = ''
            paragraph.add_run(enhanced_text)

    # Generate metadata
    print("\n\n2. Generating Realistic Metadata...")
    generator = MetadataGenerator()
    metadata = generator.generate_metadata_for_doc_object(
        doc=doc,
        document_type='template',
        author_name='Test User',
        document_title='Typography Test Document',
        subject='Testing Smart Typography Features',
        keywords='Typography, Test, Smart Quotes, Dashes, Special Characters',
    )

    # Save
    output_path = '/workspace/tests/output/sample_typography_test.docx'
    doc.save(output_path)
    print(f"\n3. Document Saved: {output_path}")

    # Validate
    print("\n4. Validating Authenticity...")
    text_content = '\n'.join([p.text for p in doc.paragraphs])

    results = validate_document_authenticity(
        metadata=metadata,
        text_content=text_content,
        file_path=output_path,
    )

    print(f"   Score: {results['total_score']}/100 ({results['level'].upper()})")
    print(f"   Status: {'✓ PASSED' if results['passed'] else '✗ FAILED'}")

    return results


def main():
    """Generate all sample documents"""
    print("\n" + "=" * 70)
    print("GENERATING SAMPLE DOCX DOCUMENTS WITH AUTHENTICITY ENHANCEMENT")
    print("=" * 70)
    print(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Output Directory: /workspace/tests/output/")

    try:
        # Create samples
        result1 = create_sample_document_1()
        result2 = create_sample_document_2()
        result3 = create_sample_document_3()

        # Summary
        print("\n\n" + "=" * 70)
        print("GENERATION COMPLETE ✓")
        print("=" * 70)

        avg_score = (result1['total_score'] + result2['total_score'] + result3['total_score']) / 3

        print(f"\nSample Documents Created: 3")
        print(f"Average Authenticity Score: {avg_score:.1f}/100")
        print(f"\nDocument Scores:")
        print(f"  1. Resume:              {result1['total_score']}/100 ({result1['level'].upper()})")
        print(f"  2. Cover Letter:        {result2['total_score']}/100 ({result2['level'].upper()})")
        print(f"  3. Typography Test:     {result3['total_score']}/100 ({result3['level'].upper()})")

        print(f"\nAll documents saved to: /workspace/tests/output/")
        print(f"\nNext Steps:")
        print(f"  1. Open documents in Microsoft Word")
        print(f"  2. Inspect File > Info > Properties")
        print(f"  3. Verify smart typography (quotes, dashes, etc.)")
        print(f"  4. Check metadata timestamps are realistic")

        return 0

    except Exception as e:
        print(f"\n✗ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
