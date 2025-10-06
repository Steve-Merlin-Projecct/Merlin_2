"""
Template Converter System for Document Generation

This module converts reference .docx files into template files with variable placeholders.
It uses pattern recognition to identify common resume/cover letter fields and replaces
them with variable placeholders for dynamic content generation.

Key Features:
- Preserves all formatting, styles, and document structure
- Intelligent pattern recognition for common fields
- Fallback system for unidentified content
- Supports both resume and cover letter templates
"""

import os
import re
import json
import logging
from datetime import datetime
from docx import Document


class TemplateConverter:
    """
    Converts reference .docx files into template files with variable placeholders

    This class handles the conversion process by:
    1. Reading reference .docx files
    2. Identifying common patterns (names, dates, emails, etc.)
    3. Replacing content with variable placeholders
    4. Preserving all formatting and document structure
    5. Saving template files for document generation
    """

    def __init__(self):
        """Initialize the template converter with pattern recognition rules"""
        self.setup_logging()
        self.define_patterns()
        self.field_counter = 1  # For generic field numbering

    def setup_logging(self):
        """Configure logging for template conversion process"""
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    def define_patterns(self):
        """
        Define pattern recognition rules for common resume/cover letter fields

        These patterns identify typical content that should be replaced with variables:
        - Personal information (names, contact details)
        - Dates and time periods
        - Educational institutions and degrees
        - Job titles and organizations
        - Skills and qualifications
        """
        self.patterns = {
            # Personal Information Patterns
            "full_name": {
                "pattern": r"^[A-Z][a-z]+ [A-Z][a-z]+$",
                "variable": "<<first_name>> <<last_name>>",
                "description": "Full name (First Last format)",
            },
            "first_last_name": {
                "pattern": r"^FirstName LastName$",
                "variable": "<<first_name>> <<last_name>>",
                "description": "Template placeholder for name",
            },
            "email": {
                "pattern": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
                "variable": "<<email>>",
                "description": "Email address",
            },
            "phone": {
                "pattern": r"(\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}",
                "variable": "<<phone_number>>",
                "description": "Phone number",
            },
            "phone_generic": {
                "pattern": r"phone number",
                "variable": "<<phone_number>>",
                "description": "Generic phone number placeholder",
            },
            # Address Patterns
            "street_address": {
                "pattern": r"^.+Street Address.*$",
                "variable": "<<street_address>>",
                "description": "Street address",
            },
            "city_state_zip": {
                "pattern": r"City, State Zip",
                "variable": "<<city>>, <<state>> <<zip_code>>",
                "description": "City, State Zip format",
            },
            "city_state": {
                "pattern": r"City, State",
                "variable": "<<city>>, <<state>>",
                "description": "City, State format",
            },
            "city_country": {
                "pattern": r"City, Country",
                "variable": "<<city>>, <<country>>",
                "description": "City, Country format",
            },
            # Education Patterns
            "university_name": {
                "pattern": r"^[A-Z][a-zA-Z\s]+ University$",
                "variable": "<<university_name>>",
                "description": "University name",
            },
            "degree_concentration": {
                "pattern": r"Degree, Concentration",
                "variable": "<<degree>>, <<concentration>>",
                "description": "Degree and concentration",
            },
            "high_school": {
                "pattern": r"High School Name",
                "variable": "<<high_school_name>>",
                "description": "High school name",
            },
            "graduation_date": {
                "pattern": r"Graduation Date",
                "variable": "<<graduation_date>>",
                "description": "Graduation date",
            },
            # Experience Patterns
            "organization": {
                "pattern": r"^Organization$",
                "variable": "<<organization>>",
                "description": "Organization name",
            },
            "position_title": {
                "pattern": r"^Position Title$",
                "variable": "<<position_title>>",
                "description": "Job position title",
            },
            "date_range": {
                "pattern": r"Month Year – Month Year",
                "variable": "<<start_date>> – <<end_date>>",
                "description": "Date range for positions",
            },
            "study_abroad_dates": {
                "pattern": r"Month Year – Month Year",
                "variable": "<<study_abroad_start>> – <<study_abroad_end>>",
                "description": "Study abroad date range",
            },
            # Skills and Interests
            "technical_skills": {
                "pattern": r"Technical: List computer software.*",
                "variable": "<<technical_skills>>",
                "description": "Technical skills list",
            },
            "languages": {
                "pattern": r"Language: First foreign languages.*",
                "variable": "<<languages>>",
                "description": "Language skills",
            },
            "laboratory_skills": {
                "pattern": r"Laboratory: List scientific.*",
                "variable": "<<laboratory_skills>>",
                "description": "Laboratory skills",
            },
            "interests": {
                "pattern": r"Interests: List activities.*",
                "variable": "<<interests>>",
                "description": "Personal interests",
            },
        }

    def convert_reference_to_template(self, reference_path, template_path, document_type="resume"):
        """
        Convert a reference .docx file to a template file with variable placeholders

        Args:
            reference_path (str): Path to the reference .docx file
            template_path (str): Path where the template will be saved
            document_type (str): Type of document ('resume' or 'coverletter')

        Returns:
            dict: Conversion results including variable mapping and statistics
        """
        try:
            self.logger.info(f"Converting {reference_path} to template...")

            # Load the reference document
            doc = Document(reference_path)

            # Track conversion statistics
            conversion_stats = {
                "total_paragraphs": 0,
                "paragraphs_modified": 0,
                "patterns_matched": {},
                "variables_created": [],
                "unmatched_content": [],
            }

            # Process each paragraph in the document
            for paragraph in doc.paragraphs:
                conversion_stats["total_paragraphs"] += 1
                original_text = paragraph.text.strip()

                if not original_text:
                    continue

                # Apply pattern matching and replacement
                new_text, matched_patterns = self.apply_pattern_matching(original_text)

                # Update statistics
                if matched_patterns:
                    conversion_stats["paragraphs_modified"] += 1
                    for pattern_name in matched_patterns:
                        conversion_stats["patterns_matched"][pattern_name] = (
                            conversion_stats["patterns_matched"].get(pattern_name, 0) + 1
                        )

                # Update paragraph text if changes were made
                if new_text != original_text:
                    self.update_paragraph_text(paragraph, new_text)
                    self.logger.info(f"Converted: '{original_text}' -> '{new_text}'")
                elif self.should_convert_to_generic_field(original_text):
                    # Convert unmatched content to generic field
                    generic_var = f"<<field_{self.field_counter}>>"
                    self.update_paragraph_text(paragraph, generic_var)
                    conversion_stats["variables_created"].append(
                        {"original": original_text, "variable": generic_var, "type": "generic"}
                    )
                    conversion_stats["unmatched_content"].append(original_text)
                    self.field_counter += 1
                    self.logger.info(f"Generic field: '{original_text}' -> '{generic_var}'")

            # Save the template document
            doc.save(template_path)

            # Create metadata file for the template
            metadata_path = template_path.replace(".docx", "_metadata.json")
            self.save_template_metadata(metadata_path, conversion_stats, document_type)

            self.logger.info(f"Template conversion completed: {template_path}")
            return conversion_stats

        except Exception as e:
            self.logger.error(f"Error converting template: {str(e)}")
            raise

    def apply_pattern_matching(self, text):
        """
        Apply pattern matching rules to identify and replace content with variables

        Args:
            text (str): Original text to process

        Returns:
            tuple: (modified_text, list_of_matched_patterns)
        """
        modified_text = text
        matched_patterns = []

        # Apply each pattern in order of specificity (more specific patterns first)
        for pattern_name, pattern_info in self.patterns.items():
            pattern = pattern_info["pattern"]
            variable = pattern_info["variable"]

            # Check if pattern matches
            if re.search(pattern, modified_text, re.IGNORECASE):
                modified_text = re.sub(pattern, variable, modified_text, flags=re.IGNORECASE)
                matched_patterns.append(pattern_name)
                break  # Only apply the first matching pattern

        return modified_text, matched_patterns

    def should_convert_to_generic_field(self, text):
        """
        Determine if unmatched content should be converted to a generic field

        Args:
            text (str): Text to evaluate

        Returns:
            bool: True if should convert to generic field
        """
        # Don't convert certain types of content
        skip_patterns = [
            r"^\s*$",  # Empty or whitespace only
            r"^\[Note:.*\]$",  # Note sections
            r"^[A-Z][a-z]+ &",  # Section headers like "Leadership &"
            r"^[A-Z][a-z]+$",  # Single words that might be section headers
            r"^Beginning with most recent.*",  # Instructional text
            r"^Begin each line with.*",  # Instructional text
            r"^With next-most recent.*",  # Instructional text
            r"^This section can be.*",  # Instructional text
            r"^If this section is.*",  # Instructional text
            r"^Do not use personal.*",  # Instructional text
            r"^Quantify where possible.*",  # Instructional text
        ]

        for skip_pattern in skip_patterns:
            if re.search(skip_pattern, text, re.IGNORECASE):
                return False

        # Convert if it's likely to be variable content
        # (short phrases, single lines, potential data fields)
        if len(text.split()) <= 10 and len(text) > 2:
            return True

        return False

    def update_paragraph_text(self, paragraph, new_text):
        """
        Update paragraph text while preserving all formatting

        Args:
            paragraph: python-docx paragraph object
            new_text (str): New text to set
        """
        # Clear existing text while preserving formatting
        for run in paragraph.runs:
            run.text = ""

        # Add new text to the first run (or create one if none exist)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

    def save_template_metadata(self, metadata_path, conversion_stats, document_type):
        """
        Save metadata about the template conversion process

        Args:
            metadata_path (str): Path to save metadata file
            conversion_stats (dict): Statistics from conversion process
            document_type (str): Type of document template
        """
        metadata = {
            "document_type": document_type,
            "conversion_date": datetime.now().isoformat(),
            "conversion_stats": conversion_stats,
            "template_version": "1.0",
            "variable_format": "double_angle_brackets",
            "notes": [
                "Template created by automated conversion process",
                "Variables use <<variable_name>> format",
                "Original formatting and structure preserved",
                "Pattern recognition used for common fields",
                "Generic fields created for unmatched content",
            ],
        }

        with open(metadata_path, "w") as f:
            json.dump(metadata, f, indent=2)

    def list_available_templates(self, template_dir):
        """
        List all available template files in a directory

        Args:
            template_dir (str): Directory containing template files

        Returns:
            list: List of available template files with metadata
        """
        templates = []

        if not os.path.exists(template_dir):
            return templates

        for filename in os.listdir(template_dir):
            if filename.endswith(".docx"):
                template_path = os.path.join(template_dir, filename)
                metadata_path = template_path.replace(".docx", "_metadata.json")

                template_info = {
                    "filename": filename,
                    "path": template_path,
                    "has_metadata": os.path.exists(metadata_path),
                }

                # Load metadata if available
                if template_info["has_metadata"]:
                    try:
                        with open(metadata_path, "r") as f:
                            metadata = json.load(f)
                            template_info["metadata"] = metadata
                    except Exception as e:
                        self.logger.warning(f"Could not load metadata for {filename}: {str(e)}")

                templates.append(template_info)

        return templates


def main():
    """
    Main function to demonstrate template conversion
    """
    converter = TemplateConverter()

    # Convert the reference resume template
    reference_path = (
        "content_template_library/reference/resume/Accessible-MCS-Resume-Template-Bullet-Points_1751349781656.docx"
    )
    template_path = "content_template_library/resumes/harvard_mcs_resume_template.docx"

    if os.path.exists(reference_path):
        print(f"Converting reference file: {reference_path}")
        stats = converter.convert_reference_to_template(reference_path, template_path, "resume")
        print(f"Conversion completed. Stats: {stats}")
    else:
        print(f"Reference file not found: {reference_path}")


if __name__ == "__main__":
    main()
