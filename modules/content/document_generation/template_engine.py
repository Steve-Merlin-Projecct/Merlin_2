"""
Template Engine for Document Generation

This module provides the core template processing engine that loads template files
and generates documents by replacing variable placeholders with actual data.
It maintains all original formatting while enabling dynamic content generation.

Key Features:
- Load and process template files with variable placeholders
- Replace variables with actual data from JSON input
- Preserve all document formatting, styles, and structure
- Handle multiple template types (resume, cover letter, etc.)
- Generate professional documents with proper metadata
"""

import os
import re
import json
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Authenticity enhancement imports
try:
    from .smart_typography import SmartTypography
    from .metadata_generator import MetadataGenerator
    AUTHENTICITY_AVAILABLE = True
except ImportError:
    logging.warning("Authenticity enhancement modules not available")
    AUTHENTICITY_AVAILABLE = False


class TemplateEngine:
    """
    Processes template files and generates documents with dynamic content

    This engine handles:
    1. Loading template files (.docx with variable placeholders)
    2. Processing variable substitution from JSON data
    3. Maintaining document formatting and structure
    4. Generating final documents with professional metadata
    """

    def __init__(self, enable_url_tracking=True, enable_authenticity=True):
        """
        Initialize the template engine with configuration

        Args:
            enable_url_tracking (bool): Enable automatic URL tracking for candidate URLs (default: True)
            enable_authenticity (bool): Enable authenticity enhancements (smart typography, realistic metadata)
        """
        self.setup_logging()
        self.template_cache = {}  # Cache loaded templates for performance
        self.variable_pattern = re.compile(r"<<([^>]+)>>")  # Pattern for <<variable_name>>
        self.job_variable_pattern = re.compile(r"\{(job_title|company_name)\}")  # Pattern for {job_title} and {company_name}

        # Enhanced formatting patterns
        self.italics_pattern = re.compile(r'\*([^*]+)\*')  # Pattern for *italics*
        self.publication_pattern = re.compile(r'\b(?:[A-Z][a-z]*(?:\s[A-Z][a-z]*)*\s(?:Journal|Magazine|Review|Times|Post|Herald|Tribune|Gazette|Chronicle|News|Weekly|Monthly|Quarterly|Report|Bulletin|Digest|Today|Business|Financial|Economic|Scientific|Academic|Medical|Legal|Technical|International|National|Global|Daily|Press|Media|Communications?|Technology|Science|Nature|Cell|PLOS|BMJ|NEJM|JAMA|IEEE|ACM|Harvard|Stanford|MIT|Oxford|Cambridge))\b')  # Pattern for publications

        # URL tracking configuration
        self.enable_url_tracking = enable_url_tracking
        self.tracked_url_cache = {}  # Cache for tracked URLs to prevent duplicate tracking entries

        # Define which variables should be converted to tracked URLs
        self.TRACKABLE_URL_VARIABLES = ['calendly_url', 'linkedin_url', 'portfolio_url']

        # Mapping from variable name to link function name for tracking system
        self.URL_VARIABLE_TO_FUNCTION = {
            'calendly_url': 'Calendly',
            'linkedin_url': 'LinkedIn',
            'portfolio_url': 'Portfolio'
        }

        # Authenticity enhancement configuration
        self.enable_authenticity = enable_authenticity and AUTHENTICITY_AVAILABLE
        if self.enable_authenticity:
            self.smart_typography = SmartTypography()
            self.metadata_generator = MetadataGenerator()
            self.logger.info("Authenticity enhancements enabled")
        else:
            self.smart_typography = None
            self.metadata_generator = None

    def setup_logging(self):
        """Configure logging for template processing"""
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    def load_template(self, template_path):
        """
        Load a template file and cache it for reuse

        Args:
            template_path (str): Path to the template .docx file

        Returns:
            Document: python-docx Document object
        """
        try:
            # Check cache first
            if template_path in self.template_cache:
                self.logger.info(f"Loading template from cache: {template_path}")
                return self.template_cache[template_path]

            # Load template from file
            self.logger.info(f"Loading template from file: {template_path}")
            doc = Document(template_path)

            # Cache the template
            self.template_cache[template_path] = doc

            return doc

        except Exception as e:
            self.logger.error(f"Error loading template {template_path}: {str(e)}")
            raise

    def generate_document(self, template_path, data, output_path=None, job_id=None, application_id=None):
        """
        Generate a document from a template using provided data

        Args:
            template_path (str): Path to the template file
            data (dict): Data dictionary containing variable values
            output_path (str): Optional path to save the generated document
            job_id (str, optional): Job UUID for URL tracking context
            application_id (str, optional): Application UUID for URL tracking context

        Returns:
            dict: Information about the generated document
        """
        try:
            self.logger.info(f"Generating document from template: {template_path}")

            # Load the template
            doc = self.load_template(template_path)

            # Track substitution statistics
            substitution_stats = {
                "variables_found": set(),
                "variables_substituted": set(),
                "variables_missing": set(),
                "total_substitutions": 0,
            }

            # Process each paragraph in the document
            for paragraph in doc.paragraphs:
                original_text = paragraph.text

                # Find all variables in the paragraph (both types)
                template_variables = self.variable_pattern.findall(original_text)
                job_variables = self.job_variable_pattern.findall(original_text)
                
                all_variables = template_variables + [f"job_variable_{var}" for var in job_variables]

                if all_variables:
                    substitution_stats["variables_found"].update(all_variables)

                    # Replace variables with actual data
                    new_text = self.substitute_variables(original_text, data, substitution_stats, job_id, application_id)

                    # Apply smart typography if enabled
                    if self.enable_authenticity and self.smart_typography:
                        new_text = self.smart_typography.enhance_paragraph_text(new_text)

                    # Apply enhanced text formatting
                    formatted_text = self.apply_enhanced_formatting(paragraph, new_text)

                    # Update paragraph if changes were made
                    if formatted_text != original_text:
                        self.update_paragraph_text(paragraph, formatted_text)
                        substitution_stats["total_substitutions"] += 1
                        self.logger.debug(f"Processed: '{original_text}' -> '{formatted_text}'")

            # Process tables if they exist
            for table in doc.tables:
                self.process_table_variables(table, data, substitution_stats, job_id, application_id)

            # Set document properties
            self.set_document_properties(doc, data)

            # Determine output path if not provided
            if output_path is None:
                output_path = self.generate_output_path(template_path, data)

            # Save the document
            doc.save(output_path)

            # Calculate final statistics
            final_stats = self.calculate_final_stats(substitution_stats)

            # Create result information
            result = {
                "template_path": template_path,
                "output_path": output_path,
                "file_size": os.path.getsize(output_path),
                "variables_processed": final_stats,
                "generation_time": datetime.now().isoformat(),
                "success": True,
            }

            self.logger.info(f"Document generated successfully: {output_path}")
            return result

        except Exception as e:
            self.logger.error(f"Error generating document: {str(e)}")
            raise

    def substitute_variables(self, text, data, stats, job_id=None, application_id=None):
        """
        Replace variable placeholders in text with actual data values
        Handles both <<template_variables>> and {job_variables} systems
        Converts URL variables to tracked redirect URLs when enabled

        Args:
            text (str): Text containing variable placeholders
            data (dict): Data dictionary with variable values
            stats (dict): Statistics tracking dictionary
            job_id (str, optional): Job UUID for URL tracking context
            application_id (str, optional): Application UUID for URL tracking context

        Returns:
            str: Text with variables replaced
        """

        def replace_template_variable(match):
            """Replace <<template_variable>> format"""
            variable_name = match.group(1).strip()

            # Handle nested data access (e.g., person.name)
            value = self.get_nested_value(data, variable_name)

            if value is not None:
                # Check if this is a trackable URL variable
                if self.enable_url_tracking and variable_name in self.TRACKABLE_URL_VARIABLES:
                    # Convert to tracked URL
                    link_function = self.URL_VARIABLE_TO_FUNCTION[variable_name]
                    tracked_url = self._get_tracked_url(
                        original_url=str(value),
                        link_function=link_function,
                        job_id=job_id,
                        application_id=application_id
                    )
                    stats["variables_substituted"].add(variable_name)
                    self.logger.info(f"Converted {variable_name} to tracked URL: {tracked_url}")
                    return tracked_url
                else:
                    # Normal variable substitution
                    stats["variables_substituted"].add(variable_name)
                    return str(value)
            else:
                stats["variables_missing"].add(variable_name)
                self.logger.warning(f"Template variable '{variable_name}' not found in data")
                return match.group(0)  # Return original placeholder

        def replace_job_variable(match):
            """Replace {job_title} and {company_name} format"""
            variable_name = match.group(1).strip()

            # Look for job-specific data
            value = None
            if variable_name == 'job_title':
                value = data.get('job_title') or data.get('position_title') or data.get('title')
            elif variable_name == 'company_name':
                value = data.get('company_name') or data.get('company') or data.get('organization')

            if value is not None:
                stats["variables_substituted"].add(f"job_variable_{variable_name}")
                self.logger.debug(f"Substituted job variable '{variable_name}' with '{value}'")
                return str(value)
            else:
                stats["variables_missing"].add(f"job_variable_{variable_name}")
                self.logger.warning(f"Job variable '{variable_name}' not found in data")
                return match.group(0)  # Return original placeholder

        # Apply template variable substitution first (<<variable>>)
        text = self.variable_pattern.sub(replace_template_variable, text)

        # Apply job variable substitution second ({job_title}, {company_name})
        text = self.job_variable_pattern.sub(replace_job_variable, text)

        return text

    def _get_tracked_url(self, original_url, link_function, job_id=None, application_id=None):
        """
        Generate tracked redirect URL using LinkTracker system

        This method converts a candidate's original URL (Calendly, LinkedIn, Portfolio)
        into a tracked redirect URL that records click events and analytics.

        Args:
            original_url (str): The actual destination URL (e.g., calendly.com/steve-glen/30min)
            link_function (str): Category of link ('Calendly', 'LinkedIn', 'Portfolio')
            job_id (str, optional): Job UUID for association tracking
            application_id (str, optional): Application UUID for association tracking

        Returns:
            str: Tracked redirect URL (e.g., https://domain.com/track/lt_abc123) or original URL if tracking fails

        Example:
            >>> engine = TemplateEngine()
            >>> tracked = engine._get_tracked_url(
            ...     original_url="https://calendly.com/steve-glen/30min",
            ...     link_function="Calendly",
            ...     job_id="550e8400-e29b-41d4-a716-446655440000"
            ... )
            >>> print(tracked)
            'https://domain.com/track/lt_calendly_abc123def456'
        """
        # Check cache first to prevent duplicate tracking entries
        cache_key = f"{job_id}:{application_id}:{link_function}:{original_url}"
        if cache_key in self.tracked_url_cache:
            cached_url = self.tracked_url_cache[cache_key]
            self.logger.debug(f"Using cached tracked URL for {link_function}: {cached_url}")
            return cached_url

        try:
            # Lazy import to handle cases where LinkTracker module might not be available
            from modules.link_tracking.link_tracker import LinkTracker

            # Initialize tracker
            tracker = LinkTracker()

            # Create tracked link with job/application context
            result = tracker.create_tracked_link(
                original_url=original_url,
                link_function=link_function,
                job_id=job_id,
                application_id=application_id,
                link_type='profile',  # All candidate URLs are profile type
                description=f'{link_function} link for job application'
            )

            # Extract redirect URL from result
            redirect_url = result['redirect_url']
            tracking_id = result['tracking_id']

            # Cache the result
            self.tracked_url_cache[cache_key] = redirect_url

            self.logger.info(
                f"Generated tracked {link_function} URL: {tracking_id} "
                f"(job_id={job_id}, app_id={application_id})"
            )

            return redirect_url

        except ImportError as e:
            self.logger.error(f"LinkTracker module not available: {e}. Using original URL.")
            return original_url  # Graceful fallback

        except Exception as e:
            self.logger.error(
                f"Failed to create tracked URL for {link_function}: {e}. "
                f"Using original URL as fallback."
            )
            return original_url  # Graceful fallback - document generation continues

    def get_nested_value(self, data, key):
        """
        Get value from nested dictionary using dot notation

        Args:
            data (dict): Data dictionary
            key (str): Key that may contain dots for nested access

        Returns:
            any: Value from dictionary or None if not found
        """
        keys = key.split(".")
        current = data

        try:
            for k in keys:
                if isinstance(current, dict) and k in current:
                    current = current[k]
                else:
                    return None
            return current
        except (KeyError, TypeError):
            return None

    def apply_enhanced_formatting(self, paragraph, text):
        """
        Apply enhanced text formatting including publication italics and Canadian spelling
        
        Args:
            paragraph: python-docx paragraph object
            text (str): Text to format
            
        Returns:
            str: Formatted text
        """
        # Apply Canadian spelling corrections
        text = self.apply_canadian_spelling(text)
        
        # Apply publication italics formatting
        text = self.apply_publication_italics(paragraph, text)
        
        return text
    
    def apply_canadian_spelling(self, text):
        """
        Apply Canadian spelling corrections to text
        
        Args:
            text (str): Text to convert
            
        Returns:
            str: Text with Canadian spellings
        """
        try:
            # Import Canadian spelling processor
            import sys
            sys.path.append('.')
            from modules.content.copywriting_evaluator.canadian_spelling_processor import CanadianSpellingProcessor
            
            processor = CanadianSpellingProcessor()
            
            # Apply spelling conversions
            conversions = processor._get_spelling_conversions()
            converted_text, changes = processor._apply_spelling_conversions(text, conversions)
            
            if changes:
                self.logger.info(f"Applied {len(changes)} Canadian spelling corrections")
                
            return converted_text
            
        except Exception as e:
            self.logger.warning(f"Canadian spelling conversion failed: {str(e)}")
            return text  # Return original text if conversion fails
    
    def apply_publication_italics(self, paragraph, text):
        """
        Apply italics formatting to publications and manual italics markers
        
        Args:
            paragraph: python-docx paragraph object  
            text (str): Text to format
            
        Returns:
            str: Text with italics markers processed
        """
        try:
            # Process manual italics markers (*text*)
            def process_manual_italics(match):
                return match.group(1)  # Remove asterisks, formatting applied to run
            
            # Remove manual italics markers from text
            processed_text = self.italics_pattern.sub(process_manual_italics, text)
            
            # Find publication names and manual italics for run-level formatting
            self.apply_run_formatting(paragraph, text)
            
            return processed_text
            
        except Exception as e:
            self.logger.warning(f"Publication italics formatting failed: {str(e)}")
            return text
    
    def apply_run_formatting(self, paragraph, text):
        """
        Apply run-level formatting for italics to specific text patterns
        
        Args:
            paragraph: python-docx paragraph object
            text (str): Text content to analyze for formatting
        """
        try:
            from docx.shared import RGBColor
            
            # Clear existing runs to rebuild with formatting
            for run in paragraph.runs[:]:
                paragraph._element.remove(run._element)
            
            current_pos = 0
            
            # Find all manual italics (*text*) and publications
            manual_italics = list(self.italics_pattern.finditer(text))
            publications = list(self.publication_pattern.finditer(text))
            
            # Combine and sort all formatting ranges
            formatting_ranges = []
            
            for match in manual_italics:
                formatting_ranges.append({
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'manual_italics',
                    'content': match.group(1)
                })
            
            for match in publications:
                formatting_ranges.append({
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'publication',
                    'content': match.group(0)
                })
            
            # Sort by start position
            formatting_ranges.sort(key=lambda x: x['start'])
            
            # Build text with appropriate formatting
            for formatting in formatting_ranges:
                # Add text before formatting
                if current_pos < formatting['start']:
                    normal_run = paragraph.add_run(text[current_pos:formatting['start']])
                
                # Add formatted text
                if formatting['type'] == 'manual_italics':
                    # Manual italics: use inner content with italics
                    italic_run = paragraph.add_run(formatting['content'])
                    italic_run.italic = True
                elif formatting['type'] == 'publication':
                    # Publication: use full content with italics
                    pub_run = paragraph.add_run(formatting['content'])
                    pub_run.italic = True
                
                current_pos = formatting['end']
            
            # Add remaining text
            if current_pos < len(text):
                final_run = paragraph.add_run(text[current_pos:])
                
        except Exception as e:
            self.logger.warning(f"Run formatting failed: {str(e)}")
            # Fallback: add text without special formatting
            paragraph.add_run(text)

    def process_table_variables(self, table, data, stats, job_id=None, application_id=None):
        """
        Process variables within document tables

        Args:
            table: python-docx table object
            data (dict): Data dictionary
            stats (dict): Statistics tracking dictionary
            job_id (str, optional): Job UUID for URL tracking context
            application_id (str, optional): Application UUID for URL tracking context
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    template_variables = self.variable_pattern.findall(original_text)
                    job_variables = self.job_variable_pattern.findall(original_text)

                    all_variables = template_variables + [f"job_variable_{var}" for var in job_variables]

                    if all_variables:
                        stats["variables_found"].update(all_variables)
                        new_text = self.substitute_variables(original_text, data, stats, job_id, application_id)

                        if new_text != original_text:
                            # Apply enhanced formatting to table cells too
                            formatted_text = self.apply_enhanced_formatting(paragraph, new_text)
                            self.update_paragraph_text(paragraph, formatted_text)
                            stats["total_substitutions"] += 1

    def update_paragraph_text(self, paragraph, new_text):
        """
        Update paragraph text while preserving formatting

        Args:
            paragraph: python-docx paragraph object
            new_text (str): New text to set
        """
        # Clear existing text while preserving formatting
        for run in paragraph.runs:
            run.text = ""

        # Add new text to the first run (or create one if none exist)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

    def set_document_properties(self, doc, data):
        """
        Set document properties and metadata with authenticity enhancements

        Args:
            doc: python-docx Document object
            data (dict): Data dictionary containing document information
        """
        # Use MetadataGenerator if authenticity is enabled
        if self.enable_authenticity and self.metadata_generator:
            # Prepare metadata parameters
            author_name = data.get("author") or (
                data.get("first_name", "") + " " + data.get("last_name", "")
            ).strip()
            document_title = data.get("title", "Generated Document")
            document_type = data.get("document_type", "resume")

            # Generate realistic metadata
            metadata = self.metadata_generator.generate_metadata_for_doc_object(
                doc=doc,
                document_type=document_type,
                author_name=author_name,
                document_title=document_title,
                subject=data.get("subject", "Professional Document"),
                keywords=data.get("keywords", "Resume, Professional, Application"),
            )

            self.logger.info(
                f"Applied authentic metadata: created={metadata['core_properties']['created'].strftime('%Y-%m-%d %H:%M')}, "
                f"editing_time={metadata['app_properties']['total_time']}min, "
                f"revision={metadata['core_properties']['revision']}"
            )

        else:
            # Fallback to original method
            doc.core_properties.title = data.get("title", "Generated Document")
            doc.core_properties.author = data.get(
                "author", data.get("first_name", "") + " " + data.get("last_name", "")
            ).strip()
            doc.core_properties.subject = data.get("subject", "Professional Document")
            doc.core_properties.keywords = data.get("keywords", "Resume, Professional, Application")
            doc.core_properties.comments = data.get(
                "comments", f"Generated on {datetime.now().strftime('%Y-%m-%d')} for professional use"
            )
            doc.core_properties.category = data.get("category", "Job Application")
            doc.core_properties.created = datetime.now()
            doc.core_properties.modified = datetime.now()

            # Set language and location
            doc.core_properties.language = data.get("language", "en-CA")

            # Set version and revision
            doc.core_properties.version = data.get("version", "1.0")
            doc.core_properties.revision = data.get("revision", 1)

    def generate_output_path(self, template_path, data):
        """
        Generate an output path for the document based on template and data

        Args:
            template_path (str): Path to the template file
            data (dict): Data dictionary

        Returns:
            str: Generated output path
        """
        # Extract template name without extension
        template_name = os.path.splitext(os.path.basename(template_path))[0]

        # Create filename with person's name if available
        name_part = ""
        if "first_name" in data or "last_name" in data:
            first_name = data.get("first_name", "").strip()
            last_name = data.get("last_name", "").strip()
            if first_name or last_name:
                name_part = f"_{first_name}_{last_name}".replace(" ", "_")

        # Create timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # Generate filename
        filename = f"{template_name}{name_part}_{timestamp}.docx"

        # Use storage directory
        storage_dir = os.path.join(os.getcwd(), "storage")
        os.makedirs(storage_dir, exist_ok=True)

        return os.path.join(storage_dir, filename)

    def calculate_final_stats(self, stats):
        """
        Calculate final statistics for the generation process

        Args:
            stats (dict): Statistics tracking dictionary

        Returns:
            dict: Final statistics summary
        """
        return {
            "total_variables_found": len(stats["variables_found"]),
            "variables_substituted": len(stats["variables_substituted"]),
            "variables_missing": len(stats["variables_missing"]),
            "total_substitutions": stats["total_substitutions"],
            "substitution_rate": len(stats["variables_substituted"]) / max(1, len(stats["variables_found"])),
            "missing_variables": list(stats["variables_missing"]),
            "processed_variables": list(stats["variables_substituted"]),
        }

    def validate_template(self, template_path):
        """
        Validate that a template file is properly formatted

        Args:
            template_path (str): Path to the template file

        Returns:
            dict: Validation results
        """
        try:
            doc = Document(template_path)

            # Find all variables in the template (both types)
            template_variables = set()
            job_variables = set()
            
            for paragraph in doc.paragraphs:
                template_variables.update(self.variable_pattern.findall(paragraph.text))
                job_variables.update(self.job_variable_pattern.findall(paragraph.text))

            # Check tables for variables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            template_variables.update(self.variable_pattern.findall(paragraph.text))
                            job_variables.update(self.job_variable_pattern.findall(paragraph.text))

            # Combine all variables for reporting
            all_variables = list(template_variables) + [f"job_variable_{var}" for var in job_variables]

            return {
                "valid": True,
                "variables_found": all_variables,
                "template_variables": list(template_variables),
                "job_variables": list(job_variables),
                "variable_count": len(all_variables),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }

        except Exception as e:
            return {"valid": False, "error": str(e), "variables_found": [], "variable_count": 0}


def main():
    """
    Main function to demonstrate template engine usage
    """
    engine = TemplateEngine()

    # Example usage
    template_path = "content_template_library/resumes/harvard_mcs_resume_template.docx"

    # Sample data
    data = {
        "first_name": "John",
        "last_name": "Doe",
        "email": "john.doe@example.com",
        "phone_number": "(555) 123-4567",
        "street_address": "123 Main Street",
        "city": "Boston",
        "state": "MA",
        "zip_code": "02101",
        "title": "Software Engineer Resume",
        "author": "John Doe",
    }

    if os.path.exists(template_path):
        result = engine.generate_document(template_path, data)
        print(f"Document generated: {result}")
    else:
        print(f"Template not found: {template_path}")


if __name__ == "__main__":
    main()
