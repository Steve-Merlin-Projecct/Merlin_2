#!/usr/bin/env python3
"""
CSV Content Mapper

Maps content from CSV mapping files to resume templates with variable substitution,
static text replacement, and content removal based on mapping specifications.

Preserves original document formatting while applying dynamic content based on:
- User profile data
- Job-specific content selection
- Static text improvements
- Content removal specifications
"""

import csv
import logging
from typing import Dict, List, Tuple, Optional
from pathlib import Path


class CSVContentMapper:
    """
    Maps CSV content specifications to template variables and transformations

    Processes CSV mapping files that specify:
    - Variables: Dynamic content from user profiles or content manager
    - Static changes: Fixed text improvements for all resumes
    - Discarded content: Text to be removed from template
    """

    def __init__(self):
        """Initialize CSV content mapper"""
        self.logger = logging.getLogger(__name__)
        self.mapping_cache = {}

        # Content categories for variable resolution
        self.variable_categories = {
            "user_profile": [
                "user_first_name",
                "user_last_name",
                "user_linkedin",
                "user_city_prov",
                "user_email",
                "user_phone",
            ],
            "education": [
                "edu_1_name",
                "edu_1_degree",
                "edu_1_concentration",
                "edu_1_specialization",
                "edu_1_location",
                "edu_1_grad_date",
            ],
            "work_experience": [
                "work_experience_1_name",
                "work_experience_1_position",
                "work_experience_1_location",
                "work_experience_1_dates",
                "work_experience_1_skill1",
                "work_experience_1_skill2",
                "work_experience_1_skill3",
                "work_experience_1_skill4",
                "work_experience_2_name",
                "work_experience_2_position",
                "work_experience_2_location",
                "work_experience_2_dates",
                "work_experience_2_skill1",
                "work_experience_2_skill2",
                "work_experience_2_skill3",
                "work_experience_2_skill4",
            ],
            "volunteer": [
                "volunteer_1_name",
                "volunteer_1_position",
                "volunteer_1_date",
                "volunteer_1_skill1",
                "volunteer_1_skill2",
            ],
            "skills": ["technical_summary", "methodology_summary", "domain_summary"],
        }

        self.logger.info("CSVContentMapper initialized for dynamic resume content mapping")

    def load_mapping_from_csv(self, csv_path: str) -> Dict:
        """
        Load content mapping from CSV file

        Args:
            csv_path: Path to CSV mapping file

        Returns:
            Dict: Parsed mapping with variables, static changes, and discarded content
        """
        try:
            mapping = {"variables": {}, "static_changes": {}, "discarded_content": [], "variable_metadata": {}}

            with open(csv_path, "r", encoding="utf-8") as file:
                reader = csv.DictReader(file)

                for row in reader:
                    original_text = row.get("Original_Text", "").strip()
                    is_variable = row.get("Is_Variable", "").upper() == "TRUE"
                    is_static = row.get("Is_Static", "").upper() == "TRUE"
                    is_discarded = row.get("Is_discarded", "").upper() == "TRUE"
                    variable_name = row.get("Variable_name", "").strip()
                    variable_intention = row.get("Variable intention", "").strip()
                    static_text = row.get("Static_Text", "").strip()

                    if not original_text:
                        continue

                    if is_variable and variable_name:
                        # Map original text to variable name for template substitution
                        mapping["variables"][original_text] = variable_name
                        mapping["variable_metadata"][variable_name] = {
                            "intention": variable_intention,
                            "original_text": original_text,
                            "category": self._categorize_variable(variable_name),
                        }
                        self.logger.debug(f"Variable mapping: '{original_text}' -> {variable_name}")

                    elif is_static and static_text:
                        # Map original text to static replacement
                        mapping["static_changes"][original_text] = static_text
                        self.logger.debug(f"Static change: '{original_text}' -> '{static_text}'")

                    elif is_discarded:
                        # Mark content for removal
                        mapping["discarded_content"].append(original_text)
                        self.logger.debug(f"Discarded content: '{original_text}'")

            # Cache the mapping
            self.mapping_cache[csv_path] = mapping

            self.logger.info(
                f"Loaded mapping: {len(mapping['variables'])} variables, "
                f"{len(mapping['static_changes'])} static changes, "
                f"{len(mapping['discarded_content'])} discarded items"
            )

            return mapping

        except Exception as e:
            self.logger.error(f"Failed to load CSV mapping from {csv_path}: {e}")
            raise

    def _categorize_variable(self, variable_name: str) -> str:
        """
        Categorize variable by type for content resolution

        Args:
            variable_name: Name of the variable

        Returns:
            str: Category name for variable resolution
        """
        for category, variables in self.variable_categories.items():
            if variable_name in variables:
                return category
        return "other"

    def apply_mapping_to_template(self, template_path: str, mapping: Dict, content_data: Dict) -> str:
        """
        Apply CSV mapping to template file, creating variable-based template

        Args:
            template_path: Path to original template file
            mapping: Mapping dictionary from load_mapping_from_csv
            content_data: User and job-specific content for variable resolution

        Returns:
            str: Path to processed template with variables applied
        """
        try:
            # Import python-docx using the same pattern as other modules
            import sys

            sys.path.append("/home/runner/workspace")

            try:
                from docx import Document
            except ImportError:
                # Try alternative import
                from python_docx import Document

            # Load the original template
            doc = Document(template_path)

            # Track changes made
            changes_made = {
                "variables_replaced": 0,
                "static_changes": 0,
                "content_discarded": 0,
                "paragraphs_modified": 0,
            }

            # Process all paragraphs in the document
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                if not original_text.strip():
                    continue

                modified = False
                new_text = original_text

                # Apply discarded content removal first
                for discarded_text in mapping.get("discarded_content", []):
                    if discarded_text in new_text:
                        new_text = new_text.replace(discarded_text, "")
                        changes_made["content_discarded"] += 1
                        modified = True
                        self.logger.debug(f"Removed discarded content: '{discarded_text}'")

                # Apply static text changes
                for original, static_replacement in mapping.get("static_changes", {}).items():
                    if original in new_text:
                        new_text = new_text.replace(original, static_replacement)
                        changes_made["static_changes"] += 1
                        modified = True
                        self.logger.debug(f"Static change: '{original}' -> '{static_replacement}'")

                # Apply variable substitutions with template variable format
                for original, variable_name in mapping.get("variables", {}).items():
                    if original in new_text:
                        # Replace with template variable format: <<variable_name>>
                        template_variable = f"<<{variable_name}>>"
                        new_text = new_text.replace(original, template_variable)
                        changes_made["variables_replaced"] += 1
                        modified = True
                        self.logger.debug(f"Variable substitution: '{original}' -> '{template_variable}'")

                # Update paragraph text if modified
                if modified:
                    # Preserve formatting by updating runs individually
                    if paragraph.runs:
                        # Clear existing runs and add new text
                        for run in paragraph.runs:
                            run.text = ""
                        paragraph.runs[0].text = new_text
                    else:
                        paragraph.text = new_text

                    changes_made["paragraphs_modified"] += 1

            # Process tables if any exist
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._apply_mapping_to_paragraph(paragraph, mapping, changes_made)

            # Generate output path
            template_name = Path(template_path).stem
            output_path = f"content_template_library/mapped_templates/{template_name}_mapped.docx"

            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)

            # Save the processed template
            doc.save(output_path)

            self.logger.info(f"Applied mapping to template: {changes_made}")
            self.logger.info(f"Processed template saved to: {output_path}")

            return output_path

        except Exception as e:
            self.logger.error(f"Failed to apply mapping to template: {e}")
            raise

    def _apply_mapping_to_paragraph(self, paragraph, mapping: Dict, changes_made: Dict) -> None:
        """
        Apply mapping transformations to a single paragraph

        Args:
            paragraph: Document paragraph object
            mapping: Mapping dictionary
            changes_made: Statistics tracking dictionary
        """
        original_text = paragraph.text
        if not original_text.strip():
            return

        modified = False
        new_text = original_text

        # Apply transformations in order: discard, static, variables
        for discarded_text in mapping.get("discarded_content", []):
            if discarded_text in new_text:
                new_text = new_text.replace(discarded_text, "")
                changes_made["content_discarded"] += 1
                modified = True

        for original, static_replacement in mapping.get("static_changes", {}).items():
            if original in new_text:
                new_text = new_text.replace(original, static_replacement)
                changes_made["static_changes"] += 1
                modified = True

        for original, variable_name in mapping.get("variables", {}).items():
            if original in new_text:
                template_variable = f"<<{variable_name}>>"
                new_text = new_text.replace(original, template_variable)
                changes_made["variables_replaced"] += 1
                modified = True

        if modified:
            if paragraph.runs:
                for run in paragraph.runs:
                    run.text = ""
                paragraph.runs[0].text = new_text
            else:
                paragraph.text = new_text
            changes_made["paragraphs_modified"] += 1

    def resolve_variables_from_content(self, mapping: Dict, content_data: Dict) -> Dict:
        """
        Resolve template variables using content data and user profiles

        Args:
            mapping: Mapping dictionary with variable metadata
            content_data: User profile and job-specific content

        Returns:
            Dict: Resolved variable values for template substitution
        """
        resolved_variables = {}

        try:
            # Extract user profile data
            user_profile = content_data.get("user_profile", {})
            job_data = content_data.get("job_data", {})
            content_selections = content_data.get("content_selections", {})

            for variable_name, metadata in mapping.get("variable_metadata", {}).items():
                category = metadata.get("category", "other")
                intention = metadata.get("intention", "")

                # Resolve based on category
                if category == "user_profile":
                    resolved_variables[variable_name] = self._resolve_user_profile_variable(variable_name, user_profile)
                elif category == "education":
                    resolved_variables[variable_name] = self._resolve_education_variable(variable_name, user_profile)
                elif category == "work_experience":
                    resolved_variables[variable_name] = self._resolve_work_experience_variable(
                        variable_name, user_profile, content_selections
                    )
                elif category == "volunteer":
                    resolved_variables[variable_name] = self._resolve_volunteer_variable(
                        variable_name, user_profile, content_selections
                    )
                elif category == "skills":
                    resolved_variables[variable_name] = self._resolve_skills_variable(
                        variable_name, user_profile, job_data
                    )
                else:
                    # Default resolution
                    resolved_variables[variable_name] = f"[{variable_name}]"

                self.logger.debug(f"Resolved {variable_name}: {resolved_variables[variable_name][:50]}...")

            self.logger.info(f"Resolved {len(resolved_variables)} variables for template")
            return resolved_variables

        except Exception as e:
            self.logger.error(f"Failed to resolve variables: {e}")
            return {}

    def _resolve_user_profile_variable(self, variable_name: str, user_profile: Dict) -> str:
        """Resolve user profile variables"""
        mapping = {
            "user_first_name": user_profile.get("first_name", "Steve"),
            "user_last_name": user_profile.get("last_name", "Glen"),
            "user_linkedin": "linkedin.com/in/steve-glen",  # Will add tracking
            "user_city_prov": f"{user_profile.get('city', 'Edmonton')}, {user_profile.get('province', 'AB')}",
            "user_email": user_profile.get("email", "therealstevenglen@gmail.com"),
            "user_phone": user_profile.get("phone", "780-884-7038"),
        }
        return mapping.get(variable_name, f"[{variable_name}]")

    def _resolve_education_variable(self, variable_name: str, user_profile: Dict) -> str:
        """Resolve education variables"""
        education = user_profile.get("education", {})
        mapping = {
            "edu_1_name": education.get("institution", "University of Alberta"),
            "edu_1_degree": education.get("degree_type", "Bachelor"),
            "edu_1_concentration": education.get("field_of_study", "Commerce"),
            "edu_1_specialization": education.get("specialization", "Entrepreneurship, Strategy, Marketing"),
            "edu_1_location": f"{education.get('city', 'Edmonton')}, {education.get('province', 'AB')}",
            "edu_1_grad_date": str(education.get("graduation_year", "2018")),
        }
        return mapping.get(variable_name, f"[{variable_name}]")

    def _resolve_work_experience_variable(
        self, variable_name: str, user_profile: Dict, content_selections: Dict
    ) -> str:
        """Resolve work experience variables"""
        experiences = user_profile.get("work_experience", [])

        if "work_experience_1_" in variable_name:
            exp = experiences[0] if experiences else {}
            if "skill" in variable_name:
                skill_num = variable_name.split("skill")[1]
                return content_selections.get(f"work_exp_1_skill_{skill_num}", f"[work experience 1 skill {skill_num}]")
            else:
                mapping = {
                    "work_experience_1_name": exp.get("company", "Odvod Media"),
                    "work_experience_1_position": exp.get("position", "Digital Strategist"),
                    "work_experience_1_location": f"{exp.get('city', 'Edmonton')} {exp.get('province', 'AB')}",
                    "work_experience_1_dates": f"{exp.get('start_year', '2020')} - {exp.get('end_year', 'current')}",
                }
                return mapping.get(variable_name, f"[{variable_name}]")

        elif "work_experience_2_" in variable_name:
            exp = experiences[1] if len(experiences) > 1 else {}
            if "skill" in variable_name:
                skill_num = variable_name.split("skill")[1]
                return content_selections.get(f"work_exp_2_skill_{skill_num}", f"[work experience 2 skill {skill_num}]")
            else:
                mapping = {
                    "work_experience_2_name": exp.get("company", "Rona"),
                    "work_experience_2_position": exp.get("position", "Visual Merchandiser"),
                    "work_experience_2_location": f"{exp.get('city', 'Edmonton')} {exp.get('province', 'AB')}",
                    "work_experience_2_dates": f"{exp.get('start_year', '2014')} - {exp.get('end_year', '2017')}",
                }
                return mapping.get(variable_name, f"[{variable_name}]")

        return f"[{variable_name}]"

    def _resolve_volunteer_variable(self, variable_name: str, user_profile: Dict, content_selections: Dict) -> str:
        """Resolve volunteer experience variables"""
        volunteers = user_profile.get("volunteer_experience", [])
        volunteer = volunteers[0] if volunteers else {}

        if "skill" in variable_name:
            skill_num = variable_name.split("skill")[1]
            return content_selections.get(f"volunteer_skill_{skill_num}", f"[volunteer skill {skill_num}]")
        else:
            mapping = {
                "volunteer_1_name": volunteer.get("organization", "[volunteer organization]"),
                "volunteer_1_position": volunteer.get("position", "[volunteer position]"),
                "volunteer_1_date": f"{volunteer.get('start_year', '')} - {volunteer.get('end_year', '')}",
            }
            return mapping.get(variable_name, f"[{variable_name}]")

    def _resolve_skills_variable(self, variable_name: str, user_profile: Dict, job_data: Dict) -> str:
        """Resolve skills variables"""
        skills = user_profile.get("skills", {})

        mapping = {
            "technical_summary": ", ".join(
                skills.get(
                    "technical",
                    [
                        "Microsoft Office Suite",
                        "Google Analytics",
                        "Adobe Creative Suite",
                        "WordPress",
                        "Mailchimp",
                        "Hootsuite",
                    ],
                )
            ),
            "methodology_summary": ", ".join(
                skills.get("methodologies", ["Agile", "Scrum", "Lean UX", "Design Thinking"])
            ),
            "domain_summary": ", ".join(
                skills.get(
                    "domains",
                    ["Digital Marketing", "Content Strategy", "Brand Management", "Public Relations", "Data Analysis"],
                )
            ),
        }
        return mapping.get(variable_name, f"[{variable_name}]")

    def get_mapping_summary(self, mapping: Dict) -> Dict:
        """
        Get summary statistics of mapping

        Args:
            mapping: Mapping dictionary

        Returns:
            Dict: Summary statistics
        """
        return {
            "total_variables": len(mapping.get("variables", {})),
            "static_changes": len(mapping.get("static_changes", {})),
            "discarded_items": len(mapping.get("discarded_content", [])),
            "variable_categories": {
                category: len(
                    [v for v in mapping.get("variable_metadata", {}).values() if v.get("category") == category]
                )
                for category in self.variable_categories.keys()
            },
        }
