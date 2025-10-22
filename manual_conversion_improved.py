#!/usr/bin/env python3
"""
Improved Manual Template Conversion Script
Converts raw Word document templates into variable-based templates
while preserving all original formatting.

This version uses a more intelligent approach by examining the document
structure and using context-aware replacements.
"""

import re
from pathlib import Path
from docx import Document
from typing import Dict, List, Set


class ImprovedTemplateConverter:
    """
    Converts Word document templates by replacing specific content with variables.
    Uses context-aware replacement to avoid false positives.
    """

    def __init__(self):
        """Initialize the converter."""
        self.converted_variables = {}

    def preserve_run_formatting(self, run):
        """Extract formatting properties from a run."""
        formatting = {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name if run.font.name else None,
            'font_size': run.font.size,
            'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
        }
        return formatting

    def apply_run_formatting(self, run, formatting):
        """Apply formatting properties to a run."""
        if formatting['bold'] is not None:
            run.bold = formatting['bold']
        if formatting['italic'] is not None:
            run.italic = formatting['italic']
        if formatting['underline'] is not None:
            run.underline = formatting['underline']
        if formatting['font_name']:
            run.font.name = formatting['font_name']
        if formatting['font_size']:
            run.font.size = formatting['font_size']
        if formatting['font_color']:
            run.font.color.rgb = formatting['font_color']

    def replace_in_paragraph(self, paragraph, old_text, new_text, case_sensitive=True):
        """
        Replace text in paragraph while preserving formatting.

        Args:
            paragraph: python-docx Paragraph object
            old_text: Text to find and replace
            new_text: Replacement text (variable placeholder)
            case_sensitive: Whether to match case exactly

        Returns:
            bool: True if replacement was made
        """
        full_text = paragraph.text

        # Check if text exists
        if case_sensitive:
            if old_text not in full_text:
                return False
            start_pos = full_text.find(old_text)
        else:
            if old_text.lower() not in full_text.lower():
                return False
            start_pos = full_text.lower().find(old_text.lower())

        if start_pos == -1:
            return False

        end_pos = start_pos + len(old_text)

        # Track replacement
        if old_text not in self.converted_variables:
            self.converted_variables[old_text] = new_text

        # Find affected runs
        current_pos = 0
        runs_to_modify = []

        for run in paragraph.runs:
            run_length = len(run.text)
            run_end = current_pos + run_length

            if current_pos <= start_pos < run_end or current_pos < end_pos <= run_end or (start_pos <= current_pos and end_pos >= run_end):
                runs_to_modify.append({
                    'run': run,
                    'start': max(0, start_pos - current_pos),
                    'end': min(run_length, end_pos - current_pos),
                    'formatting': self.preserve_run_formatting(run)
                })

            current_pos = run_end

        # Apply replacements
        if len(runs_to_modify) == 1:
            run_info = runs_to_modify[0]
            run = run_info['run']
            run.text = run.text[:run_info['start']] + new_text + run.text[run_info['end']:]
            return True
        elif len(runs_to_modify) > 1:
            for i, run_info in enumerate(runs_to_modify):
                if i == 0:
                    run_info['run'].text = run_info['run'].text[:run_info['start']] + new_text
                elif i == len(runs_to_modify) - 1:
                    run_info['run'].text = run_info['run'].text[run_info['end']:]
                else:
                    run_info['run'].text = ''
            return True

        return False

    def convert_restaurant_manager_template(self, input_path: str, output_path: str):
        """
        Convert the restaurant manager template with specific replacements.

        Args:
            input_path: Path to input .docx file
            output_path: Path to save converted .docx file
        """
        print(f"\n{'='*60}")
        print(f"Converting: Restaurant Manager Template")
        print(f"{'='*60}")

        doc = Document(input_path)
        self.converted_variables = {}

        # Define specific replacements for this template
        replacements = [
            # Header/Contact info
            ('May Riley', '<<first_name>> <<last_name>>'),
            ('4567 Main Street, Buffalo, New York 98052', '<<street_address>>, <<city>>, <<state>> <<zip_code>>'),
            ('(716) 555-0100', '<<phone>>'),
            ('m.riley@live.com', '<<email>>'),
            ('www.linkedin.com/in/m', '<<linkedin>>'),

            # Company names
            ('Contoso Bar and Grill', '<<company_name>>'),
            ('Fourth Coffee Bistro', '<<company_name>>'),
            ('Bigtown College, Chicago, Illinois', '<<education_institution>>, <<city>>, <<state>>'),

            # Job titles (in context)
            ('Restaurant Manager', '<<position_title>>'),

            # Dates
            ('September 20XX – Present', '<<start_date>> - <<end_date>>'),
            ('June 20XX – August 20XX', '<<start_date>> - <<end_date>>'),
            ('June 20XX', '<<graduation_date>>'),

            # Degrees
            ('B.S. in Business Administration', '<<degree>>'),
            ('A.A. in Hospitality Management', '<<degree>>'),
        ]

        # Process all paragraphs and tables
        count = 0
        for para in doc.paragraphs:
            for old_text, new_text in replacements:
                if self.replace_in_paragraph(para, old_text, new_text):
                    count += 1
                    print(f"  Replaced: '{old_text}' -> {new_text}")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in replacements:
                            if self.replace_in_paragraph(para, old_text, new_text):
                                count += 1
                                print(f"  Replaced in table: '{old_text}' -> {new_text}")

        doc.save(output_path)
        print(f"\nTotal replacements: {count}")
        print(f"Saved to: {output_path}")

        return doc, self.converted_variables

    def convert_accountant_template(self, input_path: str, output_path: str):
        """
        Convert the accountant template with specific replacements.

        Args:
            input_path: Path to input .docx file
            output_path: Path to save converted .docx file
        """
        print(f"\n{'='*60}")
        print(f"Converting: Accountant Template")
        print(f"{'='*60}")

        doc = Document(input_path)
        self.converted_variables = {}

        # Define specific replacements for this template
        replacements = [
            # Header/Contact info
            ('Danielle Brasseur', '<<first_name>> <<last_name>>'),
            ('4567 8th Avenue, Carson City, NV 10111', '<<street_address>>, <<city>>, <<state>> <<zip_code>>'),
            ('(313) 555-0100', '<<phone>>'),
            ('danielle@example.com', '<<email>>'),
            ('www.linkedin.com/in/danielle', '<<linkedin>>'),

            # Education
            ('Bachelor of Science in Accounting, Minor in Business Administration', '<<degree>>, Minor in <<minor>>'),
            ('University of Nevada, Reno, NV', '<<education_institution>>, <<city>>, <<state>>'),
            ('May 20XX', '<<graduation_date>>'),

            # Job titles and companies
            ('Accountant', '<<position_title>>'),
            ('Senior Accountant', '<<position_title>>'),
            ('Staff Accountant', '<<position_title>>'),
            ('Trey Research', '<<company_name>>'),
            ('Contoso Pharmaceuticals', '<<company_name>>'),

            # Dates
            ('June 20XX – Present', '<<start_date>> - <<end_date>>'),
            ('May 20XX – May 20XX', '<<start_date>> - <<end_date>>'),
            ('June 20XX – April 20XX', '<<start_date>> - <<end_date>>'),
        ]

        # Process all paragraphs and tables
        count = 0
        for para in doc.paragraphs:
            for old_text, new_text in replacements:
                if self.replace_in_paragraph(para, old_text, new_text):
                    count += 1
                    print(f"  Replaced: '{old_text}' -> {new_text}")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in replacements:
                            if self.replace_in_paragraph(para, old_text, new_text):
                                count += 1
                                print(f"  Replaced in table: '{old_text}' -> {new_text}")

        doc.save(output_path)
        print(f"\nTotal replacements: {count}")
        print(f"Saved to: {output_path}")

        return doc, self.converted_variables

    def print_summary(self, template_name: str):
        """Print summary of variables used."""
        print(f"\n{'-'*60}")
        print(f"Variables Summary for {template_name}")
        print(f"{'-'*60}")

        unique_vars = sorted(set(self.converted_variables.values()))
        for var in unique_vars:
            examples = [k for k, v in self.converted_variables.items() if v == var]
            print(f"\n{var}:")
            for example in examples:
                print(f"  - '{example}'")


def extract_document_content(doc_path: str):
    """
    Extract and display document content for manual review.

    Args:
        doc_path: Path to .docx file
    """
    print(f"\n{'='*60}")
    print(f"Content Extraction: {Path(doc_path).name}")
    print(f"{'='*60}\n")

    doc = Document(doc_path)

    print("PARAGRAPHS:")
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"  [{idx:2d}] {text}")

    print("\nTABLES:")
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx}:")
        for row_idx, row in enumerate(table.rows):
            print(f"  Row {row_idx}:")
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    # Limit output length
                    if len(text) > 100:
                        text = text[:100] + "..."
                    print(f"    Cell {cell_idx}: {text}")


def main():
    """Main conversion workflow."""
    print("="*60)
    print("IMPROVED WORD DOCUMENT TEMPLATE CONVERTER")
    print("="*60)
    print("Context-aware conversion with specific replacements\n")

    base_dir = Path("/workspace/content_template_library/manual_converted")

    # First, extract content for review
    print("\nSTEP 1: Extracting content for analysis...")

    restaurant_path = base_dir / "restaurant_manager_template.docx"
    accountant_path = base_dir / "accountant_template.docx"

    if restaurant_path.exists():
        extract_document_content(str(restaurant_path))

    if accountant_path.exists():
        extract_document_content(str(accountant_path))

    # Now convert
    print("\n\nSTEP 2: Converting templates...")

    converter = ImprovedTemplateConverter()
    results = {}

    # Convert restaurant manager template
    if restaurant_path.exists():
        try:
            output_path = base_dir / "restaurant_manager_template_converted.docx"
            doc, variables = converter.convert_restaurant_manager_template(
                str(restaurant_path),
                str(output_path)
            )
            converter.print_summary("Restaurant Manager Template")
            results['Restaurant Manager'] = {'success': True, 'output': output_path}
        except Exception as e:
            print(f"ERROR: {e}")
            import traceback
            traceback.print_exc()
            results['Restaurant Manager'] = {'success': False, 'error': str(e)}

    # Convert accountant template
    if accountant_path.exists():
        try:
            output_path = base_dir / "accountant_template_converted.docx"
            doc, variables = converter.convert_accountant_template(
                str(accountant_path),
                str(output_path)
            )
            converter.print_summary("Accountant Template")
            results['Accountant'] = {'success': True, 'output': output_path}
        except Exception as e:
            print(f"ERROR: {e}")
            import traceback
            traceback.print_exc()
            results['Accountant'] = {'success': False, 'error': str(e)}

    # Final summary
    print("\n\n" + "="*60)
    print("CONVERSION SUMMARY")
    print("="*60)

    for template_name, result in results.items():
        print(f"\n{template_name}:")
        if result['success']:
            print(f"  Status: SUCCESS")
            print(f"  Output: {result['output']}")
        else:
            print(f"  Status: FAILED")
            print(f"  Error: {result['error']}")

    print("\n" + "="*60)


if __name__ == "__main__":
    main()
