#!/usr/bin/env python3
"""
Template Variable Insertion System
Populates fully converted templates with user-provided data
Handles all 92 variables across three templates
"""

import json
import os
from datetime import datetime
from docx import Document
from typing import Dict, List, Optional, Any
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class TemplateVariableInserter:
    """Handles insertion of variables into fully converted templates"""

    def __init__(self):
        self.template_dir = "./content_template_library/manual_converted"
        self.output_dir = "./content_template_library/generated"

        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)

        # Define templates and their variables
        self.templates = {
            "restaurant_manager": {
                "file": "restaurant_manager_fully_converted.docx",
                "variables": self._get_restaurant_manager_variables()
            },
            "accountant": {
                "file": "accountant_fully_converted.docx",
                "variables": self._get_accountant_variables()
            },
            "uiux_designer": {
                "file": "uiux_designer_fully_converted.docx",
                "variables": self._get_uiux_designer_variables()
            }
        }

    def _get_restaurant_manager_variables(self) -> List[str]:
        """Get all 43 variables for Restaurant Manager template"""
        return [
            # Contact Information
            "first_name", "last_name", "street_address", "city", "state",
            "zip_code", "phone", "email", "linkedin_url",

            # Section Headers
            "profile_section_header", "experience_section_header",
            "education_section_header", "skills_section_header",
            "interests_section_header",

            # Career Overview
            "career_overview_1", "career_overview_2", "career_overview_3",
            "career_overview_4", "career_overview_5",

            # Job 1
            "position_1", "company_1", "start_date_1", "end_date_1",
            "job_1_responsibility_1", "job_1_achievement_1", "job_1_achievement_2",

            # Job 2
            "position_2", "company_2", "start_date_2", "end_date_2",
            "job_2_responsibility_1", "job_2_achievement_1",
            "job_2_achievement_2", "job_2_achievement_3",

            # Education
            "degree_1", "graduation_date_1", "institution_1",
            "institution_1_city", "institution_1_state",
            "degree_2", "graduation_date_2", "institution_2",
            "institution_2_city", "institution_2_state",

            # Skills
            "skill_1", "skill_2", "skill_3", "skill_4", "skill_5", "skill_6",

            # Interests
            "interest_1", "interest_2", "interest_3",
            "interest_4", "interest_5", "interest_6"
        ]

    def _get_accountant_variables(self) -> List[str]:
        """Get all 29 variables for Accountant template"""
        return [
            # Contact Information
            "first_name", "last_name", "street_address", "city", "state",
            "zip_code", "phone", "email", "linkedin_url",

            # Professional Summary
            "professional_summary_1", "professional_summary_2",

            # Section Headers
            "education_header", "experience_header", "skills_header",

            # Education
            "degree_1", "minor_1", "institution_1",
            "degree_label", "graduation_date_1",

            # Experience
            "position_1", "company_1", "job_1_city", "job_1_state",
            "start_date_1", "end_date_1",

            # Skills
            "technical_skill_1", "technical_skill_2", "technical_skill_3",
            "technical_skill_4", "soft_skill_1", "language_skill_1"
        ]

    def _get_uiux_designer_variables(self) -> List[str]:
        """Get all 20 variables for UI/UX Designer template"""
        return [
            # Personal
            "first_name", "last_name",

            # Professional
            "current_title",

            # Contact
            "contact_header", "email", "portfolio_website", "phone", "city",

            # Education
            "education_header", "institution_1", "degree_type_1",
            "major_1", "graduation_year_1",

            # Skills
            "skills_header", "skill_2", "skill_3", "skill_4",

            # Experience
            "experience_header", "position_2", "company_1"
        ]

    def populate_template(
        self,
        template_name: str,
        data: Dict[str, str],
        output_filename: Optional[str] = None
    ) -> str:
        """
        Populate a template with user data

        Args:
            template_name: Name of template to use
            data: Dictionary mapping variable names to values
            output_filename: Optional custom output filename

        Returns:
            Path to generated document
        """
        if template_name not in self.templates:
            raise ValueError(f"Unknown template: {template_name}")

        template_info = self.templates[template_name]
        template_path = os.path.join(self.template_dir, template_info["file"])

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")

        logger.info(f"Loading template: {template_path}")
        doc = Document(template_path)

        # Track replacements
        replacements_made = {}
        missing_variables = []

        # Get all variables for this template
        template_variables = template_info["variables"]

        # Replace variables in paragraphs
        for paragraph in doc.paragraphs:
            for var_name in template_variables:
                placeholder = f"<<{var_name}>>"
                if placeholder in paragraph.text:
                    if var_name in data and data[var_name]:
                        # Replace in all runs to preserve formatting
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, data[var_name])
                                replacements_made[var_name] = data[var_name]
                                logger.info(f"Replaced: {var_name}")
                    else:
                        missing_variables.append(var_name)

        # Replace variables in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for var_name in template_variables:
                            placeholder = f"<<{var_name}>>"
                            if placeholder in paragraph.text:
                                if var_name in data and data[var_name]:
                                    for run in paragraph.runs:
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, data[var_name])
                                            replacements_made[var_name] = data[var_name]
                                            logger.info(f"Replaced in table: {var_name}")
                                else:
                                    missing_variables.append(var_name)

        # Generate output filename
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{template_name}_{timestamp}.docx"

        output_path = os.path.join(self.output_dir, output_filename)

        # Save document
        doc.save(output_path)
        logger.info(f"Document saved: {output_path}")

        # Report statistics
        logger.info(f"Variables replaced: {len(replacements_made)}/{len(template_variables)}")

        if missing_variables:
            unique_missing = list(set(missing_variables))
            logger.warning(f"Missing variables ({len(unique_missing)}): {unique_missing[:10]}")

        return output_path

    def populate_from_json(
        self,
        template_name: str,
        json_path: str,
        output_filename: Optional[str] = None
    ) -> str:
        """
        Populate template from JSON file

        Args:
            template_name: Name of template to use
            json_path: Path to JSON file with data
            output_filename: Optional custom output filename

        Returns:
            Path to generated document
        """
        with open(json_path, 'r') as f:
            data = json.load(f)

        # Handle nested JSON structures
        flattened_data = self._flatten_json(data)

        return self.populate_template(template_name, flattened_data, output_filename)

    def _flatten_json(self, data: Dict[str, Any], parent_key: str = '') -> Dict[str, str]:
        """Flatten nested JSON structure"""
        items = []
        for k, v in data.items():
            new_key = f"{parent_key}_{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(self._flatten_json(v, new_key).items())
            elif isinstance(v, list):
                for i, item in enumerate(v):
                    if isinstance(item, dict):
                        items.extend(self._flatten_json(item, f"{new_key}_{i}").items())
                    else:
                        items.append((f"{new_key}_{i}", str(item)))
            else:
                items.append((new_key, str(v) if v is not None else ""))
        return dict(items)

    def generate_sample_data(self, template_name: str) -> Dict[str, str]:
        """
        Generate sample data for testing

        Returns complete data set with all variables populated
        """
        # Base data common to all templates
        base_data = {
            # Contact Information
            "first_name": "John",
            "last_name": "Doe",
            "street_address": "123 Main Street",
            "city": "New York",
            "state": "NY",
            "zip_code": "10001",
            "phone": "(555) 123-4567",
            "email": "john.doe@email.com",
            "linkedin_url": "linkedin.com/in/johndoe",

            # Section Headers (use defaults)
            "profile_section_header": "Profile",
            "experience_section_header": "Experience",
            "education_section_header": "Education",
            "skills_section_header": "Skills",
            "interests_section_header": "Interests",
            "contact_header": "Contact",
            "education_header": "Education",
            "experience_header": "Experience",
            "skills_header": "Skills",
        }

        if template_name == "restaurant_manager":
            base_data.update({
                # Career Overview
                "career_overview_1": "Dynamic and results-driven restaurant manager with proven leadership abilities.",
                "career_overview_2": "Expert in operations management with passion for culinary excellence.",
                "career_overview_3": "Track record of building high-performing teams and exceeding revenue targets.",
                "career_overview_4": "Consistently achieve 15% year-over-year growth.",
                "career_overview_5": "Master of customer experience and retention strategies.",

                # Job 1
                "position_1": "Senior Restaurant Manager",
                "company_1": "Fine Dining Group",
                "start_date_1": "January 2020",
                "end_date_1": "Present",
                "job_1_responsibility_1": "Lead team of 40+ staff members across front and back of house operations.",
                "job_1_achievement_1": "Reduced operational costs by 12% through strategic vendor negotiations.",
                "job_1_achievement_2": "Increased customer satisfaction scores from 4.2 to 4.8 stars.",

                # Job 2
                "position_2": "Restaurant Manager",
                "company_2": "Casual Dining Chain",
                "start_date_2": "March 2018",
                "end_date_2": "December 2019",
                "job_2_responsibility_1": "Implemented comprehensive training program for new hires.",
                "job_2_achievement_1": "Grew social media following by 45% through targeted campaigns.",
                "job_2_achievement_2": "Achieved highest health inspection score in district (98/100).",
                "job_2_achievement_3": "Reduced food waste by 20% through inventory optimization.",

                # Education
                "degree_1": "B.S. in Hospitality Management",
                "graduation_date_1": "May 2018",
                "institution_1": "Cornell University",
                "institution_1_city": "Ithaca",
                "institution_1_state": "New York",
                "degree_2": "A.A. in Culinary Arts",
                "graduation_date_2": "May 2016",
                "institution_2": "Culinary Institute of America",
                "institution_2_city": "Hyde Park",
                "institution_2_state": "New York",

                # Skills
                "skill_1": "P&L Management",
                "skill_2": "Toast POS Expert",
                "skill_3": "Team Leadership",
                "skill_4": "Crisis Management",
                "skill_5": "Menu Development",
                "skill_6": "Customer Relations",

                # Interests
                "interest_1": "Wine Sommelier Certification",
                "interest_2": "Sustainable Farming",
                "interest_3": "Food Photography",
                "interest_4": "Marathon Running",
                "interest_5": "Jazz Music",
                "interest_6": "International Travel",
            })

        elif template_name == "accountant":
            base_data.update({
                # Professional Summary
                "professional_summary_1": "CPA with 8+ years of experience in public accounting and financial reporting.",
                "professional_summary_2": "Expert in tax planning and compliance with strong analytical skills.",

                # Education
                "degree_1": "Bachelor of Science in Accounting",
                "minor_1": "Finance",
                "institution_1": "State University",
                "degree_label": "Degree",
                "graduation_date_1": "May 2016",

                # Experience
                "position_1": "Senior Accountant",
                "company_1": "Big Four Accounting Firm",
                "job_1_city": "San Francisco",
                "job_1_state": "CA",
                "start_date_1": "June 2018",
                "end_date_1": "Present",

                # Skills
                "technical_skill_1": "QuickBooks Pro",
                "technical_skill_2": "Advanced Excel",
                "technical_skill_3": "Tax Compliance",
                "technical_skill_4": "Financial Analysis",
                "soft_skill_1": "Client Relations",
                "language_skill_1": "Spanish (Fluent)",
            })

        elif template_name == "uiux_designer":
            base_data.update({
                # Professional
                "current_title": "Senior UI/UX Designer",

                # Contact
                "portfolio_website": "www.johndoe-design.com",

                # Education
                "institution_1": "Parsons School of Design",
                "degree_type_1": "MFA",
                "major_1": "Digital Design",
                "graduation_year_1": "2019",

                # Skills
                "skill_2": "User Research",
                "skill_3": "Wireframing",
                "skill_4": "Prototyping",

                # Experience
                "position_2": "UI/UX Designer",
                "company_1": "Tech Startup Inc",
            })

        return base_data

    def create_sample_json(self, output_path: str = "./data/sample_data.json"):
        """Create sample JSON file with all variables"""

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        all_data = {}
        for template_name in self.templates.keys():
            all_data[template_name] = self.generate_sample_data(template_name)

        with open(output_path, 'w') as f:
            json.dump(all_data, f, indent=2)

        logger.info(f"Sample data saved to: {output_path}")
        return output_path


def main():
    """Main function demonstrating template population"""

    inserter = TemplateVariableInserter()

    print("="*60)
    print("TEMPLATE VARIABLE INSERTION SYSTEM")
    print("="*60)
    print()

    # Show available templates
    print("Available Templates:")
    for name, info in inserter.templates.items():
        print(f"  • {name}: {len(info['variables'])} variables")
    print()

    # Generate sample data file
    print("Generating sample data file...")
    sample_file = inserter.create_sample_json()
    print(f"  ✓ Created: {sample_file}")
    print()

    # Test with sample data
    print("Testing templates with sample data...")
    for template_name in inserter.templates.keys():
        print(f"\nProcessing: {template_name}")

        try:
            # Generate sample data
            sample_data = inserter.generate_sample_data(template_name)

            # Populate template
            output_path = inserter.populate_template(
                template_name,
                sample_data,
                f"{template_name}_sample.docx"
            )

            print(f"  ✓ Generated: {output_path}")

        except Exception as e:
            print(f"  ✗ Error: {str(e)}")

    print("\n" + "="*60)
    print("Template insertion complete!")
    print(f"Output directory: {inserter.output_dir}")


if __name__ == "__main__":
    main()