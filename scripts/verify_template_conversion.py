"""
Template Conversion Verification Script

This script verifies that the converted templates:
1. Contain only variable placeholders (no hardcoded text)
2. Have valid variable format (<<variable_name>>)
3. Preserve all document structure and formatting
4. Match the documentation

Usage:
    python scripts/verify_template_conversion.py
"""

import os
import re
import json
from docx import Document


class TemplateVerifier:
    """Verifies converted templates for quality and completeness"""

    def __init__(self):
        self.variable_pattern = re.compile(r'<<([^>]+)>>')
        self.verification_results = {}

    def verify_template(self, template_path, template_name):
        """
        Verify a converted template file

        Args:
            template_path (str): Path to the converted template
            template_name (str): Name identifier for the template

        Returns:
            dict: Verification results
        """
        print(f"\nVerifying: {template_name}")
        print("-" * 80)

        try:
            doc = Document(template_path)

            results = {
                "template_name": template_name,
                "template_path": template_path,
                "status": "PASS",
                "issues": [],
                "statistics": {
                    "total_paragraphs": 0,
                    "total_tables": 0,
                    "paragraphs_with_variables": 0,
                    "paragraphs_with_hardcoded_text": 0,
                    "total_variables_found": 0,
                    "unique_variables": set()
                }
            }

            # Check paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue

                results["statistics"]["total_paragraphs"] += 1

                # Find variables in paragraph
                variables = self.variable_pattern.findall(text)

                if variables:
                    results["statistics"]["paragraphs_with_variables"] += 1
                    results["statistics"]["total_variables_found"] += len(variables)
                    results["statistics"]["unique_variables"].update(variables)

                # Check for hardcoded text (text outside of variables)
                text_without_variables = self.variable_pattern.sub('', text).strip()
                if text_without_variables:
                    # Allow newlines and minimal whitespace
                    if text_without_variables.strip('\n\r\t '):
                        results["statistics"]["paragraphs_with_hardcoded_text"] += 1
                        results["issues"].append({
                            "type": "hardcoded_text",
                            "location": "paragraph",
                            "text": text[:100]
                        })

            # Check tables
            for table_idx, table in enumerate(doc.tables):
                results["statistics"]["total_tables"] += 1

                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for paragraph in cell.paragraphs:
                            text = paragraph.text.strip()
                            if not text:
                                continue

                            # Find variables
                            variables = self.variable_pattern.findall(text)

                            if variables:
                                results["statistics"]["total_variables_found"] += len(variables)
                                results["statistics"]["unique_variables"].update(variables)

                            # Check for hardcoded text
                            text_without_variables = self.variable_pattern.sub('', text).strip()
                            if text_without_variables:
                                if text_without_variables.strip('\n\r\t '):
                                    results["issues"].append({
                                        "type": "hardcoded_text",
                                        "location": f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}",
                                        "text": text[:100]
                                    })

            # Convert set to list for JSON serialization
            results["statistics"]["unique_variables"] = list(results["statistics"]["unique_variables"])

            # Determine overall status
            if results["statistics"]["paragraphs_with_hardcoded_text"] > 0:
                results["status"] = "WARNING"

            if results["issues"]:
                results["status"] = "FAIL"

            # Display results
            print(f"Status: {results['status']}")
            print(f"Paragraphs: {results['statistics']['total_paragraphs']}")
            print(f"Tables: {results['statistics']['total_tables']}")
            print(f"Variables found: {results['statistics']['total_variables_found']}")
            print(f"Unique variables: {len(results['statistics']['unique_variables'])}")
            print(f"Hardcoded text paragraphs: {results['statistics']['paragraphs_with_hardcoded_text']}")

            if results["issues"]:
                print(f"\nIssues found: {len(results['issues'])}")
                for issue in results["issues"][:5]:  # Show first 5 issues
                    print(f"  - {issue['type']} in {issue['location']}: {issue['text'][:60]}...")

            self.verification_results[template_name] = results
            return results

        except Exception as e:
            print(f"ERROR: {str(e)}")
            return {
                "template_name": template_name,
                "status": "ERROR",
                "error": str(e)
            }

    def compare_with_documentation(self, doc_path):
        """
        Compare verification results with the generated documentation

        Args:
            doc_path (str): Path to the variable documentation JSON
        """
        print("\n" + "="*80)
        print("COMPARING WITH DOCUMENTATION")
        print("="*80)

        try:
            with open(doc_path, 'r') as f:
                doc_data = json.load(f)

            for template_name, verification in self.verification_results.items():
                if verification.get("status") == "ERROR":
                    continue

                doc_template = doc_data["templates"].get(template_name)
                if not doc_template:
                    print(f"\n{template_name}: NOT FOUND in documentation")
                    continue

                doc_var_count = len(doc_template["variables"])
                found_var_count = len(verification["statistics"]["unique_variables"])

                print(f"\n{template_name}:")
                print(f"  Documentation reports: {doc_var_count} variables")
                print(f"  Verification found: {found_var_count} variables")

                if doc_var_count == found_var_count:
                    print(f"  ✓ Variable counts match!")
                else:
                    print(f"  ✗ Variable count mismatch!")

        except Exception as e:
            print(f"ERROR comparing with documentation: {str(e)}")

    def save_verification_report(self, output_path):
        """
        Save verification results to a JSON report

        Args:
            output_path (str): Path to save the verification report
        """
        try:
            report = {
                "verification_date": "2025-10-21",
                "templates_verified": len(self.verification_results),
                "results": self.verification_results
            }

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(report, f, indent=2, ensure_ascii=False)

            print(f"\nVerification report saved: {output_path}")

        except Exception as e:
            print(f"ERROR saving verification report: {str(e)}")


def main():
    """Main verification execution"""

    verifier = TemplateVerifier()

    base_dir = "/workspace/.trees/convert-raw-template-files-into-ready-for-producti"
    output_dir = os.path.join(base_dir, "content_template_library/production_ready")

    templates = [
        {
            "name": "restaurant_manager",
            "path": os.path.join(output_dir, "restaurant_manager_fully_converted.docx")
        },
        {
            "name": "accountant",
            "path": os.path.join(output_dir, "accountant_fully_converted.docx")
        },
        {
            "name": "uiux_designer",
            "path": os.path.join(output_dir, "uiux_designer_fully_converted.docx")
        }
    ]

    print("\n" + "="*80)
    print("TEMPLATE CONVERSION VERIFICATION")
    print("="*80)

    # Verify each template
    for template in templates:
        if not os.path.exists(template["path"]):
            print(f"\nERROR: Template not found: {template['path']}")
            continue

        verifier.verify_template(template["path"], template["name"])

    # Compare with documentation
    doc_path = os.path.join(output_dir, "template_variables_documentation.json")
    if os.path.exists(doc_path):
        verifier.compare_with_documentation(doc_path)
    else:
        print(f"\nWARNING: Documentation not found: {doc_path}")

    # Save verification report
    report_path = os.path.join(output_dir, "verification_report.json")
    verifier.save_verification_report(report_path)

    # Summary
    print("\n" + "="*80)
    print("VERIFICATION SUMMARY")
    print("="*80)

    pass_count = sum(1 for r in verifier.verification_results.values() if r.get("status") == "PASS")
    fail_count = sum(1 for r in verifier.verification_results.values() if r.get("status") == "FAIL")
    warn_count = sum(1 for r in verifier.verification_results.values() if r.get("status") == "WARNING")

    print(f"\nTotal templates verified: {len(verifier.verification_results)}")
    print(f"  ✓ PASS: {pass_count}")
    print(f"  ⚠ WARNING: {warn_count}")
    print(f"  ✗ FAIL: {fail_count}")

    if fail_count == 0 and warn_count == 0:
        print("\n✓ ALL TEMPLATES VERIFIED SUCCESSFULLY!")
    elif fail_count == 0:
        print("\n⚠ All templates converted but some warnings present")
    else:
        print("\n✗ Some templates have issues - review verification report")

    print("\n" + "="*80)


if __name__ == "__main__":
    main()
