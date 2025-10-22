"""
Comprehensive Template Conversion Script for Production-Ready Templates

This script converts raw template files into fully variablized production templates
by replacing ALL text content with semantic variable placeholders. It processes
both paragraphs and table cells with intelligent semantic naming.

Features:
- Replaces all text with semantic variables (no hardcoded content)
- Uses semantic naming patterns (career_overview_1, job_1_responsibility_1, etc.)
- Processes both paragraphs and table cells
- Handles multi-line text in cells by breaking into separate variables
- Generates comprehensive documentation of all variables created
- Preserves all formatting and document structure

Usage:
    python scripts/convert_templates_to_production.py

Output:
    - restaurant_manager_fully_converted.docx
    - accountant_fully_converted.docx
    - uiux_designer_fully_converted.docx
    - template_variables_documentation.json
"""

import os
import re
import json
import logging
from datetime import datetime
from docx import Document
from collections import defaultdict


class ProductionTemplateConverter:
    """
    Converts raw templates into production-ready templates with full variable coverage

    This converter replaces ALL text content with semantic variables, ensuring
    no hardcoded text remains. It uses intelligent naming based on content type
    and position within the document.
    """

    def __init__(self):
        """Initialize the converter with logging and tracking structures"""
        self.setup_logging()
        self.variable_documentation = {}
        self.section_counters = defaultdict(int)

    def setup_logging(self):
        """Configure logging for conversion process"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)

    def convert_template_to_production(self, template_path, output_path, template_name):
        """
        Convert a template file by replacing ALL text with semantic variables

        Args:
            template_path (str): Path to the input template file
            output_path (str): Path where the converted template will be saved
            template_name (str): Name identifier for the template (for variable naming)

        Returns:
            dict: Conversion statistics and variable mapping
        """
        try:
            self.logger.info(f"Converting template: {template_path}")
            self.logger.info(f"Output will be saved to: {output_path}")

            # Load the template document
            doc = Document(template_path)

            # Reset counters for this template
            self.section_counters = defaultdict(int)
            template_variables = {}

            # Track conversion statistics
            conversion_stats = {
                "template_name": template_name,
                "total_paragraphs_processed": 0,
                "total_tables_processed": 0,
                "total_cells_processed": 0,
                "total_variables_created": 0,
                "conversion_date": datetime.now().isoformat()
            }

            # Process all paragraphs in the document
            self.logger.info("Processing paragraphs...")
            for para_idx, paragraph in enumerate(doc.paragraphs):
                original_text = paragraph.text.strip()

                if not original_text:
                    continue

                conversion_stats["total_paragraphs_processed"] += 1

                # Determine semantic variable name based on content
                variable_name, content_type = self.generate_semantic_variable(
                    original_text, template_name, para_idx
                )

                # Update paragraph with variable placeholder
                self.update_paragraph_text(paragraph, f"<<{variable_name}>>")

                # Track variable documentation
                template_variables[variable_name] = {
                    "original_text": original_text,
                    "content_type": content_type,
                    "location": "paragraph",
                    "position": para_idx
                }
                conversion_stats["total_variables_created"] += 1

                self.logger.debug(f"Para {para_idx}: '{original_text[:50]}...' -> <<{variable_name}>>")

            # Process all tables in the document
            self.logger.info("Processing tables...")
            for table_idx, table in enumerate(doc.tables):
                conversion_stats["total_tables_processed"] += 1

                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        # Process each paragraph within the cell
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            original_text = paragraph.text.strip()

                            if not original_text:
                                continue

                            conversion_stats["total_cells_processed"] += 1

                            # For multi-line content in cells, break into separate variables
                            lines = [line.strip() for line in original_text.split('\n') if line.strip()]

                            if len(lines) > 1:
                                # Multiple lines - create separate variables for each
                                new_text_parts = []
                                for line_idx, line in enumerate(lines):
                                    variable_name, content_type = self.generate_semantic_variable(
                                        line, template_name,
                                        f"table{table_idx}_row{row_idx}_cell{cell_idx}_line{line_idx}"
                                    )
                                    new_text_parts.append(f"<<{variable_name}>>")

                                    template_variables[variable_name] = {
                                        "original_text": line,
                                        "content_type": content_type,
                                        "location": f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}",
                                        "line": line_idx
                                    }
                                    conversion_stats["total_variables_created"] += 1

                                # Join with newlines to preserve structure
                                self.update_paragraph_text(paragraph, '\n'.join(new_text_parts))

                            else:
                                # Single line in cell
                                variable_name, content_type = self.generate_semantic_variable(
                                    original_text, template_name,
                                    f"table{table_idx}_row{row_idx}_cell{cell_idx}"
                                )

                                self.update_paragraph_text(paragraph, f"<<{variable_name}>>")

                                template_variables[variable_name] = {
                                    "original_text": original_text,
                                    "content_type": content_type,
                                    "location": f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}",
                                    "position": para_idx
                                }
                                conversion_stats["total_variables_created"] += 1

            # Save the converted template
            doc.save(output_path)
            self.logger.info(f"Template saved successfully: {output_path}")

            # Store variable documentation for this template
            self.variable_documentation[template_name] = {
                "variables": template_variables,
                "statistics": conversion_stats
            }

            return conversion_stats

        except Exception as e:
            self.logger.error(f"Error converting template: {str(e)}")
            raise

    def generate_semantic_variable(self, text, template_name, position_hint):
        """
        Generate a semantic variable name based on text content and context

        This method analyzes the text to determine its purpose and assigns
        an appropriate semantic variable name.

        Args:
            text (str): The original text content
            template_name (str): Name of the template being processed
            position_hint: Position or context hint for variable naming

        Returns:
            tuple: (variable_name, content_type)
        """
        # Normalize text for analysis
        text_lower = text.lower()
        text_clean = text.strip()

        # Personal information patterns
        if re.match(r'^[A-Z][a-z]+\s+[A-Z][a-z]+$', text_clean):
            return "full_name", "personal_info"

        if re.search(r'@', text):
            return "email", "contact"

        if re.search(r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}', text):
            return "phone_number", "contact"

        if re.search(r'linkedin\.com', text_lower):
            return "linkedin_url", "contact"

        if re.search(r'calendly\.com', text_lower):
            return "calendly_url", "contact"

        if re.search(r'(portfolio|website|github)', text_lower):
            return "portfolio_url", "contact"

        # Location patterns
        if re.search(r',\s*(BC|ON|AB|QC|MB|SK|NS|NB|PE|NL|YT|NT|NU)', text):
            return "location", "personal_info"

        # Date patterns
        if re.search(r'(january|february|march|april|may|june|july|august|september|october|november|december|\d{4})', text_lower):
            if 'present' in text_lower or 'current' in text_lower:
                return "date_range_current", "date"
            return "date_range", "date"

        # Profile/Summary section detection
        if any(keyword in text_lower for keyword in ['passionate', 'experienced', 'dedicated', 'professional', 'skilled']):
            self.section_counters['career_overview'] += 1
            return f"career_overview_{self.section_counters['career_overview']}", "profile"

        if any(keyword in text_lower for keyword in ['summary', 'profile', 'objective']):
            self.section_counters['professional_summary'] += 1
            return f"professional_summary_{self.section_counters['professional_summary']}", "summary"

        # Job title/position detection
        if re.search(r'^[A-Z][a-zA-Z\s]+$', text_clean) and len(text_clean.split()) <= 5:
            if any(keyword in text_lower for keyword in ['manager', 'director', 'coordinator', 'specialist', 'analyst', 'developer', 'designer', 'engineer', 'accountant']):
                self.section_counters['job_title'] += 1
                return f"job_{self.section_counters['job_title']}_title", "job_title"

        # Company/Organization name
        if re.search(r'^[A-Z][a-zA-Z\s&,\.\-]+$', text_clean) and len(text_clean.split()) <= 6:
            if not any(keyword in text_lower for keyword in ['i ', 'my ', 'the ', 'and ', 'or ']):
                self.section_counters['company'] += 1
                return f"job_{self.section_counters['company']}_company", "company"

        # Work experience bullets/responsibilities
        if text.startswith(('•', '-', '*')) or (len(text.split()) > 5 and '.' in text):
            # Determine if this is a responsibility or achievement
            if any(keyword in text_lower for keyword in ['increased', 'reduced', 'improved', 'achieved', 'generated', 'saved']):
                self.section_counters['achievement'] += 1
                job_num = (self.section_counters['achievement'] // 5) + 1  # Group by job (assuming ~5 per job)
                achievement_num = (self.section_counters['achievement'] % 5) + 1
                return f"job_{job_num}_achievement_{achievement_num}", "achievement"
            else:
                self.section_counters['responsibility'] += 1
                job_num = (self.section_counters['responsibility'] // 5) + 1  # Group by job
                resp_num = (self.section_counters['responsibility'] % 5) + 1
                return f"job_{job_num}_responsibility_{resp_num}", "responsibility"

        # Education section
        if any(keyword in text_lower for keyword in ['university', 'college', 'school', 'institute']):
            self.section_counters['education'] += 1
            return f"education_{self.section_counters['education']}_institution", "education"

        if any(keyword in text_lower for keyword in ['bachelor', 'master', 'diploma', 'degree', 'certificate']):
            self.section_counters['degree'] += 1
            return f"education_{self.section_counters['degree']}_degree", "education"

        # Skills section
        if any(keyword in text_lower for keyword in ['skills', 'proficient', 'expertise', 'knowledge']):
            self.section_counters['skill'] += 1
            return f"skill_{self.section_counters['skill']}", "skills"

        # Section headers
        if text.isupper() or (len(text.split()) <= 3 and text_clean.istitle()):
            section_type = text_lower.replace(' ', '_')
            return f"section_header_{section_type}", "section_header"

        # Generic content - use position-based naming
        self.section_counters['generic'] += 1
        return f"content_{self.section_counters['generic']}", "generic_content"

    def update_paragraph_text(self, paragraph, new_text):
        """
        Update paragraph text while preserving formatting

        Args:
            paragraph: python-docx paragraph object
            new_text (str): New text to set
        """
        # Clear existing text while preserving formatting
        for run in paragraph.runs:
            run.text = ""

        # Add new text to the first run (or create one if none exist)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

    def save_variable_documentation(self, output_path):
        """
        Save comprehensive documentation of all variables created

        Args:
            output_path (str): Path to save the JSON documentation
        """
        try:
            # Create a structured documentation format
            documentation = {
                "generation_date": datetime.now().isoformat(),
                "total_templates_processed": len(self.variable_documentation),
                "templates": self.variable_documentation,
                "variable_categories": self.generate_variable_categories(),
                "usage_guide": {
                    "description": "This file documents all variables created during template conversion",
                    "variable_format": "<<variable_name>>",
                    "naming_convention": {
                        "profile": "career_overview_N or professional_summary_N",
                        "work_experience": "job_N_title, job_N_company, job_N_responsibility_N, job_N_achievement_N",
                        "education": "education_N_institution, education_N_degree",
                        "skills": "skill_N",
                        "contact": "email, phone_number, linkedin_url, calendly_url, portfolio_url",
                        "dates": "date_range, date_range_current"
                    }
                }
            }

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(documentation, f, indent=2, ensure_ascii=False)

            self.logger.info(f"Variable documentation saved: {output_path}")

        except Exception as e:
            self.logger.error(f"Error saving documentation: {str(e)}")
            raise

    def generate_variable_categories(self):
        """
        Generate a summary of variables by category across all templates

        Returns:
            dict: Variables organized by content type
        """
        categories = defaultdict(list)

        for template_name, template_data in self.variable_documentation.items():
            for var_name, var_info in template_data["variables"].items():
                content_type = var_info.get("content_type", "unknown")
                categories[content_type].append({
                    "variable": var_name,
                    "template": template_name,
                    "original_text": var_info.get("original_text", "")[:100]  # Truncate for readability
                })

        return dict(categories)


def main():
    """
    Main execution function to convert all three templates
    """
    # Initialize converter
    converter = ProductionTemplateConverter()

    # Define input and output paths
    base_dir = "/workspace/.trees/convert-raw-template-files-into-ready-for-producti"
    input_dir = os.path.join(base_dir, "content_template_library/manual_converted")
    output_dir = os.path.join(base_dir, "content_template_library/production_ready")

    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Define templates to convert
    templates = [
        {
            "name": "restaurant_manager",
            "input": os.path.join(input_dir, "restaurant_manager_template.docx"),
            "output": os.path.join(output_dir, "restaurant_manager_fully_converted.docx")
        },
        {
            "name": "accountant",
            "input": os.path.join(input_dir, "accountant_template.docx"),
            "output": os.path.join(output_dir, "accountant_fully_converted.docx")
        },
        {
            "name": "uiux_designer",
            "input": os.path.join(input_dir, "uiux_designer_template.docx"),
            "output": os.path.join(output_dir, "uiux_designer_fully_converted.docx")
        }
    ]

    # Convert each template
    print("\n" + "="*80)
    print("COMPREHENSIVE TEMPLATE CONVERSION TO PRODUCTION")
    print("="*80 + "\n")

    conversion_results = []

    for template in templates:
        print(f"\nProcessing: {template['name']}")
        print(f"Input:  {template['input']}")
        print(f"Output: {template['output']}")
        print("-" * 80)

        try:
            if not os.path.exists(template['input']):
                print(f"ERROR: Template not found: {template['input']}")
                continue

            # Convert the template
            stats = converter.convert_template_to_production(
                template_path=template['input'],
                output_path=template['output'],
                template_name=template['name']
            )

            # Display conversion statistics
            print(f"\nConversion Statistics for {template['name']}:")
            print(f"  - Paragraphs processed: {stats['total_paragraphs_processed']}")
            print(f"  - Tables processed: {stats['total_tables_processed']}")
            print(f"  - Cells processed: {stats['total_cells_processed']}")
            print(f"  - Variables created: {stats['total_variables_created']}")
            print(f"  - Output saved: {template['output']}")

            conversion_results.append({
                "template": template['name'],
                "status": "SUCCESS",
                "stats": stats
            })

        except Exception as e:
            print(f"ERROR converting {template['name']}: {str(e)}")
            conversion_results.append({
                "template": template['name'],
                "status": "FAILED",
                "error": str(e)
            })

    # Save comprehensive variable documentation
    print("\n" + "="*80)
    print("GENERATING VARIABLE DOCUMENTATION")
    print("="*80 + "\n")

    doc_output_path = os.path.join(output_dir, "template_variables_documentation.json")
    converter.save_variable_documentation(doc_output_path)

    print(f"Documentation saved to: {doc_output_path}")

    # Summary report
    print("\n" + "="*80)
    print("CONVERSION SUMMARY")
    print("="*80 + "\n")

    for result in conversion_results:
        status_symbol = "✓" if result['status'] == "SUCCESS" else "✗"
        print(f"{status_symbol} {result['template']}: {result['status']}")
        if result['status'] == "SUCCESS":
            print(f"  Variables: {result['stats']['total_variables_created']}")

    print("\n" + "="*80)
    print("CONVERSION COMPLETE")
    print("="*80 + "\n")

    return conversion_results


if __name__ == "__main__":
    main()
