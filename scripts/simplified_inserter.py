#!/usr/bin/env python3
"""
Simplified Template Variable Inserter
Handles simplified unified variable naming (job_X_experience_Y)
"""

import json
import os
from datetime import datetime
from docx import Document
from typing import Dict, List, Optional
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class SimplifiedInserter:
    """Inserts variables using simplified naming scheme"""

    def __init__(self):
        self.template_dir = "./content_template_library/manual_converted"
        self.output_dir = "./content_template_library/generated"
        os.makedirs(self.output_dir, exist_ok=True)

        self.templates = {
            "restaurant_manager": {
                "file": "restaurant_manager_final.docx",
                "variables": self._get_restaurant_variables()
            },
            "accountant": {
                "file": "accountant_final.docx",
                "variables": self._get_accountant_variables()
            },
            "uiux_designer": {
                "file": "uiux_designer_final.docx",
                "variables": self._get_uiux_variables()
            }
        }

    def _get_restaurant_variables(self) -> List[str]:
        """Restaurant Manager template variables (simplified)

        Note: 'state' is optional - removed from fixed template for Canadian users
        """
        return [
            # Contact
            "first_name", "last_name", "street_address", "city",
            "zip_code", "phone", "email", "linkedin_url",

            # Headers
            "profile_header", "experience_header", "education_header",
            "skills_header", "interests_header",

            # Career Overview
            "career_overview_1", "career_overview_2", "career_overview_3",
            "career_overview_4", "career_overview_5",

            # Job 1 - SIMPLIFIED (all experience_X)
            "position_1", "company_1", "start_date_1", "end_date_1",
            "job_1_experience_1", "job_1_experience_2", "job_1_experience_3",

            # Job 2 - SIMPLIFIED (all experience_X)
            "position_2", "company_2", "start_date_2", "end_date_2",
            "job_2_experience_1", "job_2_experience_2",
            "job_2_experience_3", "job_2_experience_4",

            # Education
            "degree_1", "graduation_date_1", "institution_1",
            "institution_city_1", "institution_state_1",
            "degree_2", "graduation_date_2", "institution_2",
            "institution_city_2", "institution_state_2",

            # Skills
            "skill_1", "skill_2", "skill_3", "skill_4", "skill_5", "skill_6",

            # Interests
            "interest_1", "interest_2", "interest_3",
            "interest_4", "interest_5", "interest_6"
        ]

    def _get_accountant_variables(self) -> List[str]:
        """Accountant template variables (simplified)"""
        return [
            # Contact
            "first_name", "last_name", "street_address", "city", "state",
            "zip_code", "phone", "email", "linkedin_url",

            # Professional Summary
            "professional_summary_1", "professional_summary_2",

            # Headers
            "education_header", "experience_header", "skills_header",

            # Education
            "degree_1", "minor_1", "institution_1",
            "degree_label", "graduation_date_1",

            # Experience
            "position_1", "company_1", "job_city_1", "job_state_1",
            "start_date_1", "end_date_1",

            # Skills (simplified - just numbered)
            "skill_1", "skill_2", "skill_3", "skill_4", "skill_5", "skill_6"
        ]

    def _get_uiux_variables(self) -> List[str]:
        """UI/UX Designer template variables (simplified)"""
        return [
            # Personal
            "first_name", "last_name", "current_title",

            # Contact
            "contact_header", "email", "portfolio_website", "phone", "city",

            # Professional Bio
            "professional_bio_1", "professional_bio_2",

            # Education
            "education_header", "institution_1", "degree_1",
            "major_1", "graduation_year_1",

            # Skills
            "skills_header", "skill_1", "skill_2", "skill_3", "skill_4",

            # Experience Headers
            "experience_header",

            # Job 1 - SIMPLIFIED
            "position_1", "company_1", "start_date_1", "end_date_1",
            "job_1_experience_1", "job_1_experience_2",
            "job_1_experience_3", "job_1_experience_4",

            # Job 2 - SIMPLIFIED
            "position_2", "company_2", "start_date_2", "end_date_2",
            "job_2_experience_1", "job_2_experience_2", "job_2_experience_3"
        ]

    def populate_template(
        self,
        template_name: str,
        data: Dict[str, str],
        output_filename: Optional[str] = None
    ) -> str:
        """Populate template with data"""

        if template_name not in self.templates:
            raise ValueError(f"Unknown template: {template_name}")

        template_info = self.templates[template_name]
        template_path = os.path.join(self.template_dir, template_info["file"])

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")

        logger.info(f"Loading: {template_path}")
        doc = Document(template_path)

        replacements_made = {}
        template_variables = template_info["variables"]

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for var_name in template_variables:
                placeholder = f"<<{var_name}>>"
                if placeholder in paragraph.text and var_name in data and data[var_name]:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, data[var_name])
                            replacements_made[var_name] = data[var_name]

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for var_name in template_variables:
                            placeholder = f"<<{var_name}>>"
                            if placeholder in paragraph.text and var_name in data and data[var_name]:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, data[var_name])
                                        replacements_made[var_name] = data[var_name]

        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{template_name}_{timestamp}.docx"

        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)

        logger.info(f"Saved: {output_path}")
        logger.info(f"Variables replaced: {len(replacements_made)}/{len(template_variables)}")

        missing = set(template_variables) - set(replacements_made.keys())
        if missing:
            logger.warning(f"Missing ({len(missing)}): {list(missing)[:5]}")

        return output_path

    def generate_sample_data(self, template_name: str) -> Dict[str, str]:
        """Generate sample data with simplified variables"""

        base = {
            # Contact
            "first_name": "John",
            "last_name": "Doe",
            "street_address": "123 Main Street",
            "city": "New York",
            "state": "NY",
            "zip_code": "10001",
            "phone": "(555) 123-4567",
            "email": "john.doe@email.com",
            "linkedin_url": "linkedin.com/in/johndoe",

            # Headers (defaults)
            "profile_header": "Profile",
            "experience_header": "Experience",
            "education_header": "Education",
            "skills_header": "Skills",
            "interests_header": "Interests",
            "contact_header": "Contact",
        }

        if template_name == "restaurant_manager":
            base.update({
                # Career Overview (5 parts)
                "career_overview_1": "Dynamic restaurant leader with 15+ years of experience in high-volume operations.",
                "career_overview_2": "Passionate about creating exceptional dining experiences and developing talent.",
                "career_overview_3": "Expert in operational efficiency, cost control, and customer service excellence.",
                "career_overview_4": "Consistently achieve 20% revenue growth and maintain 95% customer satisfaction.",
                "career_overview_5": "Specialized in farm-to-table concepts and sustainable practices.",

                # Job 1 (unified experience bullets)
                "position_1": "General Manager",
                "company_1": "The Urban Table Group",
                "start_date_1": "March 2019",
                "end_date_1": "Present",
                "job_1_experience_1": "Oversee all operations for 3 restaurant locations with $12M annual revenue.",
                "job_1_experience_2": "Reduced labor costs by 15% through optimized scheduling and cross-training.",
                "job_1_experience_3": "Increased average check size by 22% through menu engineering and staff training.",

                # Job 2 (unified experience bullets)
                "position_2": "Restaurant Manager",
                "company_2": "Harvest & Hearth",
                "start_date_2": "January 2016",
                "end_date_2": "February 2019",
                "job_2_experience_1": "Launched successful brunch program increasing weekend revenue by 40%.",
                "job_2_experience_2": "Grew Instagram following from 500 to 15,000 through content strategy.",
                "job_2_experience_3": "Achieved perfect health department scores for 3 consecutive years.",
                "job_2_experience_4": "Reduced food waste by 25% saving $50K annually through new inventory system.",

                # Education
                "degree_1": "B.S. in Hospitality Management",
                "graduation_date_1": "May 2015",
                "institution_1": "Cornell University",
                "institution_city_1": "Ithaca",
                "institution_state_1": "New York",
                "degree_2": "A.A. in Culinary Arts",
                "graduation_date_2": "May 2013",
                "institution_2": "Culinary Institute",
                "institution_city_2": "Hyde Park",
                "institution_state_2": "New York",

                # Skills
                "skill_1": "P&L Management",
                "skill_2": "Toast POS Systems",
                "skill_3": "Team Leadership",
                "skill_4": "Crisis Management",
                "skill_5": "Multi-Unit Operations",
                "skill_6": "Guest Experience",

                # Interests
                "interest_1": "Sommelier Certification",
                "interest_2": "Sustainable Agriculture",
                "interest_3": "Food Photography",
                "interest_4": "Marathon Running",
                "interest_5": "International Cuisine",
                "interest_6": "Culinary Travel",
            })

        elif template_name == "accountant":
            base.update({
                # Professional Summary
                "professional_summary_1": "CPA with 10+ years of experience in public accounting and financial reporting.",
                "professional_summary_2": "Expert in implementing cost-saving measures that improve accuracy by 40%.",

                # Education
                "degree_1": "Bachelor of Science in Accounting",
                "minor_1": "Finance",
                "institution_1": "State University",
                "degree_label": "Degree",
                "graduation_date_1": "May 2015",

                # Experience
                "position_1": "Senior Accountant",
                "company_1": "Big Four Firm",
                "job_city_1": "San Francisco",
                "job_state_1": "CA",
                "start_date_1": "January 2018",
                "end_date_1": "Present",

                # Skills
                "skill_1": "QuickBooks Pro",
                "skill_2": "Advanced Excel",
                "skill_3": "GAAP Compliance",
                "skill_4": "Financial Analysis",
                "skill_5": "Client Relations",
                "skill_6": "Spanish (Fluent)",
            })

        elif template_name == "uiux_designer":
            base.update({
                "current_title": "Senior UI/UX Designer",
                "portfolio_website": "www.johndoe-design.com",

                # Professional Bio
                "professional_bio_1": "Passionate designer creating beautiful and functional digital experiences.",
                "professional_bio_2": "User research and data-driven design decisions to solve real problems.",

                # Education
                "institution_1": "Parsons School of Design",
                "degree_1": "MFA",
                "major_1": "Digital Design",
                "graduation_year_1": "2019",

                # Skills
                "skill_1": "UI/UX Design",
                "skill_2": "User Research",
                "skill_3": "Wireframing",
                "skill_4": "Prototyping",

                # Job 1 (unified)
                "position_1": "Senior UI/UX Designer",
                "company_1": "Tech Startup Inc",
                "start_date_1": "Jan 2020",
                "end_date_1": "Present",
                "job_1_experience_1": "Led design initiatives for 20+ products from concept to launch.",
                "job_1_experience_2": "Increased user engagement by 40% through intuitive interface design.",
                "job_1_experience_3": "Conducted research and testing with 100+ participants.",
                "job_1_experience_4": "Mentored team of 3 junior designers.",

                # Job 2 (unified)
                "position_2": "UI/UX Designer",
                "company_2": "Digital Agency",
                "start_date_2": "June 2017",
                "end_date_2": "Dec 2019",
                "job_2_experience_1": "Created wireframes and high-fidelity mockups for clients.",
                "job_2_experience_2": "Collaborated with developers to implement designs.",
                "job_2_experience_3": "Maintained design system and component library.",
            })

        return base


def main():
    """Test the simplified inserter"""

    inserter = SimplifiedInserter()

    print("="*60)
    print("SIMPLIFIED TEMPLATE VARIABLE INSERTION")
    print("="*60)
    print("\nKey Feature: Unified job_X_experience_Y format")
    print("No distinction between responsibilities and achievements\n")

    for template_name in inserter.templates.keys():
        print(f"Testing: {template_name}")

        try:
            data = inserter.generate_sample_data(template_name)
            output = inserter.populate_template(
                template_name,
                data,
                f"{template_name}_simplified_sample.docx"
            )
            print(f"  ✓ Success: {output}")

        except Exception as e:
            print(f"  ✗ Error: {e}")

    print("\n" + "="*60)
    print("✅ Complete!")
    print(f"Output: {inserter.output_dir}")


if __name__ == "__main__":
    main()