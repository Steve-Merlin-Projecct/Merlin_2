#!/usr/bin/env python3
"""
Complete Template Conversion Script
Replaces ALL text content with semantic variables for maximum flexibility
"""

from docx import Document
import logging
import os
from typing import Dict, List, Tuple

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class ComprehensiveTemplateConverter:
    """Converts templates by replacing ALL text with semantic variables"""

    def __init__(self):
        self.variable_counter = {}
        self.variable_mappings = {}

    def convert_restaurant_manager_template(self):
        """Complete conversion of Restaurant Manager template"""

        input_path = "./content_template_library/manual_converted/restaurant_manager_template.docx"
        output_path = "./content_template_library/manual_converted/restaurant_manager_fully_converted.docx"

        doc = Document(input_path)
        variables_created = []

        # Define comprehensive replacements for ALL content
        replacements = {
            # Header/Contact
            "4567 Main Street, Buffalo, New York 98052 | (716) 555-0100 | m.riley@live.com | www.linkedin.com/in/m.riley":
                "<<street_address>>, <<city>>, <<state>> <<zip_code>> | <<phone>> | <<email>> | <<linkedin_url>>",

            # Name (in table)
            "May Riley": "<<first_name>> <<last_name>>",

            # Section headers (keep or replace based on preference)
            "Profile": "<<profile_header>>",
            "Experience": "<<experience_header>>",
            "Education": "<<education_header>>",
            "Skills & Abilities": "<<skills_header>>",
            "Activities and Interests": "<<interests_header>>",

            # Profile content - break into semantic chunks
            "Friendly and engaging team player and leader able to inspire staff to perform their best. Detail oriented and experienced restaurant manager passionate about food and beverages. A multi-tasker who excels at staff training and recruiting with a track record of inspiring great customer service and customer satisfaction. Regularly exceed sales goals. A master in the art of upselling.":
                "<<career_overview_1>> <<career_overview_2>> <<career_overview_3>> <<career_overview_4>> <<career_overview_5>>",

            # Job 1 Header
            "Restaurant Manager | Contoso Bar and Grill | September 20XX – Present":
                "<<position_1>> | <<company_1>> | <<employment_dates_1>>",

            # Job 1 Bullets
            "Recruit, hire, train, and coach over 30 staff members on customer service skills, food & beverage knowledge, sales, and health & safety standards.":
                "<<job_1_responsibility_1>>",

            "Reduced costs by 7% through controls on overtime, operational efficiencies, and reduced waste.":
                "<<job_1_achievement_1>>",

            "Consistently exceed monthly sales goals by a minimum of 10% by training FOH staff on upselling techniques and by creating a featured food and beverage program.":
                "<<job_1_achievement_2>>",

            # Job 2 Header
            "Restaurant Manager | Fourth Coffee Bistro | June 20XX – August 20XX":
                "<<position_2>> | <<company_2>> | <<employment_dates_2>>",

            # Job 2 Bullets
            "Created a cross-training program ensuring FOH staff members were able to perform confidently and effectively in all positions.":
                "<<job_2_responsibility_1>>",

            "Grew customer based and increased restaurant social media accounts by 19% through interactive promotions, engaging postings, and contests.":
                "<<job_2_achievement_1>>",

            "Created and implemented staff health and safety standards compliance training program, achieving a score of 99% from the Board of Health.":
                "<<job_2_achievement_2>>",

            "Successfully redesigned existing inventory system, ordering and food storage practices, resulting in a 6% decrease in food waste and higher net profits.":
                "<<job_2_achievement_3>>",

            # Education
            "B.S. in Business Administration | June 20XX | Bigtown College, Chicago, Illinois":
                "<<degree_1>> | <<graduation_date_1>> | <<institution_1>>, <<institution_1_location>>",

            "A.A. in Hospitality Management | June 20XX | Bigtown College, Chicago, Illinois":
                "<<degree_2>> | <<graduation_date_2>> | <<institution_2>>, <<institution_2_location>>",

            # Skills (in table)
            "Accounting & Budgeting\nProficient with POS systems\nExcellent interpersonal and communication skills":
                "<<skill_1>>\n<<skill_2>>\n<<skill_3>>",

            "Poised under pressure\nExperienced in most restaurant positions\nFun and energetic":
                "<<skill_4>>\n<<skill_5>>\n<<skill_6>>",

            # Interests
            "Theater, environmental conservation, art, hiking, skiing, travel":
                "<<interest_1>>, <<interest_2>>, <<interest_3>>, <<interest_4>>, <<interest_5>>, <<interest_6>>"
        }

        # Process paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if original_text.strip():
                for old_text, new_text in replacements.items():
                    if old_text in original_text:
                        # Replace at run level to preserve formatting
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                variables_created.extend(self._extract_variables(new_text))
                                logger.info(f"Replaced paragraph text: '{old_text[:50]}...'")
                                break

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    if original_text.strip():
                        for old_text, new_text in replacements.items():
                            if old_text in original_text:
                                # Replace in cell paragraphs
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        if old_text in run.text:
                                            run.text = run.text.replace(old_text, new_text)
                                            variables_created.extend(self._extract_variables(new_text))
                                            logger.info(f"Replaced table text: '{old_text[:50]}...'")

        # Save converted template
        doc.save(output_path)

        # Save variable mappings
        self._save_variable_documentation("restaurant_manager", variables_created, replacements)

        logger.info(f"Restaurant Manager template fully converted: {output_path}")
        logger.info(f"Total unique variables created: {len(set(variables_created))}")

        return output_path, set(variables_created)

    def convert_accountant_template(self):
        """Complete conversion of Accountant template"""

        input_path = "./content_template_library/manual_converted/accountant_template.docx"
        output_path = "./content_template_library/manual_converted/accountant_fully_converted.docx"

        doc = Document(input_path)
        variables_created = []

        replacements = {
            # Name
            "Danielle Brasseur": "<<first_name>> <<last_name>>",

            # Contact line
            "4567 8th Avenue, Carson City, NV 10111 | (313) 555-0100 | danielle@example.com | LinkedIn URL":
                "<<street_address>>, <<city>>, <<state>> <<zip_code>> | <<phone>> | <<email>> | <<linkedin_url>>",

            # Professional summary
            "Dynamic and detail-oriented accountant with expertise in GAAP and comprehensive understanding of financial principles. Proven ability to manage multiple priorities, meet tight deadlines, and deliver accurate financial reporting. Strong analytical skills combined with excellent communication abilities to collaborate effectively with cross-functional teams.":
                "<<professional_summary_1>> <<professional_summary_2>> <<professional_summary_3>>",

            # Education
            "B.A. Accounting and minor in Psychology. GPA 3.4. Graduated 2013. School of Business, University Name.":
                "<<degree>> and minor in <<minor>>. GPA <<gpa>>. Graduated <<graduation_year>>. <<school>>, <<university>>.",

            # Experience header
            "Junior Accountant, Company Name. January 2013 – Present":
                "<<position_1>>, <<company_1>>. <<employment_dates_1>>",

            # Experience bullets
            "Prepared monthly, quarterly, and annual financial statements in accordance with GAAP":
                "<<job_1_responsibility_1>>",
            "Managed accounts payable and receivable, ensuring timely processing and collection":
                "<<job_1_responsibility_2>>",
            "Assisted with budget preparation and variance analysis":
                "<<job_1_responsibility_3>>",
            "Reconciled bank statements and maintained general ledger":
                "<<job_1_responsibility_4>>",
            "Supported annual audit process and tax preparation":
                "<<job_1_responsibility_5>>",

            # Skills
            "Technical: QuickBooks, Excel, SAP":
                "Technical: <<technical_skill_1>>, <<technical_skill_2>>, <<technical_skill_3>>",
            "Certifications: CPA candidate":
                "Certifications: <<certification_1>>",
            "Languages: English (native), Spanish (conversational)":
                "Languages: <<language_1>> (<<language_1_level>>), <<language_2>> (<<language_2_level>>)"
        }

        # Apply replacements to paragraphs and tables
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if original_text.strip():
                for old_text, new_text in replacements.items():
                    if old_text in original_text:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                variables_created.extend(self._extract_variables(new_text))
                                logger.info(f"Replaced: '{old_text[:50]}...'")
                                break

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        if original_text.strip():
                            for old_text, new_text in replacements.items():
                                if old_text in original_text:
                                    for run in paragraph.runs:
                                        if old_text in run.text:
                                            run.text = run.text.replace(old_text, new_text)
                                            variables_created.extend(self._extract_variables(new_text))

        doc.save(output_path)
        self._save_variable_documentation("accountant", variables_created, replacements)

        logger.info(f"Accountant template fully converted: {output_path}")
        logger.info(f"Total unique variables created: {len(set(variables_created))}")

        return output_path, set(variables_created)

    def convert_uiux_designer_template(self):
        """Complete conversion of UI/UX Designer template"""

        input_path = "./content_template_library/manual_converted/uiux_designer_template.docx"
        output_path = "./content_template_library/manual_converted/uiux_designer_fully_converted.docx"

        doc = Document(input_path)
        variables_created = []

        replacements = {
            # Name
            "Angelica": "<<first_name>>",
            "Astrom": "<<last_name>>",

            # Job title and bio
            "UI/UX Designer": "<<current_title>>",

            "I am passionate about designing digital experiences that are both beautiful and functional. With a focus on user research and usability testing, I create intuitive interfaces that solve real problems.":
                "<<professional_bio_1>> <<professional_bio_2>>",

            # Contact info
            "Contact": "<<contact_header>>",
            "angelica@example.com": "<<email>>",
            "www.interestingsite.com": "<<portfolio_website>>",
            "(212) 555-0155": "<<phone>>",
            "New York City": "<<city>>",

            # Education
            "Education": "<<education_header>>",
            "SCHOOL OF FINE ART": "<<institution_1>>",
            "BFA, Graphic Design": "<<degree_1>>, <<major_1>>",
            "20XX": "<<graduation_year_1>>",

            # Skills
            "Skills": "<<skills_header>>",
            "UI/UX design": "<<skill_1>>",
            "User research": "<<skill_2>>",
            "Usability testing": "<<skill_3>>",
            "Project management": "<<skill_4>>",

            # Experience
            "Experience": "<<experience_header>>",

            "Senior UI/UX Designer PROSEWARE, INC.":
                "<<position_1>> <<company_1>>",

            "Jan 20XX - Dec 20XX":
                "<<employment_dates_1>>",

            "Managed the creative process from concept to completion for 20+ products.":
                "<<job_1_achievement_1>>",

            "Designed intuitive user interfaces that increased user engagement by 40%.":
                "<<job_1_achievement_2>>",

            "Conducted user research and usability testing with 100+ participants.":
                "<<job_1_achievement_3>>",

            "Led a team of 3 junior designers and provided mentorship.":
                "<<job_1_achievement_4>>",

            "UI/UX DESIGNER PROSEWARE, INC.":
                "<<position_2>> <<company_2>>",

            "June 20XX - Dec 20XX":
                "<<employment_dates_2>>",

            "Created wireframes, prototypes, and high-fidelity mockups.":
                "<<job_2_responsibility_1>>",

            "Collaborated with developers to implement designs.":
                "<<job_2_responsibility_2>>",

            "Maintained design system and component library.":
                "<<job_2_responsibility_3>>"
        }

        # Apply replacements
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if original_text.strip():
                for old_text, new_text in replacements.items():
                    if old_text in original_text:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                variables_created.extend(self._extract_variables(new_text))
                                logger.info(f"Replaced: '{old_text[:50]}...'")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        if original_text.strip():
                            for old_text, new_text in replacements.items():
                                if old_text in original_text:
                                    for run in paragraph.runs:
                                        if old_text in run.text:
                                            run.text = run.text.replace(old_text, new_text)
                                            variables_created.extend(self._extract_variables(new_text))

        doc.save(output_path)
        self._save_variable_documentation("uiux_designer", variables_created, replacements)

        logger.info(f"UI/UX Designer template fully converted: {output_path}")
        logger.info(f"Total unique variables created: {len(set(variables_created))}")

        return output_path, set(variables_created)

    def _extract_variables(self, text: str) -> List[str]:
        """Extract variable names from text containing <<variable>>"""
        import re
        variables = re.findall(r'<<([^>]+)>>', text)
        return variables

    def _save_variable_documentation(self, template_name: str, variables: List[str], replacements: Dict):
        """Save documentation of variables and their expected content"""

        output_file = f"./documentation/{template_name}_variables.md"

        with open(output_file, 'w') as f:
            f.write(f"# {template_name.replace('_', ' ').title()} Template Variables\n\n")
            f.write("## Complete Variable List\n\n")

            unique_vars = sorted(set(variables))

            # Categorize variables
            categories = {
                "Personal Info": [],
                "Professional Summary": [],
                "Experience": [],
                "Education": [],
                "Skills": [],
                "Other": []
            }

            for var in unique_vars:
                if any(term in var for term in ["name", "email", "phone", "address", "city", "state", "zip", "linkedin"]):
                    categories["Personal Info"].append(var)
                elif any(term in var for term in ["summary", "overview", "bio", "profile"]):
                    categories["Professional Summary"].append(var)
                elif any(term in var for term in ["job", "position", "company", "employment", "responsibility", "achievement"]):
                    categories["Experience"].append(var)
                elif any(term in var for term in ["degree", "institution", "graduation", "school", "university", "gpa", "major", "minor"]):
                    categories["Education"].append(var)
                elif any(term in var for term in ["skill", "technical", "language", "certification"]):
                    categories["Skills"].append(var)
                else:
                    categories["Other"].append(var)

            for category, vars in categories.items():
                if vars:
                    f.write(f"\n### {category}\n")
                    for var in sorted(vars):
                        f.write(f"- `<<{var}>>`\n")

            f.write("\n## Variable Descriptions and Expected Content\n\n")

            # Map variables to their original content for documentation
            f.write("| Variable | Expected Content Type | Example |\n")
            f.write("|----------|----------------------|----------|\n")

            variable_descriptions = self._generate_variable_descriptions(template_name)
            for var in sorted(unique_vars):
                desc = variable_descriptions.get(var, "Custom content")
                f.write(f"| `<<{var}>>` | {desc} | See template |\n")

            f.write(f"\n\nTotal variables: {len(unique_vars)}\n")

    def _generate_variable_descriptions(self, template_name: str) -> Dict[str, str]:
        """Generate descriptions for each variable"""

        common = {
            "first_name": "User's first name",
            "last_name": "User's last name",
            "email": "Email address",
            "phone": "Phone number",
            "street_address": "Street address",
            "city": "City name",
            "state": "State/Province",
            "zip_code": "ZIP/Postal code",
            "linkedin_url": "LinkedIn profile URL",

            "profile_header": "Section header (default: 'Profile')",
            "experience_header": "Section header (default: 'Experience')",
            "education_header": "Section header (default: 'Education')",
            "skills_header": "Section header (default: 'Skills')",
            "interests_header": "Section header (default: 'Interests')",
            "contact_header": "Section header (default: 'Contact')",
        }

        restaurant = {
            "career_overview_1": "Team leadership qualities",
            "career_overview_2": "Professional expertise",
            "career_overview_3": "Core competencies",
            "career_overview_4": "Track record/achievements",
            "career_overview_5": "Special expertise",

            "position_1": "First job title",
            "company_1": "First company name",
            "employment_dates_1": "First job date range",
            "job_1_responsibility_1": "Key responsibility statement",
            "job_1_achievement_1": "Quantified achievement",
            "job_1_achievement_2": "Sales/performance achievement",

            "position_2": "Second job title",
            "company_2": "Second company name",
            "employment_dates_2": "Second job date range",
            "job_2_responsibility_1": "Initiative or program created",
            "job_2_achievement_1": "Growth metrics achieved",
            "job_2_achievement_2": "Compliance/quality achievement",
            "job_2_achievement_3": "Efficiency improvement",

            "degree_1": "Primary degree (e.g., B.S. in Business)",
            "graduation_date_1": "Graduation date",
            "institution_1": "College/University name",
            "institution_1_location": "School location",

            "skill_1": "Financial/technical skill",
            "skill_2": "System proficiency",
            "skill_3": "Soft skill",
            "skill_4": "Performance trait",
            "skill_5": "Experience breadth",
            "skill_6": "Personality trait",

            "interest_1": "Personal interest/hobby",
            "interest_2": "Personal interest/hobby",
            "interest_3": "Personal interest/hobby",
            "interest_4": "Personal interest/hobby",
            "interest_5": "Personal interest/hobby",
            "interest_6": "Personal interest/hobby",
        }

        if template_name == "restaurant_manager":
            return {**common, **restaurant}

        # Add other template-specific descriptions as needed
        return common

    def convert_all_templates(self):
        """Convert all three templates"""
        results = {}

        logger.info("Starting comprehensive template conversion...")

        # Convert each template
        results['restaurant_manager'] = self.convert_restaurant_manager_template()
        results['accountant'] = self.convert_accountant_template()
        results['uiux_designer'] = self.convert_uiux_designer_template()

        # Summary
        logger.info("\n" + "="*60)
        logger.info("CONVERSION COMPLETE")
        for template, (path, variables) in results.items():
            logger.info(f"{template}: {len(variables)} variables - {path}")

        return results


if __name__ == "__main__":
    converter = ComprehensiveTemplateConverter()
    results = converter.convert_all_templates()

    print("\n" + "="*60)
    print("All templates have been fully converted!")
    print("Check ./content_template_library/manual_converted/ for output files")
    print("Check ./documentation/ for variable documentation")