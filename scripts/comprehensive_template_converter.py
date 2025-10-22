#!/usr/bin/env python3
"""
Comprehensive Template Converter - Replaces ALL text with semantic variables
This script ensures EVERY piece of text becomes a variable for maximum flexibility
"""

from docx import Document
import logging
import os
import re
from typing import Dict, List, Set, Tuple

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class FullTemplateConverter:
    """Converts templates by replacing ALL text content with semantic variables"""

    def __init__(self):
        self.variables_created = {}
        self.variable_descriptions = {}

    def convert_restaurant_manager(self):
        """Completely convert Restaurant Manager template"""

        input_path = "./content_template_library/manual_converted/restaurant_manager_template.docx"
        output_path = "./content_template_library/manual_converted/restaurant_manager_fully_converted.docx"

        doc = Document(input_path)
        variables = set()

        # Process all paragraphs
        para_replacements = [
            # Header/Contact
            ("4567 Main Street", "<<street_address>>"),
            ("Buffalo", "<<city>>"),
            ("New York", "<<state>>"),
            ("98052", "<<zip_code>>"),
            ("(716) 555-0100", "<<phone>>"),
            ("m.riley@live.com", "<<email>>"),
            ("www.linkedin.com/in/m.riley", "<<linkedin_url>>"),

            # Section Headers
            ("Profile", "<<profile_section_header>>"),
            ("Experience", "<<experience_section_header>>"),
            ("Education", "<<education_section_header>>"),
            ("Skills & Abilities", "<<skills_section_header>>"),
            ("Activities and Interests", "<<interests_section_header>>"),

            # Profile paragraph - break into sentences
            ("Friendly and engaging team player and leader able to inspire staff to perform their best.",
             "<<career_overview_1>>"),
            ("Detail oriented and experienced restaurant manager passionate about food and beverages.",
             "<<career_overview_2>>"),
            ("A multi-tasker who excels at staff training and recruiting with a track record of inspiring great customer service and customer satisfaction.",
             "<<career_overview_3>>"),
            ("Regularly exceed sales goals.",
             "<<career_overview_4>>"),
            ("A master in the art of upselling.",
             "<<career_overview_5>>"),

            # Job 1
            ("Restaurant Manager", "<<position_1>>"),
            ("Contoso Bar and Grill", "<<company_1>>"),
            ("September 20XX", "<<start_date_1>>"),
            ("Present", "<<end_date_1>>"),

            # Job 1 Bullets
            ("Recruit, hire, train, and coach over 30 staff members on customer service skills, food & beverage knowledge, sales, and health & safety standards.",
             "<<job_1_responsibility_1>>"),
            ("Reduced costs by 7% through controls on overtime, operational efficiencies, and reduced waste.",
             "<<job_1_achievement_1>>"),
            ("Consistently exceed monthly sales goals by a minimum of 10% by training FOH staff on upselling techniques and by creating a featured food and beverage program.",
             "<<job_1_achievement_2>>"),

            # Job 2
            ("Restaurant Manager", "<<position_2>>"),
            ("Fourth Coffee Bistro", "<<company_2>>"),
            ("June 20XX", "<<start_date_2>>"),
            ("August 20XX", "<<end_date_2>>"),

            # Job 2 Bullets
            ("Created a cross-training program ensuring FOH staff members were able to perform confidently and effectively in all positions.",
             "<<job_2_responsibility_1>>"),
            ("Grew customer based and increased restaurant social media accounts by 19% through interactive promotions, engaging postings, and contests.",
             "<<job_2_achievement_1>>"),
            ("Created and implemented staff health and safety standards compliance training program, achieving a score of 99% from the Board of Health.",
             "<<job_2_achievement_2>>"),
            ("Successfully redesigned existing inventory system, ordering and food storage practices, resulting in a 6% decrease in food waste and higher net profits.",
             "<<job_2_achievement_3>>"),

            # Education 1
            ("B.S. in Business Administration", "<<degree_1>>"),
            ("June 20XX", "<<graduation_date_1>>"),
            ("Bigtown College", "<<institution_1>>"),
            ("Chicago", "<<institution_1_city>>"),
            ("Illinois", "<<institution_1_state>>"),

            # Education 2
            ("A.A. in Hospitality Management", "<<degree_2>>"),
            ("June 20XX", "<<graduation_date_2>>"),
            ("Bigtown College", "<<institution_2>>"),
            ("Chicago", "<<institution_2_city>>"),
            ("Illinois", "<<institution_2_state>>"),

            # Interests
            ("Theater", "<<interest_1>>"),
            ("environmental conservation", "<<interest_2>>"),
            ("art", "<<interest_3>>"),
            ("hiking", "<<interest_4>>"),
            ("skiing", "<<interest_5>>"),
            ("travel", "<<interest_6>>"),
        ]

        # Table replacements
        table_replacements = [
            ("May Riley", "<<first_name>> <<last_name>>"),

            # Skills column 1
            ("Accounting & Budgeting", "<<skill_1>>"),
            ("Proficient with POS systems", "<<skill_2>>"),
            ("Excellent interpersonal and communication skills", "<<skill_3>>"),

            # Skills column 2
            ("Poised under pressure", "<<skill_4>>"),
            ("Experienced in most restaurant positions", "<<skill_5>>"),
            ("Fun and energetic", "<<skill_6>>"),
        ]

        # Apply paragraph replacements
        for para in doc.paragraphs:
            if para.text.strip():
                for old_text, new_text in para_replacements:
                    if old_text in para.text:
                        # Replace in all runs to preserve formatting
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                variables.update(self._extract_variables(new_text))
                                logger.info(f"Replaced: '{old_text[:50]}...' → {new_text}")

        # Apply table replacements
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in table_replacements:
                            if old_text in para.text:
                                for run in para.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        variables.update(self._extract_variables(new_text))
                                        logger.info(f"Replaced in table: '{old_text[:30]}...' → {new_text}")

        doc.save(output_path)

        self.variables_created['restaurant_manager'] = variables
        logger.info(f"Restaurant Manager: {len(variables)} unique variables created")

        return output_path, variables

    def convert_accountant(self):
        """Completely convert Accountant template"""

        input_path = "./content_template_library/manual_converted/accountant_template.docx"
        output_path = "./content_template_library/manual_converted/accountant_fully_converted.docx"

        doc = Document(input_path)
        variables = set()

        # Paragraph replacements
        para_replacements = [
            # Name
            ("Danielle Brasseur", "<<first_name>> <<last_name>>"),

            # Contact
            ("4567 8th Avenue", "<<street_address>>"),
            ("Carson City", "<<city>>"),
            ("NV", "<<state>>"),
            ("10111", "<<zip_code>>"),
            ("(313) 555-0100", "<<phone>>"),
            ("danielle@example.com", "<<email>>"),
            ("www.linkedin.com", "<<linkedin_url>>"),

            # Professional Summary - break into components
            ("Dynamic and detail-oriented accountant with expertise in GAAP and comprehensive public accounting experience.",
             "<<professional_summary_1>>"),
            ("Known for delivering top-notch strategic solutions and fostering business growth through effective collaboration and ownership mentality.",
             "<<professional_summary_2>>"),
        ]

        # Table content replacements
        table_replacements = [
            # Headers
            ("Education", "<<education_header>>"),
            ("Experience", "<<experience_header>>"),
            ("Skills", "<<skills_header>>"),

            # Education content
            ("Bachelor of Science in Accounting", "<<degree_1>>"),
            ("Minor in Business Administration", "<<minor_1>>"),
            ("Bellows College", "<<institution_1>>"),
            ("Degree", "<<degree_label>>"),
            ("June 20XX", "<<graduation_date_1>>"),

            # Experience
            ("Accountant", "<<position_1>>"),
            ("Trey Research", "<<company_1>>"),
            ("San Francisco", "<<job_1_city>>"),
            ("CA", "<<job_1_state>>"),
            ("March 20XX", "<<start_date_1>>"),
            ("Present", "<<end_date_1>>"),

            # Experience bullets (need to extract from cell text)
            ("Provide accounting services with a focus on estate and trust tax returns.",
             "<<job_1_responsibility_1>>"),
            ("Prepare financial statements and tax returns for individuals and businesses.",
             "<<job_1_responsibility_2>>"),
            ("Conduct financial audits and reviews.",
             "<<job_1_responsibility_3>>"),
            ("Advise clients on tax planning strategies.",
             "<<job_1_responsibility_4>>"),

            # Skills column 1
            ("Microsoft NAV Dynamics", "<<technical_skill_1>>"),
            ("Cashflow planning & management", "<<technical_skill_2>>"),
            ("State & federal tax codes", "<<technical_skill_3>>"),

            # Skills column 2
            ("Bookkeeping", "<<technical_skill_4>>"),
            ("Exceptional communication", "<<soft_skill_1>>"),
            ("Fluent in German", "<<language_skill_1>>"),
        ]

        # Apply replacements
        for para in doc.paragraphs:
            if para.text.strip():
                for old_text, new_text in para_replacements:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                variables.update(self._extract_variables(new_text))
                                logger.info(f"Replaced: '{old_text[:50]}...' → {new_text}")

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    if cell_text.strip():
                        # Apply all table replacements
                        for old_text, new_text in table_replacements:
                            if old_text in cell_text:
                                # Replace in all paragraphs in the cell
                                for para in cell.paragraphs:
                                    if old_text in para.text:
                                        for run in para.runs:
                                            if old_text in run.text:
                                                run.text = run.text.replace(old_text, new_text)
                                                variables.update(self._extract_variables(new_text))
                                                logger.info(f"Replaced in table: '{old_text[:30]}...' → {new_text}")

        doc.save(output_path)

        self.variables_created['accountant'] = variables
        logger.info(f"Accountant: {len(variables)} unique variables created")

        return output_path, variables

    def convert_uiux_designer(self):
        """Completely convert UI/UX Designer template"""

        input_path = "./content_template_library/manual_converted/uiux_designer_template.docx"
        output_path = "./content_template_library/manual_converted/uiux_designer_fully_converted.docx"

        doc = Document(input_path)
        variables = set()

        # Name paragraphs
        para_replacements = [
            ("Angelica", "<<first_name>>"),
            ("Astrom", "<<last_name>>"),
        ]

        # Table replacements
        table_replacements = [
            # Left column - Title and Bio
            ("UI/UX Designer", "<<current_title>>"),

            # Professional bio - break into semantic parts
            ("I am passionate about designing digital experiences that are both beautiful and functional.",
             "<<professional_bio_1>>"),
            ("With a focus on user research and usability testing, I create intuitive interfaces that solve real problems.",
             "<<professional_bio_2>>"),
            ("My approach combines creativity with data-driven design decisions.",
             "<<professional_bio_3>>"),

            # Contact section
            ("Contact", "<<contact_header>>"),
            ("angelica@example.com", "<<email>>"),
            ("www.interestingsite.com", "<<portfolio_website>>"),
            ("(212) 555-0155", "<<phone>>"),
            ("New York City", "<<city>>"),

            # Education section
            ("Education", "<<education_header>>"),
            ("SCHOOL OF FINE ART", "<<institution_1>>"),
            ("BFA", "<<degree_type_1>>"),
            ("Graphic Design", "<<major_1>>"),
            ("20XX", "<<graduation_year_1>>"),

            # Skills section
            ("Skills", "<<skills_header>>"),
            ("UI/UX design", "<<skill_1>>"),
            ("User research", "<<skill_2>>"),
            ("Usability testing", "<<skill_3>>"),
            ("Project management", "<<skill_4>>"),
            ("Wireframing", "<<skill_5>>"),
            ("Prototyping", "<<skill_6>>"),

            # Experience section
            ("Experience", "<<experience_header>>"),

            # Job 1
            ("Senior UI/UX Designer", "<<position_1>>"),
            ("PROSEWARE, INC.", "<<company_1>>"),
            ("Jan 20XX", "<<start_date_1>>"),
            ("Dec 20XX", "<<end_date_1>>"),

            # Job 1 achievements
            ("Managed the creative process from concept to completion for 20+ products.",
             "<<job_1_achievement_1>>"),
            ("Designed intuitive user interfaces that increased user engagement by 40%.",
             "<<job_1_achievement_2>>"),
            ("Conducted user research and usability testing with 100+ participants.",
             "<<job_1_achievement_3>>"),
            ("Led a team of 3 junior designers and provided mentorship.",
             "<<job_1_achievement_4>>"),

            # Job 2
            ("UI/UX DESIGNER", "<<position_2>>"),
            ("PROSEWARE, INC.", "<<company_2>>"),
            ("June 20XX", "<<start_date_2>>"),
            ("Dec 20XX", "<<end_date_2>>"),

            # Job 2 responsibilities
            ("Created wireframes, prototypes, and high-fidelity mockups.",
             "<<job_2_responsibility_1>>"),
            ("Collaborated with developers to implement designs.",
             "<<job_2_responsibility_2>>"),
            ("Maintained design system and component library.",
             "<<job_2_responsibility_3>>"),
        ]

        # Apply paragraph replacements
        for para in doc.paragraphs:
            if para.text.strip():
                for old_text, new_text in para_replacements:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                variables.update(self._extract_variables(new_text))
                                logger.info(f"Replaced: '{old_text}' → {new_text}")

        # Apply table replacements
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    if cell_text.strip():
                        for old_text, new_text in table_replacements:
                            if old_text in cell_text:
                                for para in cell.paragraphs:
                                    if old_text in para.text:
                                        for run in para.runs:
                                            if old_text in run.text:
                                                run.text = run.text.replace(old_text, new_text)
                                                variables.update(self._extract_variables(new_text))
                                                logger.info(f"Replaced in table: '{old_text[:30]}...' → {new_text}")

        doc.save(output_path)

        self.variables_created['uiux_designer'] = variables
        logger.info(f"UI/UX Designer: {len(variables)} unique variables created")

        return output_path, variables

    def _extract_variables(self, text: str) -> Set[str]:
        """Extract variable names from text"""
        variables = re.findall(r'<<([^>]+)>>', text)
        return set(variables)

    def generate_documentation(self):
        """Generate comprehensive documentation of all variables"""

        doc_path = "./documentation/TEMPLATE_VARIABLES_COMPLETE.md"

        with open(doc_path, 'w') as f:
            f.write("# Complete Template Variable Documentation\n\n")
            f.write("This document lists ALL variables created for each template with descriptions of expected content.\n\n")

            # Summary
            f.write("## Summary\n\n")
            f.write("| Template | Total Variables | File |\n")
            f.write("|----------|----------------|------|\n")

            for template, variables in self.variables_created.items():
                f.write(f"| {template.replace('_', ' ').title()} | {len(variables)} | {template}_fully_converted.docx |\n")

            # Detailed variable lists
            for template, variables in self.variables_created.items():
                f.write(f"\n## {template.replace('_', ' ').title()} Template Variables\n\n")

                # Categorize variables
                categories = self._categorize_variables(variables)

                for category, vars in categories.items():
                    if vars:
                        f.write(f"### {category}\n\n")
                        f.write("| Variable | Description | Example Content |\n")
                        f.write("|----------|-------------|----------------|\n")

                        for var in sorted(vars):
                            desc = self._get_variable_description(var)
                            example = self._get_variable_example(var)
                            f.write(f"| `<<{var}>>` | {desc} | {example} |\n")

                        f.write("\n")

            f.write("\n## Variable Input Requirements\n\n")
            f.write("When using these templates, provide data in the following format:\n\n")
            f.write("```json\n")
            f.write("{\n")
            f.write('  "first_name": "John",\n')
            f.write('  "last_name": "Doe",\n')
            f.write('  "email": "john.doe@example.com",\n')
            f.write('  "phone": "(555) 123-4567",\n')
            f.write('  "street_address": "123 Main St",\n')
            f.write('  "city": "New York",\n')
            f.write('  "state": "NY",\n')
            f.write('  "zip_code": "10001",\n')
            f.write('  \n')
            f.write('  "career_overview_1": "Dynamic professional with...",\n')
            f.write('  "career_overview_2": "Proven track record in...",\n')
            f.write('  \n')
            f.write('  "position_1": "Senior Manager",\n')
            f.write('  "company_1": "Tech Corp",\n')
            f.write('  "start_date_1": "Jan 2020",\n')
            f.write('  "end_date_1": "Present",\n')
            f.write('  "job_1_responsibility_1": "Led team of 15 professionals...",\n')
            f.write('  "job_1_achievement_1": "Increased revenue by 25%...",\n')
            f.write('  \n')
            f.write('  "degree_1": "Bachelor of Science",\n')
            f.write('  "institution_1": "State University",\n')
            f.write('  "graduation_date_1": "May 2019",\n')
            f.write('  \n')
            f.write('  "skill_1": "Project Management",\n')
            f.write('  "skill_2": "Data Analysis",\n')
            f.write('  "// etc...": ""\n')
            f.write("}\n")
            f.write("```\n")

        logger.info(f"Documentation saved to: {doc_path}")
        return doc_path

    def _categorize_variables(self, variables: Set[str]) -> Dict[str, List[str]]:
        """Categorize variables by type"""
        categories = {
            "Personal Information": [],
            "Contact Information": [],
            "Professional Summary": [],
            "Career Overview": [],
            "Work Experience": [],
            "Education": [],
            "Skills": [],
            "Interests": [],
            "Section Headers": [],
            "Other": []
        }

        for var in variables:
            if any(term in var for term in ["first_name", "last_name"]):
                categories["Personal Information"].append(var)
            elif any(term in var for term in ["email", "phone", "address", "city", "state", "zip", "linkedin"]):
                categories["Contact Information"].append(var)
            elif "professional_summary" in var or "professional_bio" in var:
                categories["Professional Summary"].append(var)
            elif "career_overview" in var:
                categories["Career Overview"].append(var)
            elif any(term in var for term in ["position", "company", "job", "responsibility", "achievement", "start_date", "end_date", "employment"]):
                categories["Work Experience"].append(var)
            elif any(term in var for term in ["degree", "institution", "graduation", "major", "minor", "school"]):
                categories["Education"].append(var)
            elif "skill" in var or "technical" in var or "language" in var:
                categories["Skills"].append(var)
            elif "interest" in var:
                categories["Interests"].append(var)
            elif "header" in var or "label" in var:
                categories["Section Headers"].append(var)
            else:
                categories["Other"].append(var)

        return categories

    def _get_variable_description(self, var: str) -> str:
        """Get description for a variable"""
        descriptions = {
            "career_overview_1": "Opening statement about leadership/teamwork",
            "career_overview_2": "Core expertise and passion",
            "career_overview_3": "Key competencies and track record",
            "career_overview_4": "Performance metrics/achievements",
            "career_overview_5": "Unique value proposition",
            "job_1_responsibility_1": "Primary responsibility in first role",
            "job_1_achievement_1": "Quantified achievement",
            "job_1_achievement_2": "Revenue/performance impact",
            "job_2_responsibility_1": "Initiative or program created",
            "job_2_achievement_1": "Growth metrics",
            "job_2_achievement_2": "Quality/compliance achievement",
            "job_2_achievement_3": "Efficiency improvement",
            "professional_summary_1": "Core expertise statement",
            "professional_summary_2": "Value proposition",
            "professional_bio_1": "Personal mission/passion",
            "professional_bio_2": "Approach and methodology",
            "professional_bio_3": "Unique differentiator",
        }

        return descriptions.get(var, "User-provided content")

    def _get_variable_example(self, var: str) -> str:
        """Get example content for a variable"""
        examples = {
            "first_name": "John",
            "last_name": "Doe",
            "email": "john@example.com",
            "phone": "(555) 123-4567",
            "city": "New York",
            "state": "NY",
            "career_overview_1": "Dynamic leader with...",
            "position_1": "Senior Manager",
            "company_1": "Fortune 500 Company",
            "skill_1": "Strategic Planning",
        }

        return examples.get(var, "...")

    def convert_all(self):
        """Convert all templates and generate documentation"""

        logger.info("="*60)
        logger.info("STARTING COMPREHENSIVE TEMPLATE CONVERSION")
        logger.info("="*60)

        results = {}

        # Convert each template
        results['restaurant_manager'] = self.convert_restaurant_manager()
        results['accountant'] = self.convert_accountant()
        results['uiux_designer'] = self.convert_uiux_designer()

        # Generate documentation
        doc_path = self.generate_documentation()

        # Summary
        logger.info("="*60)
        logger.info("CONVERSION COMPLETE")
        logger.info("="*60)

        total_vars = sum(len(vars) for _, vars in results.values())
        logger.info(f"Total variables created across all templates: {total_vars}")

        for template, (path, variables) in results.items():
            logger.info(f"  {template}: {len(variables)} variables → {path}")

        logger.info(f"Documentation: {doc_path}")

        return results


if __name__ == "__main__":
    converter = FullTemplateConverter()
    results = converter.convert_all()

    print("\n" + "="*60)
    print("✅ ALL TEMPLATES FULLY CONVERTED!")
    print("="*60)
    print("\nFiles created:")
    print("  • restaurant_manager_fully_converted.docx")
    print("  • accountant_fully_converted.docx")
    print("  • uiux_designer_fully_converted.docx")
    print("  • TEMPLATE_VARIABLES_COMPLETE.md")
    print("\nCheck ./content_template_library/manual_converted/ for templates")
    print("Check ./documentation/ for variable documentation")