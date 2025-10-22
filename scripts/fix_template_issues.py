#!/usr/bin/env python3
"""
Fix issues identified in FINDINGS_REPORT.md:
1. Remove state/province variable
2. Add missing variables (company_1, start_date_1, degree_1, etc.)
3. Fix position_1 duplication (Job 2 should use position_2)
4. Replace hardcoded values with variables
"""

from docx import Document
import re


def replace_text_in_runs(paragraph, old_text, new_text):
    """Replace text while preserving formatting by working at run level"""
    full_text = ''.join(run.text for run in paragraph.runs)

    if old_text in full_text:
        # Find which runs contain the text
        char_index = 0
        for run in paragraph.runs:
            run_length = len(run.text)

            # Check if replacement starts in this run
            if char_index <= full_text.find(old_text) < char_index + run_length:
                # Build new paragraph text with replacement
                new_full_text = full_text.replace(old_text, new_text, 1)

                # Clear all runs and put new text in first run
                for run in paragraph.runs:
                    run.text = ''
                paragraph.runs[0].text = new_full_text
                return True

            char_index += run_length

    return False


def fix_template(input_path, output_path):
    """Fix all identified issues in template"""

    print("=" * 60)
    print("FIXING TEMPLATE ISSUES")
    print("=" * 60)
    print()

    doc = Document(input_path)

    # Track replacements
    replacements_made = []

    # Define all fixes needed
    fixes = [
        # Remove state variable (replace pattern with city, no state)
        (", <<state>> <<zip_code>>", " <<zip_code>>", "Remove state/province variable"),

        # Fix Job 1 - add missing variables
        ("Contoso Bar and Grill", "<<company_1>>", "Add company_1 variable"),
        ("September 20XX", "<<start_date_1>>", "Add start_date_1 variable"),

        # Fix Job 2 - add missing variables and fix position duplication
        # This is tricky - we need to replace the SECOND occurrence of position_1
        ("June 20XX", "<<start_date_2>>", "Add start_date_2 variable"),

        # Fix Education - add missing variables
        ("B.S. in Business Administration", "<<degree_1>>", "Add degree_1 variable"),
        ("A.A. in Hospitality Management", "<<degree_2>>", "Add degree_2 variable"),
    ]

    # Apply fixes to all paragraphs
    for old_text, new_text, description in fixes:
        for para in doc.paragraphs:
            if replace_text_in_runs(para, old_text, new_text):
                replacements_made.append(description)
                print(f"✅ {description}")

    # Apply fixes to all table cells
    for old_text, new_text, description in fixes:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if replace_text_in_runs(para, old_text, new_text):
                            if description not in replacements_made:
                                replacements_made.append(description)
                                print(f"✅ {description}")

    # Special handling for position_1 duplication in Job 2
    # Need to find and replace the SECOND occurrence
    print()
    print("Fixing position_1 duplication in Job 2...")

    position_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if '<<position_1>>' in text and '<<company_2>>' in text:
                        # This is Job 2 - replace position_1 with position_2
                        replace_text_in_runs(para, '<<position_1>>', '<<position_2>>')
                        print(f"✅ Fixed position_1 → position_2 in Job 2")
                        replacements_made.append("Fix position_2 in Job 2")
                        break

    # Handle graduation dates - need to distinguish between two entries
    print()
    print("Fixing graduation dates...")

    graduation_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if 'June 20XX' in para.text:
                        graduation_count += 1
                        if graduation_count == 1:
                            replace_text_in_runs(para, 'June 20XX', '<<graduation_date_1>>')
                            print(f"✅ Added graduation_date_1 variable")
                        elif graduation_count == 2:
                            replace_text_in_runs(para, 'June 20XX', '<<graduation_date_2>>')
                            print(f"✅ Added graduation_date_2 variable")

    # Save fixed template
    doc.save(output_path)

    print()
    print("=" * 60)
    print(f"TEMPLATE FIXED: {output_path}")
    print(f"Total fixes applied: {len(set(replacements_made))}")
    print("=" * 60)

    return replacements_made


if __name__ == "__main__":
    fixes = fix_template(
        './content_template_library/manual_converted/restaurant_manager_final.docx',
        './content_template_library/manual_converted/restaurant_manager_fixed.docx'
    )

    print()
    print("Summary of fixes:")
    for fix in set(fixes):
        print(f"  - {fix}")
