#!/usr/bin/env python3
"""
Simplified Template Converter - Uses unified variable naming
All work experience uses <<job_X_experience_Y>> format
"""

from docx import Document
import logging
import os
import re
from typing import Dict, List, Set

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class SimplifiedTemplateConverter:
    """Converts templates with simplified, unified variable naming"""

    def __init__(self):
        self.variables_created = {}

    def convert_restaurant_manager(self):
        """Convert Restaurant Manager with simplified variables"""

        input_path = "./content_template_library/manual_converted/restaurant_manager_template.docx"
        output_path = "./content_template_library/manual_converted/restaurant_manager_final.docx"

        doc = Document(input_path)
        variables = set()

        # Simplified replacements - job experience uses sequential numbering
        para_replacements = [
            # Header/Contact
            ("4567 Main Street", "<<street_address>>"),
            ("Buffalo", "<<city>>"),
            ("New York", "<<state>>"),
            ("98052", "<<zip_code>>"),
            ("(716) 555-0100", "<<phone>>"),
            ("m.riley@live.com", "<<email>>"),
            ("www.linkedin.com/in/m.riley", "<<linkedin_url>>"),

            # Section Headers
            ("Profile", "<<profile_header>>"),
            ("Experience", "<<experience_header>>"),
            ("Education", "<<education_header>>"),
            ("Skills & Abilities", "<<skills_header>>"),
            ("Activities and Interests", "<<interests_header>>"),

            # Career Overview (5 sentences)
            ("Friendly and engaging team player and leader able to inspire staff to perform their best.",
             "<<career_overview_1>>"),
            ("Detail oriented and experienced restaurant manager passionate about food and beverages.",
             "<<career_overview_2>>"),
            ("A multi-tasker who excels at staff training and recruiting with a track record of inspiring great customer service and customer satisfaction.",
             "<<career_overview_3>>"),
            ("Regularly exceed sales goals.",
             "<<career_overview_4>>"),
            ("A master in the art of upselling.",
             "<<career_overview_5>>"),

            # Job 1 Header
            ("Restaurant Manager", "<<position_1>>"),
            ("Contoso Bar and Grill", "<<company_1>>"),
            ("September 20XX", "<<start_date_1>>"),
            ("Present", "<<end_date_1>>"),

            # Job 1 Experience (all unified as experience_X)
            ("Recruit, hire, train, and coach over 30 staff members on customer service skills, food & beverage knowledge, sales, and health & safety standards.",
             "<<job_1_experience_1>>"),
            ("Reduced costs by 7% through controls on overtime, operational efficiencies, and reduced waste.",
             "<<job_1_experience_2>>"),
            ("Consistently exceed monthly sales goals by a minimum of 10% by training FOH staff on upselling techniques and by creating a featured food and beverage program.",
             "<<job_1_experience_3>>"),

            # Job 2 Header
            ("Restaurant Manager", "<<position_2>>"),
            ("Fourth Coffee Bistro", "<<company_2>>"),
            ("June 20XX", "<<start_date_2>>"),
            ("August 20XX", "<<end_date_2>>"),

            # Job 2 Experience (all unified)
            ("Created a cross-training program ensuring FOH staff members were able to perform confidently and effectively in all positions.",
             "<<job_2_experience_1>>"),
            ("Grew customer based and increased restaurant social media accounts by 19% through interactive promotions, engaging postings, and contests.",
             "<<job_2_experience_2>>"),
            ("Created and implemented staff health and safety standards compliance training program, achieving a score of 99% from the Board of Health.",
             "<<job_2_experience_3>>"),
            ("Successfully redesigned existing inventory system, ordering and food storage practices, resulting in a 6% decrease in food waste and higher net profits.",
             "<<job_2_experience_4>>"),

            # Education
            ("B.S. in Business Administration", "<<degree_1>>"),
            ("June 20XX", "<<graduation_date_1>>"),
            ("Bigtown College", "<<institution_1>>"),
            ("Chicago", "<<institution_city_1>>"),
            ("Illinois", "<<institution_state_1>>"),

            ("A.A. in Hospitality Management", "<<degree_2>>"),
            ("June 20XX", "<<graduation_date_2>>"),
            ("Bigtown College", "<<institution_2>>"),
            ("Chicago", "<<institution_city_2>>"),
            ("Illinois", "<<institution_state_2>>"),

            # Interests
            ("Theater", "<<interest_1>>"),
            ("environmental conservation", "<<interest_2>>"),
            ("art", "<<interest_3>>"),
            ("hiking", "<<interest_4>>"),
            ("skiing", "<<interest_5>>"),
            ("travel", "<<interest_6>>"),
        ]

        table_replacements = [
            ("May Riley", "<<first_name>> <<last_name>>"),

            # Skills
            ("Accounting & Budgeting", "<<skill_1>>"),
            ("Proficient with POS systems", "<<skill_2>>"),
            ("Excellent interpersonal and communication skills", "<<skill_3>>"),
            ("Poised under pressure", "<<skill_4>>"),
            ("Experienced in most restaurant positions", "<<skill_5>>"),
            ("Fun and energetic", "<<skill_6>>"),
        ]

        # Apply replacements
        for para in doc.paragraphs:
            if para.text.strip():
                for old_text, new_text in para_replacements:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                variables.update(self._extract_variables(new_text))
                                logger.info(f"Replaced: '{old_text[:40]}...'")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in table_replacements:
                            if old_text in para.text:
                                for run in para.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        variables.update(self._extract_variables(new_text))
                                        logger.info(f"Replaced in table: '{old_text[:30]}...'")

        doc.save(output_path)
        self.variables_created['restaurant_manager'] = variables
        logger.info(f"✓ Restaurant Manager: {len(variables)} variables → {output_path}")
        return output_path, variables

    def convert_accountant(self):
        """Convert Accountant with simplified variables"""

        input_path = "./content_template_library/manual_converted/accountant_template.docx"
        output_path = "./content_template_library/manual_converted/accountant_final.docx"

        doc = Document(input_path)
        variables = set()

        para_replacements = [
            ("Danielle Brasseur", "<<first_name>> <<last_name>>"),
            ("4567 8th Avenue", "<<street_address>>"),
            ("Carson City", "<<city>>"),
            ("NV", "<<state>>"),
            ("10111", "<<zip_code>>"),
            ("(313) 555-0100", "<<phone>>"),
            ("danielle@example.com", "<<email>>"),
            ("www.linkedin.com", "<<linkedin_url>>"),

            # Professional Summary (2 parts)
            ("Dynamic and detail-oriented accountant with expertise in GAAP and comprehensive public accounting experience.",
             "<<professional_summary_1>>"),
            ("Known for delivering top-notch strategic solutions and fostering business growth through effective collaboration and ownership mentality.",
             "<<professional_summary_2>>"),
        ]

        table_replacements = [
            ("Education", "<<education_header>>"),
            ("Experience", "<<experience_header>>"),
            ("Skills", "<<skills_header>>"),

            # Education
            ("Bachelor of Science in Accounting", "<<degree_1>>"),
            ("Minor in Business Administration", "<<minor_1>>"),
            ("Bellows College", "<<institution_1>>"),
            ("Degree", "<<degree_label>>"),
            ("June 20XX", "<<graduation_date_1>>"),

            # Experience
            ("Accountant", "<<position_1>>"),
            ("Trey Research", "<<company_1>>"),
            ("San Francisco", "<<job_city_1>>"),
            ("CA", "<<job_state_1>>"),
            ("March 20XX", "<<start_date_1>>"),
            ("Present", "<<end_date_1>>"),

            # Skills
            ("Microsoft NAV Dynamics", "<<skill_1>>"),
            ("Cashflow planning & management", "<<skill_2>>"),
            ("State & federal tax codes", "<<skill_3>>"),
            ("Bookkeeping", "<<skill_4>>"),
            ("Exceptional communication", "<<skill_5>>"),
            ("Fluent in German", "<<skill_6>>"),
        ]

        for para in doc.paragraphs:
            if para.text.strip():
                for old_text, new_text in para_replacements:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                variables.update(self._extract_variables(new_text))
                                logger.info(f"Replaced: '{old_text[:40]}...'")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in table_replacements:
                            if old_text in para.text:
                                for run in para.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        variables.update(self._extract_variables(new_text))

        doc.save(output_path)
        self.variables_created['accountant'] = variables
        logger.info(f"✓ Accountant: {len(variables)} variables → {output_path}")
        return output_path, variables

    def convert_uiux_designer(self):
        """Convert UI/UX Designer with simplified variables"""

        input_path = "./content_template_library/manual_converted/uiux_designer_template.docx"
        output_path = "./content_template_library/manual_converted/uiux_designer_final.docx"

        doc = Document(input_path)
        variables = set()

        para_replacements = [
            ("Angelica", "<<first_name>>"),
            ("Astrom", "<<last_name>>"),
        ]

        table_replacements = [
            ("UI/UX Designer", "<<current_title>>"),

            # Professional bio
            ("I am passionate about designing digital experiences that are both beautiful and functional.",
             "<<professional_bio_1>>"),
            ("With a focus on user research and usability testing, I create intuitive interfaces that solve real problems.",
             "<<professional_bio_2>>"),

            # Contact
            ("Contact", "<<contact_header>>"),
            ("angelica@example.com", "<<email>>"),
            ("www.interestingsite.com", "<<portfolio_website>>"),
            ("(212) 555-0155", "<<phone>>"),
            ("New York City", "<<city>>"),

            # Education
            ("Education", "<<education_header>>"),
            ("SCHOOL OF FINE ART", "<<institution_1>>"),
            ("BFA", "<<degree_1>>"),
            ("Graphic Design", "<<major_1>>"),
            ("20XX", "<<graduation_year_1>>"),

            # Skills
            ("Skills", "<<skills_header>>"),
            ("UI/UX design", "<<skill_1>>"),
            ("User research", "<<skill_2>>"),
            ("Usability testing", "<<skill_3>>"),
            ("Project management", "<<skill_4>>"),

            # Experience
            ("Experience", "<<experience_header>>"),
            ("Senior UI/UX Designer", "<<position_1>>"),
            ("PROSEWARE, INC.", "<<company_1>>"),
            ("Jan 20XX", "<<start_date_1>>"),
            ("Dec 20XX", "<<end_date_1>>"),

            # Job 1 Experience (unified)
            ("Managed the creative process from concept to completion for 20+ products.",
             "<<job_1_experience_1>>"),
            ("Designed intuitive user interfaces that increased user engagement by 40%.",
             "<<job_1_experience_2>>"),
            ("Conducted user research and usability testing with 100+ participants.",
             "<<job_1_experience_3>>"),
            ("Led a team of 3 junior designers and provided mentorship.",
             "<<job_1_experience_4>>"),

            # Job 2
            ("UI/UX DESIGNER", "<<position_2>>"),
            ("PROSEWARE, INC.", "<<company_2>>"),
            ("June 20XX", "<<start_date_2>>"),
            ("Dec 20XX", "<<end_date_2>>"),

            # Job 2 Experience (unified)
            ("Created wireframes, prototypes, and high-fidelity mockups.",
             "<<job_2_experience_1>>"),
            ("Collaborated with developers to implement designs.",
             "<<job_2_experience_2>>"),
            ("Maintained design system and component library.",
             "<<job_2_experience_3>>"),
        ]

        for para in doc.paragraphs:
            if para.text.strip():
                for old_text, new_text in para_replacements:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                variables.update(self._extract_variables(new_text))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in table_replacements:
                            if old_text in para.text:
                                for run in para.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        variables.update(self._extract_variables(new_text))

        doc.save(output_path)
        self.variables_created['uiux_designer'] = variables
        logger.info(f"✓ UI/UX Designer: {len(variables)} variables → {output_path}")
        return output_path, variables

    def _extract_variables(self, text: str) -> Set[str]:
        """Extract variable names from text"""
        return set(re.findall(r'<<([^>]+)>>', text))

    def generate_documentation(self):
        """Generate simplified documentation"""

        doc_path = "./documentation/SIMPLIFIED_VARIABLES.md"

        with open(doc_path, 'w') as f:
            f.write("# Simplified Template Variables\n\n")
            f.write("All templates use unified, simplified variable naming.\n\n")

            f.write("## Key Simplifications\n\n")
            f.write("- **Work Experience**: All bullets use `<<job_X_experience_Y>>` format\n")
            f.write("- **No distinction** between responsibilities and achievements\n")
            f.write("- **Sequential numbering**: Experience bullets numbered 1, 2, 3, etc.\n\n")

            f.write("## Summary\n\n")
            f.write("| Template | Variables | File |\n")
            f.write("|----------|-----------|------|\n")

            for template, variables in self.variables_created.items():
                f.write(f"| {template.replace('_', ' ').title()} | {len(variables)} | {template}_final.docx |\n")

            for template, variables in self.variables_created.items():
                f.write(f"\n## {template.replace('_', ' ').title()}\n\n")

                # Categorize
                categories = {
                    "Contact": [],
                    "Career Overview / Summary": [],
                    "Work Experience": [],
                    "Education": [],
                    "Skills": [],
                    "Interests": [],
                    "Headers": []
                }

                for var in sorted(variables):
                    if any(x in var for x in ['name', 'email', 'phone', 'address', 'city', 'state', 'zip', 'linkedin']):
                        categories["Contact"].append(var)
                    elif 'overview' in var or 'summary' in var or 'bio' in var:
                        categories["Career Overview / Summary"].append(var)
                    elif 'job' in var or 'position' in var or 'company' in var or 'start_date' in var or 'end_date' in var:
                        categories["Work Experience"].append(var)
                    elif 'degree' in var or 'institution' in var or 'graduation' in var or 'major' in var or 'minor' in var:
                        categories["Education"].append(var)
                    elif 'skill' in var:
                        categories["Skills"].append(var)
                    elif 'interest' in var:
                        categories["Interests"].append(var)
                    elif 'header' in var:
                        categories["Headers"].append(var)

                for category, vars_list in categories.items():
                    if vars_list:
                        f.write(f"### {category}\n")
                        for var in vars_list:
                            f.write(f"- `<<{var}>>`\n")
                        f.write("\n")

            # Usage example
            f.write("\n## Example: Work Experience Variables\n\n")
            f.write("```json\n")
            f.write("{\n")
            f.write('  "job_1_experience_1": "Led team of 40 staff members across operations",\n')
            f.write('  "job_1_experience_2": "Reduced costs by 15% through strategic planning",\n')
            f.write('  "job_1_experience_3": "Increased customer satisfaction from 4.2 to 4.8 stars",\n')
            f.write('  \n')
            f.write('  "job_2_experience_1": "Implemented training program for new hires",\n')
            f.write('  "job_2_experience_2": "Grew social media following by 45%",\n')
            f.write('  "job_2_experience_3": "Achieved highest inspection score in district"\n')
            f.write("}\n")
            f.write("```\n\n")
            f.write("**Note**: No distinction between responsibilities and achievements - all come from the same source.\n")

        logger.info(f"Documentation: {doc_path}")
        return doc_path

    def convert_all(self):
        """Convert all templates"""

        logger.info("="*60)
        logger.info("SIMPLIFIED TEMPLATE CONVERSION")
        logger.info("="*60)

        results = {}
        results['restaurant_manager'] = self.convert_restaurant_manager()
        results['accountant'] = self.convert_accountant()
        results['uiux_designer'] = self.convert_uiux_designer()

        doc_path = self.generate_documentation()

        logger.info("="*60)
        logger.info("CONVERSION COMPLETE")
        logger.info("="*60)

        total = sum(len(vars) for _, vars in results.values())
        logger.info(f"Total variables: {total}")

        return results


if __name__ == "__main__":
    converter = SimplifiedTemplateConverter()
    results = converter.convert_all()

    print("\n" + "="*60)
    print("✅ SIMPLIFIED CONVERSION COMPLETE")
    print("="*60)
    print("\nKey Changes:")
    print("  • job_X_experience_Y (unified format)")
    print("  • No responsibility vs achievement distinction")
    print("  • All experience bullets from same source")
    print("\nFiles:")
    print("  • restaurant_manager_final.docx")
    print("  • accountant_final.docx")
    print("  • uiux_designer_final.docx")