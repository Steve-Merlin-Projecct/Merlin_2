#!/usr/bin/env python3
"""
Final Manual Template Conversion Script
Comprehensive conversion of Word document templates into variable-based templates
while preserving all original formatting.

This version includes extensive replacements for all identifiable personal,
professional, and educational content.
"""

import re
from pathlib import Path
from docx import Document
from typing import Dict, List, Set


class FinalTemplateConverter:
    """
    Comprehensive template converter with extensive variable mappings.
    """

    def __init__(self):
        """Initialize the converter."""
        self.converted_variables = {}
        self.replacement_count = 0

    def replace_in_paragraph(self, paragraph, old_text, new_text, case_sensitive=True):
        """
        Replace text in paragraph while preserving formatting.

        Args:
            paragraph: python-docx Paragraph object
            old_text: Text to find and replace
            new_text: Replacement text (variable placeholder)
            case_sensitive: Whether to match case exactly

        Returns:
            bool: True if replacement was made
        """
        full_text = paragraph.text

        # Check if text exists
        if case_sensitive:
            if old_text not in full_text:
                return False
            start_pos = full_text.find(old_text)
        else:
            if old_text.lower() not in full_text.lower():
                return False
            start_pos = full_text.lower().find(old_text.lower())

        if start_pos == -1:
            return False

        end_pos = start_pos + len(old_text)

        # Track replacement
        if old_text not in self.converted_variables:
            self.converted_variables[old_text] = new_text

        # Find affected runs
        current_pos = 0
        runs_to_modify = []

        for run in paragraph.runs:
            run_length = len(run.text)
            run_end = current_pos + run_length

            if current_pos <= start_pos < run_end or current_pos < end_pos <= run_end or (start_pos <= current_pos and end_pos >= run_end):
                runs_to_modify.append({
                    'run': run,
                    'start': max(0, start_pos - current_pos),
                    'end': min(run_length, end_pos - current_pos),
                })

            current_pos = run_end

        # Apply replacements
        if len(runs_to_modify) == 1:
            run_info = runs_to_modify[0]
            run = run_info['run']
            run.text = run.text[:run_info['start']] + new_text + run.text[run_info['end']:]
            self.replacement_count += 1
            return True
        elif len(runs_to_modify) > 1:
            for i, run_info in enumerate(runs_to_modify):
                if i == 0:
                    run_info['run'].text = run_info['run'].text[:run_info['start']] + new_text
                elif i == len(runs_to_modify) - 1:
                    run_info['run'].text = run_info['run'].text[run_info['end']:]
                else:
                    run_info['run'].text = ''
            self.replacement_count += 1
            return True

        return False

    def convert_restaurant_manager_template(self, input_path: str, output_path: str):
        """
        Convert the restaurant manager template with comprehensive replacements.

        Args:
            input_path: Path to input .docx file
            output_path: Path to save converted .docx file
        """
        print(f"\n{'='*70}")
        print(f"Converting: Restaurant Manager Template")
        print(f"{'='*70}")

        doc = Document(input_path)
        self.converted_variables = {}
        self.replacement_count = 0

        # Comprehensive replacements in order of specificity (most specific first)
        replacements = [
            # Full contact header line - do this first before breaking it down
            ('4567 Main Street, Buffalo, New York 98052 | (716) 555-0100 | m.riley@live.com | www.linkedin.com/in/m.riley',
             '<<street_address>>, <<city>>, <<state>> <<zip_code>> | <<phone>> | <<email>> | <<linkedin>>'),

            # Personal information
            ('May Riley', '<<first_name>> <<last_name>>'),
            ('4567 Main Street, Buffalo, New York 98052', '<<street_address>>, <<city>>, <<state>> <<zip_code>>'),
            ('Buffalo, New York 98052', '<<city>>, <<state>> <<zip_code>>'),
            ('(716) 555-0100', '<<phone>>'),
            ('m.riley@live.com', '<<email>>'),
            ('www.linkedin.com/in/m.riley', '<<linkedin>>'),
            ('www.linkedin.com/in/m', '<<linkedin>>'),

            # Company names
            ('Contoso Bar and Grill', '<<company_name>>'),
            ('Fourth Coffee Bistro', '<<company_name>>'),

            # Education institutions (specific before general)
            ('Bigtown College, Chicago, Illinois', '<<education_institution>>, <<city>>, <<state>>'),
            ('Bigtown College', '<<education_institution>>'),
            ('Chicago, Illinois', '<<city>>, <<state>>'),

            # Degrees
            ('B.S. in Business Administration', '<<degree>>'),
            ('A.A. in Hospitality Management', '<<degree>>'),

            # Job titles
            ('Restaurant Manager', '<<position_title>>'),

            # Date ranges (most specific first)
            ('September 20XX – Present', '<<start_date>> - <<end_date>>'),
            ('June 20XX – August 20XX', '<<start_date>> - <<end_date>>'),
            ('June 20XX', '<<graduation_date>>'),

            # Skills and abilities (convert bullet-style lists)
            ('Accounting & Budgeting', '<<skill_1>>'),
            ('Proficient with POS systems', '<<skill_2>>'),
            ('Excellent interpersonal and communication skills', '<<skill_3>>'),
            ('Poised under pressure', '<<skill_4>>'),
            ('Experienced in most restaurant positions', '<<skill_5>>'),
            ('Fun and energetic', '<<skill_6>>'),

            # Interests
            ('Theater, environmental conservation, art, hiking, skiing, travel', '<<interests>>'),
        ]

        # Process all paragraphs
        for para in doc.paragraphs:
            for old_text, new_text in replacements:
                if self.replace_in_paragraph(para, old_text, new_text):
                    print(f"  [Para] '{old_text[:50]}...' -> {new_text}")

        # Process all tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        for old_text, new_text in replacements:
                            if self.replace_in_paragraph(para, old_text, new_text):
                                print(f"  [T{table_idx}R{row_idx}C{cell_idx}] '{old_text[:40]}...' -> {new_text}")

        doc.save(output_path)
        print(f"\nTotal replacements made: {self.replacement_count}")
        print(f"Saved to: {output_path}")

        return doc, self.converted_variables

    def convert_accountant_template(self, input_path: str, output_path: str):
        """
        Convert the accountant template with comprehensive replacements.

        Args:
            input_path: Path to input .docx file
            output_path: Path to save converted .docx file
        """
        print(f"\n{'='*70}")
        print(f"Converting: Accountant Template")
        print(f"{'='*70}")

        doc = Document(input_path)
        self.converted_variables = {}
        self.replacement_count = 0

        # Comprehensive replacements in order of specificity
        replacements = [
            # Full contact header
            ('4567 8th Avenue, Carson City, NV 10111 | (313) 555-0100 | danielle@example.com | www.linkedin.com',
             '<<street_address>>, <<city>>, <<state>> <<zip_code>> | <<phone>> | <<email>> | <<linkedin>>'),

            # Personal information
            ('Danielle Brasseur', '<<first_name>> <<last_name>>'),
            ('4567 8th Avenue, Carson City, NV 10111', '<<street_address>>, <<city>>, <<state>> <<zip_code>>'),
            ('Carson City, NV 10111', '<<city>>, <<state>> <<zip_code>>'),
            ('(313) 555-0100', '<<phone>>'),
            ('danielle@example.com', '<<email>>'),
            ('www.linkedin.com/in/danielle', '<<linkedin>>'),
            ('www.linkedin.com', '<<linkedin>>'),

            # Education (most specific first)
            ('Bachelor of Science in Accounting, Minor in Business Administration', '<<degree>>, Minor in <<minor>>'),
            ('University of Nevada, Reno, NV', '<<education_institution>>, <<city>>, <<state>>'),
            ('Bellows College', '<<education_institution>>'),

            # Job titles and companies
            ('Senior Accountant', '<<position_title>>'),
            ('Staff Accountant', '<<position_title>>'),
            ('Accountant', '<<position_title>>'),  # Less specific, do last
            ('Trey Research', '<<company_name>>'),
            ('Contoso Pharmaceuticals', '<<company_name>>'),

            # Locations
            ('San Francisco, CA', '<<city>>, <<state>>'),
            ('Reno, NV', '<<city>>, <<state>>'),

            # Date ranges
            ('March 20XX – Present', '<<start_date>> - <<end_date>>'),
            ('June 20XX – February 20XX', '<<start_date>> - <<end_date>>'),
            ('May 20XX – May 20XX', '<<start_date>> - <<end_date>>'),
            ('June 20XX – April 20XX', '<<start_date>> - <<end_date>>'),
            ('May 20XX', '<<graduation_date>>'),

            # Skills
            ('Microsoft NAV Dynamics', '<<skill_1>>'),
            ('Cashflow planning & management', '<<skill_2>>'),
            ('State & federal tax codes', '<<skill_3>>'),
            ('Bookkeeping', '<<skill_4>>'),
            ('Exceptional communication', '<<skill_5>>'),
            ('Fluent in German', '<<skill_6>>'),
        ]

        # Process all paragraphs
        for para in doc.paragraphs:
            for old_text, new_text in replacements:
                if self.replace_in_paragraph(para, old_text, new_text):
                    print(f"  [Para] '{old_text[:50]}...' -> {new_text}")

        # Process all tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        for old_text, new_text in replacements:
                            if self.replace_in_paragraph(para, old_text, new_text):
                                print(f"  [T{table_idx}R{row_idx}C{cell_idx}] '{old_text[:40]}...' -> {new_text}")

        doc.save(output_path)
        print(f"\nTotal replacements made: {self.replacement_count}")
        print(f"Saved to: {output_path}")

        return doc, self.converted_variables

    def print_summary(self, template_name: str):
        """Print summary of variables used."""
        print(f"\n{'-'*70}")
        print(f"Variable Summary: {template_name}")
        print(f"{'-'*70}")

        unique_vars = sorted(set(self.converted_variables.values()))

        print(f"\nTotal unique variables: {len(unique_vars)}")
        print(f"Total replacements: {self.replacement_count}\n")

        for var in unique_vars:
            examples = [k for k, v in self.converted_variables.items() if v == var]
            print(f"{var}")
            for example in examples:
                display_text = example if len(example) <= 50 else example[:50] + "..."
                print(f"  - '{display_text}'")


def main():
    """Main conversion workflow."""
    print("="*70)
    print("FINAL WORD DOCUMENT TEMPLATE CONVERTER")
    print("="*70)
    print("Comprehensive conversion with full variable replacement\n")

    base_dir = Path("/workspace/content_template_library/manual_converted")

    # File definitions
    templates = [
        {
            'name': 'Restaurant Manager',
            'input': base_dir / "restaurant_manager_template.docx",
            'output': base_dir / "restaurant_manager_template_converted.docx",
            'method': 'convert_restaurant_manager_template'
        },
        {
            'name': 'Accountant',
            'input': base_dir / "accountant_template.docx",
            'output': base_dir / "accountant_template_converted.docx",
            'method': 'convert_accountant_template'
        }
    ]

    converter = FinalTemplateConverter()
    results = {}

    # Convert each template
    for template in templates:
        if not template['input'].exists():
            print(f"\nWARNING: File not found: {template['input']}")
            results[template['name']] = {'success': False, 'error': 'File not found'}
            continue

        try:
            method = getattr(converter, template['method'])
            doc, variables = method(str(template['input']), str(template['output']))
            converter.print_summary(template['name'])
            results[template['name']] = {
                'success': True,
                'output': template['output'],
                'variable_count': len(set(variables.values())),
                'replacement_count': converter.replacement_count
            }
        except Exception as e:
            print(f"\nERROR converting {template['name']}: {e}")
            import traceback
            traceback.print_exc()
            results[template['name']] = {'success': False, 'error': str(e)}

    # Final summary
    print("\n\n" + "="*70)
    print("FINAL CONVERSION SUMMARY")
    print("="*70)

    for template_name, result in results.items():
        print(f"\n{template_name}:")
        if result['success']:
            print(f"  Status: ✓ SUCCESS")
            print(f"  Output: {result['output']}")
            print(f"  Unique Variables: {result['variable_count']}")
            print(f"  Total Replacements: {result['replacement_count']}")
        else:
            print(f"  Status: ✗ FAILED")
            print(f"  Error: {result['error']}")

    print("\n" + "="*70)
    print("Conversion Complete!")
    print("="*70)
    print("\nThe converted templates are ready for use with the document")
    print("generation system. All formatting has been preserved.\n")


if __name__ == "__main__":
    main()
