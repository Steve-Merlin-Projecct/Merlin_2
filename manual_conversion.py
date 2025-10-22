#!/usr/bin/env python3
"""
Manual Template Conversion Script
Converts raw Word document templates into variable-based templates
while preserving all original formatting.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from typing import Dict, List, Tuple


class TemplateConverter:
    """Converts Word document templates by replacing content with variables."""

    def __init__(self):
        """Initialize the converter with variable mapping patterns."""
        # Comprehensive mapping of content patterns to variable placeholders
        self.variable_mappings = {
            # Personal Information
            r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b': '<<first_name>> <<last_name>>',  # Full names
            r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Boulevard|Blvd)\b': '<<street_address>>',
            r'\b[A-Z][a-z]+,\s*[A-Z]{2}\s+\d{5}(?:-\d{4})?\b': '<<city>>, <<state>> <<zip_code>>',
            r'\b\(\d{3}\)\s*\d{3}-\d{4}\b': '<<phone>>',  # (555) 555-5555
            r'\b\d{3}-\d{3}-\d{4}\b': '<<phone>>',  # 555-555-5555
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b': '<<email>>',
            r'linkedin\.com/in/[\w-]+': '<<linkedin>>',

            # Professional Content - Job Titles
            r'\bRestaurant Manager\b': '<<position_title>>',
            r'\bAccountant\b': '<<position_title>>',
            r'\bSenior Accountant\b': '<<position_title>>',
            r'\bStaff Accountant\b': '<<position_title>>',
            r'\bAssistant Manager\b': '<<position_title>>',

            # Date Ranges
            r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*[-–]\s*(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b': '<<start_date>> - <<end_date>>',
            r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*[-–]\s*Present\b': '<<start_date>> - <<end_date>>',
            r'\b\d{1,2}/\d{4}\s*[-–]\s*\d{1,2}/\d{4}\b': '<<start_date>> - <<end_date>>',
            r'\b\d{1,2}/\d{4}\s*[-–]\s*Present\b': '<<start_date>> - <<end_date>>',

            # Education
            r'\bBachelor of Science in Accounting\b': '<<degree>>',
            r'\bBachelor\'?s Degree in\s+[A-Za-z\s]+\b': '<<degree>>',
            r'\bMaster\'?s Degree in\s+[A-Za-z\s]+\b': '<<degree>>',
            r'\bUniversity of\s+[A-Za-z\s]+\b': '<<education_institution>>',
            r'\b[A-Z][a-z]+\s+University\b': '<<education_institution>>',
        }

        self.converted_variables = {}

    def preserve_run_formatting(self, run):
        """
        Extract formatting properties from a run.

        Args:
            run: A python-docx Run object

        Returns:
            Dict containing formatting properties
        """
        formatting = {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name if run.font.name else None,
            'font_size': run.font.size,
            'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
        }
        return formatting

    def apply_run_formatting(self, run, formatting):
        """
        Apply formatting properties to a run.

        Args:
            run: A python-docx Run object
            formatting: Dict containing formatting properties
        """
        if formatting['bold'] is not None:
            run.bold = formatting['bold']
        if formatting['italic'] is not None:
            run.italic = formatting['italic']
        if formatting['underline'] is not None:
            run.underline = formatting['underline']
        if formatting['font_name']:
            run.font.name = formatting['font_name']
        if formatting['font_size']:
            run.font.size = formatting['font_size']
        if formatting['font_color']:
            run.font.color.rgb = formatting['font_color']

    def replace_text_preserve_formatting(self, paragraph, pattern, replacement):
        """
        Replace text in a paragraph while preserving formatting.

        Args:
            paragraph: A python-docx Paragraph object
            pattern: Regex pattern to match
            replacement: Variable placeholder to insert

        Returns:
            bool: True if replacement was made
        """
        # Get full paragraph text
        full_text = paragraph.text

        # Check if pattern matches
        match = re.search(pattern, full_text, re.IGNORECASE)
        if not match:
            return False

        # Track the replacement
        matched_text = match.group(0)
        if matched_text not in self.converted_variables:
            self.converted_variables[matched_text] = replacement

        # Find which run(s) contain the matched text
        start_pos = match.start()
        end_pos = match.end()

        current_pos = 0
        runs_to_modify = []

        for run in paragraph.runs:
            run_length = len(run.text)
            run_end = current_pos + run_length

            # Check if this run overlaps with the match
            if current_pos <= start_pos < run_end or current_pos < end_pos <= run_end or (start_pos <= current_pos and end_pos >= run_end):
                runs_to_modify.append({
                    'run': run,
                    'start': max(0, start_pos - current_pos),
                    'end': min(run_length, end_pos - current_pos),
                    'formatting': self.preserve_run_formatting(run)
                })

            current_pos = run_end

        # If the match spans a single run, simple replacement
        if len(runs_to_modify) == 1:
            run_info = runs_to_modify[0]
            run = run_info['run']
            run.text = run.text[:run_info['start']] + replacement + run.text[run_info['end']:]
            return True

        # If match spans multiple runs, more complex
        elif len(runs_to_modify) > 1:
            # Clear all but the first run
            for i, run_info in enumerate(runs_to_modify):
                if i == 0:
                    # First run: keep text before match, add replacement
                    run_info['run'].text = run_info['run'].text[:run_info['start']] + replacement
                elif i == len(runs_to_modify) - 1:
                    # Last run: keep text after match
                    run_info['run'].text = run_info['run'].text[run_info['end']:]
                else:
                    # Middle runs: clear completely
                    run_info['run'].text = ''
            return True

        return False

    def convert_document(self, input_path: str, output_path: str, template_name: str):
        """
        Convert a Word document template by replacing content with variables.

        Args:
            input_path: Path to input .docx file
            output_path: Path to save converted .docx file
            template_name: Name of template for reporting
        """
        print(f"\n{'='*60}")
        print(f"Converting: {template_name}")
        print(f"{'='*60}")

        # Load the document
        doc = Document(input_path)

        # Reset variables tracking for this document
        self.converted_variables = {}

        # Process each paragraph
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue

            # Try each pattern
            for pattern, replacement in self.variable_mappings.items():
                if self.replace_text_preserve_formatting(paragraph, pattern, replacement):
                    print(f"  Paragraph {para_idx}: Replaced with {replacement}")

        # Process tables (common in resumes)
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue

                        for pattern, replacement in self.variable_mappings.items():
                            if self.replace_text_preserve_formatting(para, pattern, replacement):
                                print(f"  Table {table_idx}, Row {row_idx}, Cell {cell_idx}: Replaced with {replacement}")

        # Save the converted document
        doc.save(output_path)
        print(f"\nSaved converted template to: {output_path}")

        # Report variables used
        print(f"\n{'-'*60}")
        print(f"Variables used in {template_name}:")
        print(f"{'-'*60}")
        unique_vars = sorted(set(self.converted_variables.values()))
        for var in unique_vars:
            examples = [k for k, v in self.converted_variables.items() if v == var]
            print(f"  {var}")
            for example in examples[:3]:  # Show up to 3 examples
                print(f"    - Replaced: '{example}'")

        return doc, self.converted_variables


def inspect_document_structure(doc_path: str):
    """
    Inspect a document to understand its structure before conversion.

    Args:
        doc_path: Path to the .docx file
    """
    print(f"\n{'='*60}")
    print(f"Inspecting: {doc_path}")
    print(f"{'='*60}\n")

    doc = Document(doc_path)

    print("PARAGRAPHS:")
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  [{idx}] {para.text[:100]}...")
            if para.runs:
                for run_idx, run in enumerate(para.runs):
                    if run.text.strip():
                        print(f"      Run {run_idx}: Bold={run.bold}, Italic={run.italic}, "
                              f"Font={run.font.name}, Size={run.font.size}")

    print("\nTABLES:")
    for table_idx, table in enumerate(doc.tables):
        print(f"  Table {table_idx}: {len(table.rows)} rows x {len(table.columns)} columns")
        for row_idx, row in enumerate(table.rows[:3]):  # Show first 3 rows
            print(f"    Row {row_idx}:")
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()[:50]
                if cell_text:
                    print(f"      Cell {cell_idx}: {cell_text}...")


def main():
    """Main conversion workflow."""
    print("="*60)
    print("WORD DOCUMENT TEMPLATE CONVERTER")
    print("="*60)
    print("This script converts raw templates into variable-based templates")
    print("while preserving all original formatting.\n")

    # Define file paths
    base_dir = Path("/workspace/content_template_library/manual_converted")

    templates = [
        {
            'name': 'Restaurant Manager Template',
            'input': base_dir / "restaurant_manager_template.docx",
            'output': base_dir / "restaurant_manager_template_converted.docx"
        },
        {
            'name': 'Accountant Template',
            'input': base_dir / "accountant_template.docx",
            'output': base_dir / "accountant_template_converted.docx"
        }
    ]

    # Initialize converter
    converter = TemplateConverter()

    # Inspect documents first (optional, for debugging)
    print("\nSTEP 1: Inspecting document structures...")
    for template in templates:
        if template['input'].exists():
            inspect_document_structure(str(template['input']))
        else:
            print(f"WARNING: File not found: {template['input']}")

    # Convert documents
    print("\n\nSTEP 2: Converting documents...")
    results = {}

    for template in templates:
        if template['input'].exists():
            try:
                doc, variables = converter.convert_document(
                    str(template['input']),
                    str(template['output']),
                    template['name']
                )
                results[template['name']] = {
                    'success': True,
                    'variables': variables,
                    'output': template['output']
                }
            except Exception as e:
                print(f"ERROR converting {template['name']}: {e}")
                import traceback
                traceback.print_exc()
                results[template['name']] = {
                    'success': False,
                    'error': str(e)
                }
        else:
            results[template['name']] = {
                'success': False,
                'error': 'File not found'
            }

    # Final summary
    print("\n\n" + "="*60)
    print("CONVERSION SUMMARY")
    print("="*60)

    for template_name, result in results.items():
        print(f"\n{template_name}:")
        if result['success']:
            print(f"  Status: SUCCESS")
            print(f"  Output: {result['output']}")
            print(f"  Variables: {len(set(result['variables'].values()))} unique variables")
        else:
            print(f"  Status: FAILED")
            print(f"  Error: {result['error']}")

    print("\n" + "="*60)
    print("Conversion complete!")
    print("="*60)


if __name__ == "__main__":
    main()
