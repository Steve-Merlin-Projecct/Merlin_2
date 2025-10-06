import os
import json
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from replit import object_storage
import logging

# Import database manager for job tracking
try:
    from .database_manager import DatabaseManager
    database_available = True
except ImportError:
    database_available = False
    logging.warning("Database manager not available - job tracking disabled")

class BaseDocumentGenerator:
    """
    Base class for all document generators
    Provides shared functionality for document creation and storage
    """
    
    def __init__(self):
        """Initialize the base generator with object storage and database tracking"""
        try:
            # Initialize Replit Object Storage
            self.storage = object_storage.Client()
            logging.info("Document generator connected to Replit Object Storage (default bucket)")
        except Exception as e:
            logging.error(f"Failed to connect to object storage: {e}")
            self.storage = None
        
        # Initialize database manager for job tracking
        self.db_manager = None
        if database_available:
            try:
                self.db_manager = DatabaseManager()
                logging.info("Database tracking enabled for document generation")
            except Exception as e:
                logging.warning(f"Database tracking disabled due to error: {e}")
                self.db_manager = None
    
    def _create_document_with_metadata(self, title, author, subject=None, keywords=None, category="Professional Document"):
        """Create a new document with professional metadata"""
        doc = Document()
        
        # Set document properties for authenticity
        doc.core_properties.title = title
        doc.core_properties.author = author
        doc.core_properties.subject = subject or title
        doc.core_properties.keywords = keywords or "Professional Document"
        doc.core_properties.category = category
        doc.core_properties.comments = f"Generated in Edmonton, Alberta, Canada using Microsoft Word"
        doc.core_properties.last_modified_by = author
        doc.core_properties.revision = 3
        doc.core_properties.version = "3.2"
        doc.core_properties.language = "en-CA"
        doc.core_properties.identifier = "Microsoft Office Word"
        doc.core_properties.content_status = "Final"
        
        return doc
    
    def _create_job_record(self, document_type, webhook_data, title, author):
        """Create initial job record in database"""
        if not self.db_manager:
            return None
        
        try:
            job_id = self.db_manager.create_job(
                document_type=document_type,
                webhook_data=webhook_data,
                title=title,
                author=author
            )
            logging.info(f"Created job record: {job_id} for {document_type}")
            return job_id
        except Exception as e:
            logging.error(f"Failed to create job record: {e}")
            return None
    
    def _save_document(self, doc, base_filename, title, author, job_id=None, document_type=None):
        """Save document to object storage and return file information"""
        try:
            # Generate unique filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            filename = f"{timestamp}_{base_filename}_{unique_id}.docx"
            
            # Save to bytes
            from io import BytesIO
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            # Upload to object storage
            storage_path = f"documents/{filename}"
            
            if self.storage:
                # Upload to Replit Object Storage
                self.storage.upload_from_bytes(
                    storage_path,
                    doc_bytes.getvalue()
                )
                logging.info(f"Document uploaded to object storage: {storage_path}")
                storage_type = "object_storage"
            else:
                # Fallback to local storage
                os.makedirs("storage", exist_ok=True)
                local_path = f"storage/{filename}"
                doc.save(local_path)
                logging.info(f"Document saved locally: {local_path}")
                storage_type = "local"
                storage_path = local_path
            
            # Return file information
            file_size = len(doc_bytes.getvalue())
            file_info = {
                "filename": filename,
                "file_path": storage_path,
                "file_size": file_size,
                "file_size_mb": round(file_size / (1024 * 1024), 2),
                "title": title,
                "author": author,
                "created_at": datetime.now().isoformat(),
                "storage_type": storage_type,
                "object_storage_path": storage_path if storage_type == "object_storage" else None,
                "download_url": f"/download/{filename}",
                "direct_download_url": f"/download/{filename}"
            }
            
            # Update job record with success at the very end
            if job_id and self.db_manager:
                try:
                    self.db_manager.update_job_success(job_id, file_info)
                    logging.info(f"Updated job {job_id} with success status")
                except Exception as db_error:
                    logging.error(f"Failed to update job success status: {db_error}")
            
            return file_info
            
        except Exception as e:
            logging.error(f"Error saving document: {e}")
            # Update job record with failure at the very end
            if job_id and self.db_manager:
                try:
                    self.db_manager.update_job_failure(
                        job_id, 
                        error_code='SAVE_ERROR',
                        error_message=str(e),
                        error_details={'document_type': document_type, 'filename': base_filename}
                    )
                    logging.info(f"Updated job {job_id} with failure status")
                except Exception as db_error:
                    logging.error(f"Failed to update job failure status: {db_error}")
            
            raise Exception(f"Failed to save document: {str(e)}")
    
    def _handle_generation_error(self, job_id, document_type, error_code, error_message, error_details=None):
        """Handle generation errors and update job record"""
        if job_id and self.db_manager:
            try:
                self.db_manager.update_job_failure(
                    job_id,
                    error_code=error_code,
                    error_message=error_message,
                    error_details=error_details
                )
                logging.info(f"Updated job {job_id} with error: {error_code}")
            except Exception as e:
                logging.error(f"Failed to update job error status: {e}")
    
    def _add_section_heading(self, doc, heading_text, level=1):
        """Add a section heading to the document"""
        heading = doc.add_heading(heading_text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return heading
    
    def _add_paragraph_with_formatting(self, doc, text, bold_prefix=None, font_size=11):
        """Add a paragraph with optional bold prefix"""
        para = doc.add_paragraph()
        
        if bold_prefix:
            bold_run = para.add_run(bold_prefix)
            bold_run.bold = True
            
        text_run = para.add_run(text)
        text_run.font.size = Pt(font_size)
        
        return para
    
    def cleanup_old_files(self, max_age_hours=24):
        """
        Clean up old generated files
        
        Args:
            max_age_hours (int): Maximum age of files to keep in hours
        """
        # Implementation would depend on storage system
        # For now, this is a placeholder for future implementation
        pass