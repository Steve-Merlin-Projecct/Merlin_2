import os
import logging
import uuid
import io
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from replit.object_storage import Client

class DocumentGenerator:
    """
    Handles Word document generation from webhook data
    """
    
    def __init__(self):
        self.storage_dir = os.path.join(os.getcwd(), 'storage')
        # Initialize Replit Object Storage client (using default bucket)
        try:
            self.storage_client = Client()
            logging.info("Connected to Replit Object Storage (default bucket)")
        except Exception as e:
            logging.error(f"Failed to initialize object storage: {str(e)}")
            self.storage_client = None
        
    def generate_document(self, webhook_data):
        """
        Generate a Word document from webhook data
        
        Args:
            webhook_data (dict): Data received from Make.com webhook
            
        Returns:
            dict: File information including path, filename, and size
        """
        try:
            # Create new document
            doc = Document()
            
            # Extract data from webhook payload with Marketing Manager defaults
            title = webhook_data.get('title', 'Marketing Manager Application')
            content = webhook_data.get('content', '')
            author = webhook_data.get('author', 'Steve Glen')
            date = webhook_data.get('date', datetime.now().strftime('%Y-%m-%d'))
            sections = webhook_data.get('sections', [])
            formatting = webhook_data.get('formatting', {})
            
            # Set enhanced document properties for authenticity with realistic timestamps
            doc.core_properties.title = title
            doc.core_properties.author = author
            
            # Create realistic revision history - document appears to have been worked on over time
            created_time = datetime.now() - timedelta(days=7, hours=3, minutes=15)  # Created ~1 week ago
            modified_time = datetime.now() - timedelta(hours=2, minutes=30)  # Last modified a few hours ago
            
            doc.core_properties.created = created_time
            doc.core_properties.modified = modified_time
            doc.core_properties.last_modified_by = author
            
            # Add professional metadata with Marketing Manager defaults
            doc.core_properties.subject = webhook_data.get('subject', f"Steve Glen - Marketing Manager Application Materials")
            doc.core_properties.keywords = webhook_data.get('keywords', 'marketing manager, marketing communications, journalism, public relations, data analysis, strategy, business strategy, strategic communications, resume, cover letter, job application, steve glen')
            doc.core_properties.comments = webhook_data.get('comments', 'Professional application materials for Marketing Manager position - Steve Glen')
            doc.core_properties.category = webhook_data.get('category', 'Job Application')
            doc.core_properties.version = webhook_data.get('version', '3.2')
            
            # Add realistic Microsoft Office metadata for authenticity
            doc.core_properties.language = webhook_data.get('language', 'en-CA')
            doc.core_properties.identifier = webhook_data.get('identifier', 'Microsoft Office Word')
            doc.core_properties.content_status = webhook_data.get('content_status', 'Final')
            
            # Add location/origin information
            location = webhook_data.get('location', 'Edmonton, Alberta, Canada')
            if 'location_meta' not in webhook_data:
                # Add location info to comments for authenticity
                current_comments = doc.core_properties.comments
                doc.core_properties.comments = f"{current_comments} | Created in {location}"
            
            # Note: Company property is not available in python-docx core_properties
            # Company info can be included in document content or custom properties if needed
            
            # Add title
            title_paragraph = doc.add_heading(title, level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            if author or date:
                meta_paragraph = doc.add_paragraph()
                meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if author:
                    meta_paragraph.add_run(f"Author: {author}")
                if author and date:
                    meta_paragraph.add_run(" | ")
                if date:
                    meta_paragraph.add_run(f"Date: {date}")
                meta_paragraph.add_run("\n")
            
            # Add main content
            if content:
                doc.add_paragraph(content)
                doc.add_paragraph()  # Add spacing
            
            # Add sections if provided
            if sections and isinstance(sections, list):
                for section in sections:
                    if isinstance(section, dict):
                        section_title = section.get('title', 'Section')
                        section_content = section.get('content', '')
                        section_level = section.get('level', 2)
                        
                        # Add section heading
                        doc.add_heading(section_title, level=min(section_level, 4))
                        
                        # Add section content
                        if section_content:
                            doc.add_paragraph(section_content)
                            
                        # Add subsections if any
                        subsections = section.get('subsections', [])
                        if subsections and isinstance(subsections, list):
                            for subsection in subsections:
                                if isinstance(subsection, dict):
                                    subsection_title = subsection.get('title', 'Subsection')
                                    subsection_content = subsection.get('content', '')
                                    
                                    doc.add_heading(subsection_title, level=3)
                                    if subsection_content:
                                        doc.add_paragraph(subsection_content)
                    elif isinstance(section, str):
                        # Simple string section
                        doc.add_paragraph(section)
            
            # Apply formatting if specified
            if formatting and isinstance(formatting, dict):
                # Note: Advanced formatting would require more complex docx manipulation
                # This is a basic implementation
                font_size = formatting.get('font_size')
                if font_size:
                    logging.info(f"Font size formatting requested: {font_size} (basic implementation)")
            
            # Generate unique filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            unique_id = str(uuid.uuid4())[:8]
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_title = safe_title.replace(' ', '_')[:50]  # Limit length
            
            filename = f"{timestamp}_{safe_title}_{unique_id}.docx"
            
            # Upload directly to Replit Object Storage
            object_storage_path = None
            file_path = None
            file_size = 0
            
            if self.storage_client:
                try:
                    # Save document to bytes buffer for upload
                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    file_size = len(doc_buffer.getvalue())
                    
                    # Upload to object storage
                    object_storage_path = f"documents/{filename}"
                    self.storage_client.upload_from_bytes(object_storage_path, doc_buffer.getvalue())
                    logging.info(f"Document uploaded to object storage: {object_storage_path}")
                        
                except Exception as storage_error:
                    logging.error(f"Object storage upload failed: {str(storage_error)}")
                    # Fallback to local storage only if cloud fails
                    file_path = os.path.join(self.storage_dir, filename)
                    doc.save(file_path)
                    file_size = os.path.getsize(file_path)
                    logging.info(f"Using local storage fallback for: {filename}")
            else:
                logging.warning("Object storage not available, using local storage")
                file_path = os.path.join(self.storage_dir, filename)
                doc.save(file_path)
                file_size = os.path.getsize(file_path)
            
            file_info = {
                "filename": filename,
                "file_path": object_storage_path if object_storage_path else file_path,
                "file_size": file_size,
                "file_size_mb": round(file_size / (1024 * 1024), 2),
                "download_url": f"/download/{filename}",
                "direct_download_url": f"/download/{filename}",  # For Make.com to download directly
                "title": title,
                "author": author,
                "created_at": datetime.now().isoformat(),
                "storage_type": "object_storage" if object_storage_path else "local",
                "object_storage_path": object_storage_path
            }
            
            logging.info(f"Document saved: {filename} ({file_size} bytes)")
            return file_info
            
        except Exception as e:
            logging.error(f"Document generation error: {str(e)}")
            raise Exception(f"Failed to generate document: {str(e)}")
    
    def cleanup_old_files(self, max_age_hours=24):
        """
        Clean up old generated files
        
        Args:
            max_age_hours (int): Maximum age of files to keep in hours
        """
        try:
            current_time = datetime.now()
            files_deleted = 0
            
            for filename in os.listdir(self.storage_dir):
                if filename.endswith('.docx'):
                    file_path = os.path.join(self.storage_dir, filename)
                    file_age = current_time - datetime.fromtimestamp(os.path.getctime(file_path))
                    
                    if file_age.total_seconds() > (max_age_hours * 3600):
                        os.remove(file_path)
                        files_deleted += 1
                        logging.info(f"Deleted old file: {filename}")
            
            if files_deleted > 0:
                logging.info(f"Cleanup completed: {files_deleted} files deleted")
                
        except Exception as e:
            logging.error(f"File cleanup error: {str(e)}")
