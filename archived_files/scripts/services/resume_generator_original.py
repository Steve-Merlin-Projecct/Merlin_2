import os
import logging
import uuid
import io
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from replit.object_storage import Client

class ResumeGenerator:
    """
    Handles structured resume generation from webhook data
    Based on Harvard MCS Resume Template format
    """
    
    def __init__(self):
        """Initialize the resume generator with object storage"""
        try:
            self.storage_client = Client()
            logging.info("Resume generator connected to Replit Object Storage (default bucket)")
        except Exception as e:
            logging.error(f"Failed to connect to object storage: {e}")
            self.storage_client = None

    def generate_resume(self, resume_data):
        """
        Generate a structured resume from comprehensive resume data
        
        Args:
            resume_data (dict): Structured resume data with all sections
            
        Returns:
            dict: File information including path, filename, and size
        """
        
        try:
            # Create new document
            doc = Document()
            
            # Set document margins (narrow margins for more content)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Extract personal information with Steve Glen defaults
            personal = resume_data.get('personal', {})
            first_name = personal.get('first_name', 'Steve')
            last_name = personal.get('last_name', 'Glen')
            full_name = f"{first_name} {last_name}"
            
            # Set enhanced document properties for authenticity
            doc.core_properties.title = resume_data.get('title', f"{full_name} - Marketing Manager Resume")
            doc.core_properties.author = full_name
            
            # Create realistic revision history
            created_time = datetime.now() - timedelta(days=7, hours=3, minutes=15)
            modified_time = datetime.now() - timedelta(hours=2, minutes=30)
            
            doc.core_properties.created = created_time
            doc.core_properties.modified = modified_time
            doc.core_properties.last_modified_by = full_name
            
            # Add enhanced metadata
            doc.core_properties.subject = resume_data.get('subject', f"{full_name} - Professional Resume")
            doc.core_properties.keywords = resume_data.get('keywords', 'marketing manager, marketing communications, journalism, public relations, data analysis, strategy, business strategy, strategic communications, resume, steve glen')
            doc.core_properties.comments = resume_data.get('comments', f'Professional resume for {full_name} - Marketing Manager position | Created in Edmonton, Alberta, Canada')
            doc.core_properties.category = resume_data.get('category', 'Professional Resume')
            doc.core_properties.version = resume_data.get('version', '3.2')
            doc.core_properties.language = resume_data.get('language', 'en-CA')
            doc.core_properties.identifier = resume_data.get('identifier', 'Microsoft Office Word')
            doc.core_properties.content_status = resume_data.get('content_status', 'Final')
            
            # Generate header section
            self._create_header(doc, personal)
            
            # Generate professional summary if provided
            professional_summary = resume_data.get('professional_summary', '')
            if professional_summary:
                self._create_professional_summary(doc, professional_summary)
            
            # Generate education section
            education_data = resume_data.get('education', [])
            if education_data:
                self._create_education_section(doc, education_data)
            
            # Generate experience section
            experience_data = resume_data.get('experience', [])
            if experience_data:
                self._create_experience_section(doc, experience_data)
                
            # Generate leadership & activities section
            leadership_data = resume_data.get('leadership', [])
            if leadership_data:
                self._create_leadership_section(doc, leadership_data)
                
            # Generate skills & interests section
            skills_data = resume_data.get('skills', {})
            if skills_data:
                self._create_skills_section(doc, skills_data)
            
            # Generate unique filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            unique_id = str(uuid.uuid4())[:8]
            safe_name = "".join(c for c in full_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_name = safe_name.replace(' ', '_')
            
            filename = f"{timestamp}_{safe_name}_Resume_{unique_id}.docx"
            
            # Upload to object storage
            doc_stream = io.BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)
            
            file_size = len(doc_stream.getvalue())
            
            if self.storage_client:
                # Upload to Replit Object Storage
                storage_path = f"documents/{filename}"
                doc_stream.seek(0)
                self.storage_client.upload_from_bytes(storage_path, doc_stream.getvalue())
                logging.info(f"Resume uploaded to object storage: {storage_path}")
            else:
                # Fallback: save locally (not recommended for production)
                os.makedirs('storage', exist_ok=True)
                local_path = os.path.join('storage', filename)
                doc_stream.seek(0)
                with open(local_path, 'wb') as f:
                    f.write(doc_stream.getvalue())
                logging.info(f"Resume saved locally: {local_path}")
            
            logging.info(f"Resume generated successfully: {filename}")
            
            return {
                'filename': filename,
                'file_path': f"documents/{filename}" if self.storage_client else f"storage/{filename}",
                'file_size': file_size,
                'file_size_mb': round(file_size / (1024 * 1024), 2),
                'title': doc.core_properties.title,
                'author': full_name,
                'created_at': datetime.now().isoformat(),
                'storage_type': 'object_storage' if self.storage_client else 'local',
                'object_storage_path': f"documents/{filename}" if self.storage_client else None
            }
            
        except Exception as e:
            logging.error(f"Error generating resume: {str(e)}")
            raise
    
    def _create_header(self, doc, personal):
        """Create the header section with name and contact info"""
        # Full name (centered, large, bold)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(f"{personal.get('first_name', 'Steve')} {personal.get('last_name', 'Glen')}")
        name_run.bold = True
        name_run.font.size = Pt(16)
        
        # Contact information (centered, smaller)
        contact_parts = []
        if personal.get('address'):
            contact_parts.append(personal['address'])
        if personal.get('city') and personal.get('province'):
            contact_parts.append(f"{personal['city']}, {personal['province']} {personal.get('postal_code', '')}")
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
            
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(' • '.join(contact_parts))
            contact_run.font.size = Pt(11)
        
        # Add spacing
        doc.add_paragraph()
    
    def _create_education_section(self, doc, education_data):
        """Create the education section"""
        # Section heading
        edu_heading = doc.add_heading('Education', level=1)
        edu_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for edu in education_data:
            # Institution name and location
            if edu.get('institution'):
                inst_para = doc.add_paragraph()
                inst_run = inst_para.add_run(edu['institution'])
                inst_run.bold = True
                
                if edu.get('location'):
                    # Add location on same line, right-aligned
                    location_run = inst_para.add_run(f"\t{edu['location']}")
                    # Note: True right alignment requires more complex formatting
            
            # Degree and graduation date (GPA removed as requested)
            degree_para = doc.add_paragraph()
            degree_parts = []
            if edu.get('degree'):
                degree_parts.append(edu['degree'])
            if edu.get('concentration'):
                degree_parts.append(edu['concentration'])
                
            if degree_parts:
                degree_run = degree_para.add_run('. '.join(degree_parts))
                
                if edu.get('graduation_date'):
                    grad_run = degree_para.add_run(f"\t{edu['graduation_date']}")
            
            # Optional fields
            if edu.get('thesis'):
                thesis_para = doc.add_paragraph()
                thesis_para.add_run(f"Thesis: {edu['thesis']}")
                
            if edu.get('coursework'):
                course_para = doc.add_paragraph()
                course_para.add_run(f"Relevant Coursework: {edu['coursework']}")
                
            if edu.get('study_abroad'):
                abroad_para = doc.add_paragraph()
                abroad_para.add_run(f"Study Abroad: {edu['study_abroad']}")
                
            # Add spacing between education entries
            doc.add_paragraph()
    
    def _create_professional_summary(self, doc, summary):
        """Create the professional summary section"""
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(summary)
        summary_run.font.size = Pt(11)
        
        # Add spacing
        doc.add_paragraph()
    
    def _create_experience_section(self, doc, experience_data):
        """Create the experience section"""
        # Section heading
        exp_heading = doc.add_heading('Experience', level=1)
        exp_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for exp in experience_data:
            # Organization name and location
            if exp.get('organization'):
                org_para = doc.add_paragraph()
                org_run = org_para.add_run(exp['organization'])
                org_run.bold = True
                
                if exp.get('location'):
                    location_run = org_para.add_run(f"\t{exp['location']}")
            
            # Position title and dates
            if exp.get('position'):
                pos_para = doc.add_paragraph()
                pos_run = pos_para.add_run(exp['position'])
                pos_run.italic = True
                
                if exp.get('start_date') or exp.get('end_date'):
                    start = exp.get('start_date', '')
                    end = exp.get('end_date', 'Present')
                    date_run = pos_para.add_run(f"\t{start} – {end}")
            
            # Bullet points for responsibilities/achievements
            if exp.get('bullets'):
                for bullet in exp['bullets']:
                    bullet_para = doc.add_paragraph()
                    bullet_para.style = 'List Bullet'
                    bullet_para.add_run(bullet)
            
            # Add spacing between experience entries
            doc.add_paragraph()
    
    def _create_leadership_section(self, doc, leadership_data):
        """Create the leadership & activities section"""
        # Section heading
        leader_heading = doc.add_heading('Leadership & Activities', level=1)
        leader_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for activity in leadership_data:
            # Organization name and location
            if activity.get('organization'):
                org_para = doc.add_paragraph()
                org_run = org_para.add_run(activity['organization'])
                org_run.bold = True
                
                if activity.get('location'):
                    location_run = org_para.add_run(f"\t{activity['location']}")
            
            # Position title and dates
            if activity.get('position'):
                pos_para = doc.add_paragraph()
                pos_run = pos_para.add_run(activity['position'])
                pos_run.italic = True
                
                if activity.get('start_date') or activity.get('end_date'):
                    start = activity.get('start_date', '')
                    end = activity.get('end_date', 'Present')
                    date_run = pos_para.add_run(f"\t{start} – {end}")
            
            # Description or bullet points
            if activity.get('description'):
                desc_para = doc.add_paragraph()
                desc_para.add_run(activity['description'])
            elif activity.get('bullets'):
                for bullet in activity['bullets']:
                    bullet_para = doc.add_paragraph()
                    bullet_para.style = 'List Bullet'
                    bullet_para.add_run(bullet)
            
            # Add spacing
            doc.add_paragraph()
    
    def _create_skills_section(self, doc, skills_data):
        """Create the skills section based on Steve Glen's actual skills structure"""
        # Section heading
        skills_heading = doc.add_heading('Skills', level=1)
        skills_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Digital Marketing & Strategy
        if skills_data.get('digital_marketing'):
            dm_para = doc.add_paragraph()
            dm_para.add_run('Digital Marketing & Strategy: ').bold = True
            dm_para.add_run(skills_data['digital_marketing'])
        
        # Technical Expertise
        if skills_data.get('technical_expertise'):
            tech_para = doc.add_paragraph()
            tech_para.add_run('Technical Expertise: ').bold = True
            tech_para.add_run(skills_data['technical_expertise'])
        
        # Business Analytics & IT
        if skills_data.get('business_analytics'):
            ba_para = doc.add_paragraph()
            ba_para.add_run('Business Analytics & IT: ').bold = True
            ba_para.add_run(skills_data['business_analytics'])
        
        # Productivity Software
        if skills_data.get('productivity_software'):
            prod_para = doc.add_paragraph()
            prod_para.add_run('Productivity Software: ').bold = True
            prod_para.add_run(skills_data['productivity_software'])
        
        # Legacy support for original fields
        if skills_data.get('technical'):
            tech_para = doc.add_paragraph()
            tech_para.add_run('Technical: ').bold = True
            tech_para.add_run(skills_data['technical'])
        
        if skills_data.get('language'):
            lang_para = doc.add_paragraph()
            lang_para.add_run('Language: ').bold = True
            lang_para.add_run(skills_data['language'])
        
        if skills_data.get('interests'):
            int_para = doc.add_paragraph()
            int_para.add_run('Interests: ').bold = True
            int_para.add_run(skills_data['interests'])