import os
import json
from datetime import datetime
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .base_generator import BaseDocumentGenerator
import logging

class ResumeGenerator(BaseDocumentGenerator):
    """
    Handles structured resume generation from webhook data
    Based on Harvard MCS Resume Template format
    """
    
    def __init__(self):
        """Initialize the resume generator"""
        super().__init__()
        self.template_path = "steve_glen_resume_template.json"
    
    def generate_resume(self, resume_data):
        """
        Generate a structured resume from comprehensive resume data
        
        Args:
            resume_data (dict): Structured resume data with all sections
            
        Returns:
            dict: File information including path, filename, and size
        """
        # Create job record at the very beginning
        job_id = None
        document_type = "resume"
        
        try:
            # Load default template data
            default_data = self._load_template_data()
            
            # Merge provided data with defaults
            merged_data = self._merge_with_defaults(resume_data, default_data)
            
            # Extract metadata
            title = merged_data.get('title', 'Steve Glen - Marketing Manager Resume')
            author = f"{merged_data['personal']['first_name']} {merged_data['personal']['last_name']}"
            subject = merged_data.get('subject', 'Steve Glen - Professional Marketing Manager Resume')
            keywords = merged_data.get('keywords', 'Marketing Communications, Journalism, Public Relations, Data Analysis, Strategy, Business Strategy, Strategic Communications')
            
            # Create job record with complete resume data
            job_id = self._create_job_record(
                document_type=document_type,
                webhook_data=resume_data,  # Store original input
                title=title,
                author=author
            )
            
            # Create document with professional metadata
            doc = self._create_document_with_metadata(
                title=title,
                author=author,
                subject=subject,
                keywords=keywords,
                category="Resume"
            )
            
            # Generate document sections
            personal = merged_data.get('personal', {})
            self._create_header(doc, personal)
            
            # Professional summary
            professional_summary = merged_data.get('professional_summary', '')
            if professional_summary:
                self._create_professional_summary(doc, professional_summary)
            
            # Education section
            education_data = merged_data.get('education', [])
            if education_data:
                self._create_education_section(doc, education_data)
            
            # Experience section
            experience_data = merged_data.get('experience', [])
            if experience_data:
                self._create_experience_section(doc, experience_data)
            
            # Leadership & activities section
            leadership_data = merged_data.get('leadership', [])
            if leadership_data:
                self._create_leadership_section(doc, leadership_data)
            
            # Additional experience section
            additional_experience = merged_data.get('additional_experience', [])
            if additional_experience:
                self._create_additional_experience_section(doc, additional_experience)
            
            # Certifications section
            certifications = merged_data.get('certifications', [])
            if certifications:
                self._create_certifications_section(doc, certifications)
            
            # Skills section
            skills_data = merged_data.get('skills', {})
            if skills_data:
                self._create_skills_section(doc, skills_data)
            
            # Save document with database tracking
            base_filename = f"{merged_data['personal']['first_name']}_{merged_data['personal']['last_name']}_Resume"
            file_info = self._save_document(
                doc, 
                base_filename, 
                title, 
                author, 
                job_id=job_id, 
                document_type=document_type
            )
            
            logging.info(f"Resume generated successfully: {file_info['filename']}")
            return file_info
            
        except Exception as e:
            logging.error(f"Error generating resume: {e}")
            # Handle generation error with database tracking
            self._handle_generation_error(
                job_id=job_id,
                document_type=document_type,
                error_code='GENERATION_ERROR',
                error_message=str(e),
                error_details={'resume_data': resume_data}
            )
            raise Exception(f"Failed to generate resume: {str(e)}")
    
    def _load_template_data(self):
        """Load default template data from JSON file"""
        try:
            with open(self.template_path, 'r') as f:
                return json.load(f)
        except FileNotFoundError:
            logging.warning(f"Template file not found: {self.template_path}")
            return {}
        except json.JSONDecodeError as e:
            logging.error(f"Invalid JSON in template file: {e}")
            return {}
    
    def _merge_with_defaults(self, provided_data, default_data):
        """Merge provided data with template defaults"""
        merged = default_data.copy()
        
        # Deep merge for nested dictionaries
        for key, value in provided_data.items():
            if isinstance(value, dict) and key in merged and isinstance(merged[key], dict):
                merged[key].update(value)
            else:
                merged[key] = value
        
        return merged
    
    def _create_header(self, doc, personal):
        """Create the header section with name and contact info"""
        # Name (larger font)
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(f"{personal.get('first_name', '')} {personal.get('last_name', '')}")
        name_run.font.size = Pt(16)
        name_run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact information
        contact_parts = []
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        if personal.get('city') and personal.get('province'):
            contact_parts.append(f"{personal['city']}, {personal['province']}")
        if personal.get('linkedin'):
            contact_parts.append(personal['linkedin'])
        
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(11)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
    
    def _create_professional_summary(self, doc, summary):
        """Create the professional summary section"""
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(summary)
        summary_run.font.size = Pt(11)
        
        # Add spacing
        doc.add_paragraph()
    
    def _create_education_section(self, doc, education_data):
        """Create the education section"""
        self._add_section_heading(doc, 'Education')
        
        for edu in education_data:
            # Institution name and location
            if edu.get('institution'):
                inst_para = doc.add_paragraph()
                inst_run = inst_para.add_run(edu['institution'])
                inst_run.bold = True
                inst_run.font.size = Pt(11)
                
                if edu.get('location'):
                    location_run = inst_para.add_run(f"\t{edu['location']}")
            
            # Degree and graduation date (GPA removed as requested)
            degree_para = doc.add_paragraph()
            degree_parts = []
            if edu.get('degree'):
                degree_parts.append(edu['degree'])
            if edu.get('concentration'):
                degree_parts.append(edu['concentration'])
                
            if degree_parts:
                degree_run = degree_para.add_run('. '.join(degree_parts))
                
                if edu.get('graduation_date'):
                    grad_run = degree_para.add_run(f"\t{edu['graduation_date']}")
            
            # Relevant coursework
            if edu.get('relevant_coursework'):
                course_para = doc.add_paragraph()
                course_para.add_run('Relevant Coursework: ').bold = True
                course_para.add_run(edu['relevant_coursework'])
            
            # Honors
            if edu.get('honors'):
                honors_para = doc.add_paragraph()
                honors_para.add_run('Honors: ').bold = True
                honors_para.add_run(edu['honors'])
            
            # Study abroad
            if edu.get('study_abroad'):
                abroad_para = doc.add_paragraph()
                abroad_para.add_run(f"Study Abroad: {edu['study_abroad']}")
                
            # Add spacing between education entries
            doc.add_paragraph()
    
    def _create_experience_section(self, doc, experience_data):
        """Create the experience section"""
        self._add_section_heading(doc, 'Experience')
        
        for exp in experience_data:
            # Company and location
            if exp.get('company'):
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(exp['company'])
                company_run.bold = True
                company_run.font.size = Pt(11)
                
                if exp.get('location'):
                    location_run = company_para.add_run(f"\t{exp['location']}")
            
            # Position and dates
            if exp.get('position'):
                position_para = doc.add_paragraph()
                position_run = position_para.add_run(exp['position'])
                position_run.italic = True
                
                if exp.get('start_date'):
                    date_text = exp['start_date']
                    if exp.get('end_date'):
                        date_text += f" - {exp['end_date']}"
                    date_run = position_para.add_run(f"\t{date_text}")
            
            # Description or bullet points
            if exp.get('description'):
                desc_para = doc.add_paragraph()
                desc_para.add_run(exp['description'])
            
            if exp.get('bullets'):
                for bullet in exp['bullets']:
                    bullet_para = doc.add_paragraph()
                    bullet_para.add_run(f"• {bullet}")
            
            # Add spacing between experience entries
            doc.add_paragraph()
    
    def _create_leadership_section(self, doc, leadership_data):
        """Create the leadership & activities section"""
        self._add_section_heading(doc, 'Leadership & Activities')
        
        for leader in leadership_data:
            # Organization and location
            if leader.get('organization'):
                org_para = doc.add_paragraph()
                org_run = org_para.add_run(leader['organization'])
                org_run.bold = True
                org_run.font.size = Pt(11)
                
                if leader.get('location'):
                    location_run = org_para.add_run(f"\t{leader['location']}")
            
            # Position and dates
            if leader.get('position'):
                position_para = doc.add_paragraph()
                position_run = position_para.add_run(leader['position'])
                position_run.italic = True
                
                if leader.get('start_date'):
                    date_text = leader['start_date']
                    if leader.get('end_date'):
                        date_text += f" - {leader['end_date']}"
                    date_run = position_para.add_run(f"\t{date_text}")
            
            # Description or bullet points
            if leader.get('description'):
                desc_para = doc.add_paragraph()
                desc_para.add_run(leader['description'])
            
            if leader.get('bullets'):
                for bullet in leader['bullets']:
                    bullet_para = doc.add_paragraph()
                    bullet_para.add_run(f"• {bullet}")
            
            # Add spacing
            doc.add_paragraph()
    
    def _create_additional_experience_section(self, doc, additional_experience):
        """Create the additional experience section"""
        self._add_section_heading(doc, 'Additional Experience')
        
        for exp in additional_experience:
            exp_para = doc.add_paragraph()
            
            if exp.get('role'):
                role_run = exp_para.add_run(exp['role'])
                role_run.bold = True
                
                if exp.get('details'):
                    exp_para.add_run(f": {exp['details']}")
                elif exp.get('year'):
                    exp_para.add_run(f" ({exp['year']})")
                elif exp.get('years'):
                    exp_para.add_run(f" ({exp['years']})")
        
        # Add spacing
        doc.add_paragraph()
    
    def _create_certifications_section(self, doc, certifications):
        """Create the certifications section"""
        self._add_section_heading(doc, 'Certifications')
        
        for cert_group in certifications:
            if cert_group.get('provider'):
                provider_para = doc.add_paragraph()
                provider_run = provider_para.add_run(cert_group['provider'])
                provider_run.bold = True
                
                if cert_group.get('certifications'):
                    for cert in cert_group['certifications']:
                        cert_para = doc.add_paragraph()
                        cert_para.add_run(f"• {cert}")
        
        # Add spacing
        doc.add_paragraph()
    
    def _create_skills_section(self, doc, skills_data):
        """Create the skills section based on Steve Glen's actual skills structure"""
        self._add_section_heading(doc, 'Skills')
        
        # Digital Marketing & Strategy
        if skills_data.get('digital_marketing'):
            self._add_paragraph_with_formatting(doc, skills_data['digital_marketing'], 'Digital Marketing & Strategy: ')
        
        # Technical Expertise
        if skills_data.get('technical_expertise'):
            self._add_paragraph_with_formatting(doc, skills_data['technical_expertise'], 'Technical Expertise: ')
        
        # Business Analytics & IT
        if skills_data.get('business_analytics'):
            self._add_paragraph_with_formatting(doc, skills_data['business_analytics'], 'Business Analytics & IT: ')
        
        # Productivity Software
        if skills_data.get('productivity_software'):
            self._add_paragraph_with_formatting(doc, skills_data['productivity_software'], 'Productivity Software: ')
        
        # Legacy support for original fields
        if skills_data.get('technical'):
            self._add_paragraph_with_formatting(doc, skills_data['technical'], 'Technical: ')
        
        if skills_data.get('language'):
            self._add_paragraph_with_formatting(doc, skills_data['language'], 'Language: ')
        
        if skills_data.get('interests'):
            self._add_paragraph_with_formatting(doc, skills_data['interests'], 'Interests: ')