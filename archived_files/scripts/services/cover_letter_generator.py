import os
import json
from datetime import datetime
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .base_generator import BaseDocumentGenerator
import logging

class CoverLetterGenerator(BaseDocumentGenerator):
    """
    Handles cover letter generation from webhook data
    Creates professional cover letters with proper formatting
    """
    
    def __init__(self):
        """Initialize the cover letter generator"""
        super().__init__()
        self.template_path = "steve_glen_cover_letter_template.json"
    
    def generate_cover_letter(self, letter_data):
        """
        Generate a professional cover letter from provided data
        
        Args:
            letter_data (dict): Cover letter data including personal info, company details, content
            
        Returns:
            dict: File information including path, filename, and size
        """
        # Create job record at the very beginning
        job_id = None
        document_type = "cover_letter"
        
        try:
            # Load default template data
            default_data = self._load_template_data()
            
            # Merge provided data with defaults
            merged_data = self._merge_with_defaults(letter_data, default_data)
            
            # Extract metadata
            company_name = merged_data.get('company', {}).get('name', 'Company')
            position = merged_data.get('position', 'Marketing Manager')
            author = f"{merged_data['personal']['first_name']} {merged_data['personal']['last_name']}"
            
            title = f"{author} - {position} Cover Letter - {company_name}"
            subject = f"Cover Letter - {position} Application"
            keywords = "Cover Letter, Job Application, Marketing, Professional Communication"
            
            # Create job record with complete cover letter data
            job_id = self._create_job_record(
                document_type=document_type,
                webhook_data=letter_data,  # Store original input
                title=title,
                author=author
            )
            
            # Create document with professional metadata
            doc = self._create_document_with_metadata(
                title=title,
                author=author,
                subject=subject,
                keywords=keywords,
                category="Cover Letter"
            )
            
            # Set document margins for professional letter format
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Generate document sections
            self._create_header(doc, merged_data.get('personal', {}))
            self._create_date_and_recipient(doc, merged_data)
            self._create_letter_body(doc, merged_data)
            self._create_closing(doc, merged_data)
            
            # Save document with database tracking
            safe_company_name = company_name.replace(' ', '_').replace(',', '').replace('.', '')
            base_filename = f"{merged_data['personal']['first_name']}_{merged_data['personal']['last_name']}_Cover_Letter_{safe_company_name}"
            file_info = self._save_document(
                doc, 
                base_filename, 
                title, 
                author, 
                job_id=job_id, 
                document_type=document_type
            )
            
            logging.info(f"Cover letter generated successfully: {file_info['filename']}")
            return file_info
            
        except Exception as e:
            logging.error(f"Error generating cover letter: {e}")
            # Handle generation error with database tracking
            self._handle_generation_error(
                job_id=job_id,
                document_type=document_type,
                error_code='GENERATION_ERROR',
                error_message=str(e),
                error_details={'letter_data': letter_data}
            )
            raise Exception(f"Failed to generate cover letter: {str(e)}")
    
    def _load_template_data(self):
        """Load default template data from JSON file"""
        try:
            with open(self.template_path, 'r') as f:
                return json.load(f)
        except FileNotFoundError:
            logging.warning(f"Template file not found: {self.template_path}, using built-in defaults")
            return self._get_default_template()
        except json.JSONDecodeError as e:
            logging.error(f"Invalid JSON in template file: {e}")
            return self._get_default_template()
    
    def _get_default_template(self):
        """Return default template data for Steve Glen"""
        return {
            "personal": {
                "first_name": "Steve",
                "last_name": "Glen",
                "email": "therealstevenglen@gmail.com",
                "phone": "780-884-7038",
                "city": "Edmonton",
                "province": "Alberta",
                "postal_code": "",
                "address": ""
            },
            "company": {
                "name": "",
                "address": "",
                "city": "",
                "province": "",
                "postal_code": ""
            },
            "recipient": {
                "name": "Hiring Manager",
                "title": "",
                "department": ""
            },
            "position": "Marketing Manager",
            "opening_paragraph": "I am writing to express my strong interest in the {position} position at {company_name}. With over 14 years of experience in marketing communications, content creation, and business strategy, I am confident in my ability to contribute effectively to your team's success.",
            "body_paragraphs": [
                "In my current role as Digital Strategist at Odvod Media, I have successfully managed comprehensive marketing campaigns that have generated over 6 million article reads and 18 million ad views. My expertise spans digital marketing strategy, content management, SEO optimization, and social media marketing across multiple platforms. I have extensive experience with tools including Google Analytics, Adobe Creative Suite, WordPress, and various marketing automation platforms.",
                "My technical expertise complements my marketing skills, with proficiency in data analysis using RStudio and Python, multimedia content creation using DaVinci Resolve and Adobe After Effects, and business analytics dashboard design. I hold Google certifications in Cloud Digital Leader, Data Analytics, and Google Ads, along with specialized training in machine learning foundations and business analytics from the University of Alberta."
            ],
            "closing_paragraph": "I am excited about the opportunity to bring my proven track record of successful marketing campaigns and analytical expertise to {company_name}. I would welcome the opportunity to discuss how my skills and experience can contribute to your marketing objectives. Thank you for considering my application.",
            "signature": "Sincerely,",
            "title": "Steve Glen - Marketing Manager Cover Letter",
            "subject": "Cover Letter - Marketing Manager Application",
            "keywords": "Cover Letter, Marketing Manager, Digital Marketing, Content Strategy, Business Analytics, Google Certifications"
        }
    
    def _merge_with_defaults(self, provided_data, default_data):
        """Merge provided data with template defaults"""
        merged = default_data.copy()
        
        # Deep merge for nested dictionaries
        for key, value in provided_data.items():
            if isinstance(value, dict) and key in merged and isinstance(merged[key], dict):
                merged[key].update(value)
            else:
                merged[key] = value
        
        return merged
    
    def _create_header(self, doc, personal):
        """Create the header section with personal contact information"""
        # Name (larger font, centered)
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(f"{personal.get('first_name', '')} {personal.get('last_name', '')}")
        name_run.font.size = Pt(16)
        name_run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact information (centered)
        contact_parts = []
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(11)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Address (centered)
        address_parts = []
        if personal.get('address'):
            address_parts.append(personal['address'])
        if personal.get('city') and personal.get('province'):
            city_prov = f"{personal['city']}, {personal['province']}"
            if personal.get('postal_code'):
                city_prov += f" {personal['postal_code']}"
            address_parts.append(city_prov)
        
        if address_parts:
            address_para = doc.add_paragraph()
            address_run = address_para.add_run(" | ".join(address_parts))
            address_run.font.size = Pt(11)
            address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
    
    def _create_date_and_recipient(self, doc, letter_data):
        """Create the date and recipient address section"""
        # Date (right-aligned)
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(letter_data.get('date', datetime.now().strftime("%B %d, %Y")))
        date_run.font.size = Pt(11)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add spacing
        doc.add_paragraph()
        
        # Recipient information (left-aligned)
        recipient = letter_data.get('recipient', {})
        company = letter_data.get('company', {})
        
        if recipient.get('name'):
            recipient_name = recipient['name']
            if recipient.get('title'):
                recipient_name += f", {recipient['title']}"
            
            recipient_para = doc.add_paragraph()
            recipient_run = recipient_para.add_run(recipient_name)
            recipient_run.font.size = Pt(11)
        
        if recipient.get('department'):
            dept_para = doc.add_paragraph()
            dept_run = dept_para.add_run(recipient['department'])
            dept_run.font.size = Pt(11)
        
        if company.get('name'):
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(company['name'])
            company_run.font.size = Pt(11)
        
        # Company address
        if company.get('address'):
            addr_para = doc.add_paragraph()
            addr_run = addr_para.add_run(company['address'])
            addr_run.font.size = Pt(11)
        
        if company.get('city') and company.get('province'):
            city_prov = f"{company['city']}, {company['province']}"
            if company.get('postal_code'):
                city_prov += f" {company['postal_code']}"
            
            city_para = doc.add_paragraph()
            city_run = city_para.add_run(city_prov)
            city_run.font.size = Pt(11)
        
        # Add spacing
        doc.add_paragraph()
    
    def _create_letter_body(self, doc, letter_data):
        """Create the main body of the cover letter"""
        # Salutation
        recipient_name = letter_data.get('recipient', {}).get('name', 'Hiring Manager')
        salutation = f"Dear {recipient_name},"
        
        salutation_para = doc.add_paragraph()
        salutation_run = salutation_para.add_run(salutation)
        salutation_run.font.size = Pt(11)
        
        # Add spacing
        doc.add_paragraph()
        
        # Replace placeholders in content
        company_name = letter_data.get('company', {}).get('name', 'your company')
        position = letter_data.get('position', 'the position')
        
        # Opening paragraph
        opening = letter_data.get('opening_paragraph', '')
        if opening:
            opening = opening.replace('{company_name}', company_name)
            opening = opening.replace('{position}', position)
            
            opening_para = doc.add_paragraph()
            opening_run = opening_para.add_run(opening)
            opening_run.font.size = Pt(11)
            doc.add_paragraph()
        
        # Body paragraphs
        body_paragraphs = letter_data.get('body_paragraphs', [])
        for paragraph_text in body_paragraphs:
            paragraph_text = paragraph_text.replace('{company_name}', company_name)
            paragraph_text = paragraph_text.replace('{position}', position)
            
            body_para = doc.add_paragraph()
            body_run = body_para.add_run(paragraph_text)
            body_run.font.size = Pt(11)
            doc.add_paragraph()
        
        # Closing paragraph
        closing = letter_data.get('closing_paragraph', '')
        if closing:
            closing = closing.replace('{company_name}', company_name)
            closing = closing.replace('{position}', position)
            
            closing_para = doc.add_paragraph()
            closing_run = closing_para.add_run(closing)
            closing_run.font.size = Pt(11)
            doc.add_paragraph()
    
    def _create_closing(self, doc, letter_data):
        """Create the closing section of the cover letter"""
        # Signature line
        signature = letter_data.get('signature', 'Sincerely,')
        signature_para = doc.add_paragraph()
        signature_run = signature_para.add_run(signature)
        signature_run.font.size = Pt(11)
        
        # Add space for signature
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Typed name
        name = f"{letter_data['personal']['first_name']} {letter_data['personal']['last_name']}"
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.font.size = Pt(11)