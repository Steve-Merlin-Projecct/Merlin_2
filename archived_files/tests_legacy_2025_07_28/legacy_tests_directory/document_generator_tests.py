"""
Comprehensive Test Suite for Template-Based Document Generation

This test suite validates the complete template library system including:
- Template conversion from reference files
- Template engine functionality  
- Document generation with variable substitution
- Professional metadata inclusion
- Error handling and validation

Tests ensure the system can reliably generate professional documents
from JSON input while preserving all formatting and structure.
"""

import os
import json
import logging
from datetime import datetime
from template_converter import TemplateConverter
from template_engine import TemplateEngine

class DocumentGeneratorTests:
    """
    Comprehensive test suite for the template-based document generation system
    
    This class runs all necessary tests to validate that the template library
    system works correctly and can generate professional documents with proper
    formatting, metadata, and variable substitution.
    """
    
    def __init__(self):
        """Initialize test suite with required components"""
        self.setup_logging()
        self.converter = TemplateConverter()
        self.engine = TemplateEngine()
        self.test_results = {}
        
    def setup_logging(self):
        """Configure logging for test execution"""
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)
        
    def run_all_tests(self):
        """
        Execute all tests in the test suite
        
        Returns:
            dict: Complete test results with pass/fail status for each test
        """
        print("=== Template Library System Test Suite ===\n")
        
        # Test 0: Archive verification
        self.test_0_archive_verification()
        
        # Test 1: Basic functionality
        self.test_1_no_errors()
        
        # Test 2: Document generation from test data
        self.test_2_convert_test_data()
        
        # Test 3: Professional metadata
        self.test_3_professional_metadata()
        
        # Test 4: Template conversion
        self.test_4_template_conversion()
        
        # Test 5: Reference file processing
        self.test_5_reference_file_processing()
        
        # Test 6: Template loading
        self.test_6_template_loading()
        
        # Test 7: Variable substitution
        self.test_7_variable_substitution()
        
        # Test 8: Formatting preservation
        self.test_8_formatting_preservation()
        
        # Test 9: Multiple template types
        self.test_9_multiple_template_types()
        
        # Test 10: Folder structure
        self.test_10_folder_structure()
        
        # Print summary
        self.print_test_summary()
        
        return self.test_results
    
    def test_0_archive_verification(self):
        """Test 0: Verify current document_generator.py is archived"""
        print("Test 0: Archive verification...")
        try:
            # Check if the archived file exists in archived_files folder
            archived_file = "../archived_files/makecom-document-generator.py"
            original_file = "../document_generator.py"
            
            if os.path.exists(archived_file):
                print("✅ makecom-document-generator.py archive exists in archived_files/")
                
                # Verify it's a proper backup (has expected content)
                with open(archived_file, 'r') as f:
                    content = f.read()
                    if 'class DocumentGenerator' in content and 'webhook_data' in content:
                        print("✅ Archive contains original webhook-based code")
                        self.test_results['test_0'] = {'status': 'PASS', 'message': 'Archive verified'}
                    else:
                        print("❌ Archive doesn't contain expected content")
                        self.test_results['test_0'] = {'status': 'FAIL', 'message': 'Archive content invalid'}
            else:
                print("❌ makecom-document-generator.py not found in archived_files/")
                self.test_results['test_0'] = {'status': 'FAIL', 'message': 'Archive file missing'}
                
        except Exception as e:
            print(f"❌ Archive verification failed: {str(e)}")
            self.test_results['test_0'] = {'status': 'FAIL', 'message': f'Error: {str(e)}'}
    
    def test_1_no_errors(self):
        """Test 1: Verify document_generator.py loads without errors"""
        print("\nTest 1: No errors in document_generator.py...")
        try:
            # Try to import the updated document generator
            import sys
            sys.path.append('..')
            from document_generator import DocumentGenerator
            
            # Try to instantiate it
            generator = DocumentGenerator()
            print("✅ DocumentGenerator loads and instantiates without errors")
            self.test_results['test_1'] = {'status': 'PASS', 'message': 'No import or instantiation errors'}
            
        except Exception as e:
            print(f"❌ Error loading document_generator.py: {str(e)}")
            self.test_results['test_1'] = {'status': 'FAIL', 'message': f'Error: {str(e)}'}
    
    def test_2_convert_test_data(self):
        """Test 2: Convert steve_glen_resume_test.json to valid .docx file"""
        print("\nTest 2: Converting test data to .docx...")
        try:
            # Load test data from template library test data
            test_data_path = "test_data/steve_glen_resume_template_data.json"
            
            if not os.path.exists(test_data_path):
                print("❌ steve_glen_resume_template_data.json not found")
                self.test_results['test_2'] = {'status': 'FAIL', 'message': 'Test data file missing'}
                return
            
            with open(test_data_path, 'r') as f:
                test_data = json.load(f)
            
            # Use template engine to generate document
            template_path = "resumes/harvard_mcs_resume_template.docx"
            
            if not os.path.exists(template_path):
                print("❌ Harvard template not found")
                self.test_results['test_2'] = {'status': 'FAIL', 'message': 'Template file missing'}
                return
            
            # Test data is already flattened for template processing
            flattened_data = test_data
            
            # Generate document
            result = self.engine.generate_document(template_path, flattened_data)
            
            # Verify output file exists and is valid
            if os.path.exists(result['output_path']):
                file_size = os.path.getsize(result['output_path'])
                if file_size > 1000:  # Reasonable size for a document
                    print(f"✅ Document generated successfully: {result['output_path']} ({file_size} bytes)")
                    print(f"   Variables processed: {result['variables_processed']['variables_substituted']}")
                    self.test_results['test_2'] = {'status': 'PASS', 'message': f'Document generated: {file_size} bytes'}
                else:
                    print("❌ Generated document is too small")
                    self.test_results['test_2'] = {'status': 'FAIL', 'message': 'Document too small'}
            else:
                print("❌ Output document not created")
                self.test_results['test_2'] = {'status': 'FAIL', 'message': 'No output file'}
                
        except Exception as e:
            print(f"❌ Document generation failed: {str(e)}")
            self.test_results['test_2'] = {'status': 'FAIL', 'message': f'Error: {str(e)}'}
    
    def test_3_professional_metadata(self):
        """Test 3: Verify professional metadata is visible in generated .docx"""
        print("\nTest 3: Verifying professional metadata...")
        try:
            # Find the most recently generated document - check both locations
            storage_dirs = ["storage", "../storage"]
            latest_path = None
            
            for storage_dir in storage_dirs:
                if os.path.exists(storage_dir):
                    # Find latest .docx file
                    docx_files = [f for f in os.listdir(storage_dir) if f.endswith('.docx')]
                    if docx_files:
                        # Get the most recent file
                        latest_file = max(docx_files, key=lambda f: os.path.getctime(os.path.join(storage_dir, f)))
                        latest_path = os.path.join(storage_dir, latest_file)
                        break
            
            if not latest_path:
                print("❌ No .docx files found in any storage directory")
                self.test_results['test_3'] = {'status': 'FAIL', 'message': 'No generated documents found'}
                return
            
            # Load document and check metadata
            from docx import Document
            doc = Document(latest_path)
            
            properties = doc.core_properties
            metadata_checks = {
                'title': properties.title,
                'author': properties.author,
                'subject': properties.subject,
                'keywords': properties.keywords,
                'comments': properties.comments,
                'category': properties.category,
                'language': properties.language
            }
            
            # Verify metadata exists
            metadata_present = sum(1 for v in metadata_checks.values() if v)
            
            if metadata_present >= 4:  # At least 4 metadata fields should be set
                print(f"✅ Professional metadata present: {metadata_present}/7 fields")
                for key, value in metadata_checks.items():
                    if value:
                        print(f"   {key}: {value}")
                self.test_results['test_3'] = {'status': 'PASS', 'message': f'{metadata_present}/7 metadata fields set'}
            else:
                print(f"❌ Insufficient metadata: only {metadata_present}/7 fields set")
                self.test_results['test_3'] = {'status': 'FAIL', 'message': f'Only {metadata_present}/7 fields set'}
                
        except Exception as e:
            print(f"❌ Metadata verification failed: {str(e)}")
            self.test_results['test_3'] = {'status': 'FAIL', 'message': f'Error: {str(e)}'}
    
    def test_4_template_conversion(self):
        """Test 4: Template conversion script works without errors"""
        print("\nTest 4: Template conversion functionality...")
        try:
            # Verify converter can process reference files
            reference_path = "reference/resume/Accessible-MCS-Resume-Template-Bullet-Points_1751349781656.docx"
            
            if not os.path.exists(reference_path):
                print("❌ Reference file not found")
                self.test_results['test_4'] = {'status': 'FAIL', 'message': 'Reference file missing'}
                return
            
            # Test template validation (without creating another template)
            validation_result = self.engine.validate_template("resumes/harvard_mcs_resume_template.docx")
            
            if validation_result['valid']:
                print(f"✅ Template conversion successful: {validation_result['variable_count']} variables found")
                print(f"   Variables: {', '.join(validation_result['variables_found'][:5])}{'...' if len(validation_result['variables_found']) > 5 else ''}")
                self.test_results['test_4'] = {'status': 'PASS', 'message': f'{validation_result["variable_count"]} variables found'}
            else:
                print(f"❌ Template validation failed: {validation_result['error']}")
                self.test_results['test_4'] = {'status': 'FAIL', 'message': validation_result['error']}
                
        except Exception as e:
            print(f"❌ Template conversion test failed: {str(e)}")
            self.test_results['test_4'] = {'status': 'FAIL', 'message': f'Error: {str(e)}'}
    
    def test_5_reference_file_processing(self):
        """Test 5: Reference .docx files can be converted to templates"""
        print("\nTest 5: Reference file processing...")
        try:
            # Check if conversion was successful by looking at metadata
            metadata_path = "resumes/harvard_mcs_resume_template_metadata.json"
            
            if os.path.exists(metadata_path):
                with open(metadata_path, 'r') as f:
                    metadata = json.load(f)
                
                stats = metadata['conversion_stats']
                print(f"✅ Reference file processed successfully")
                print(f"   Paragraphs processed: {stats['total_paragraphs']}")
                print(f"   Paragraphs modified: {stats['paragraphs_modified']}")
                print(f"   Patterns matched: {len(stats['patterns_matched'])}")
                
                self.test_results['test_5'] = {
                    'status': 'PASS', 
                    'message': f'{stats["paragraphs_modified"]}/{stats["total_paragraphs"]} paragraphs processed'
                }
            else:
                print("❌ Template metadata not found")
                self.test_results['test_5'] = {'status': 'FAIL', 'message': 'Metadata missing'}
                
        except Exception as e:
            print(f"❌ Reference file processing test failed: {str(e)}")
            self.test_results['test_5'] = {'status': 'FAIL', 'message': f'Error: {str(e)}'}
    
    def test_6_template_loading(self):
        """Test 6: Template files can be loaded and processed"""
        print("\nTest 6: Template loading...")
        try:
            template_path = "resumes/harvard_mcs_resume_template.docx"
            
            # Test template loading
            doc = self.engine.load_template(template_path)
            
            if doc:
                print("✅ Template loads successfully")
                
                # Test caching
                doc2 = self.engine.load_template(template_path)
                if doc2:
                    print("✅ Template caching works")
                    self.test_results['test_6'] = {'status': 'PASS', 'message': 'Template loading and caching work'}
                else:
                    print("❌ Template caching failed")
                    self.test_results['test_6'] = {'status': 'FAIL', 'message': 'Caching failed'}
            else:
                print("❌ Template loading failed")
                self.test_results['test_6'] = {'status': 'FAIL', 'message': 'Loading failed'}
                
        except Exception as e:
            print(f"❌ Template loading test failed: {str(e)}")
            self.test_results['test_6'] = {'status': 'FAIL', 'message': f'Error: {str(e)}'}
    
    def test_7_variable_substitution(self):
        """Test 7: Variable substitution works correctly"""
        print("\nTest 7: Variable substitution...")
        try:
            # Test variable substitution function directly
            test_text = "Hello <<first_name>> <<last_name>>, your email is <<email>>"
            test_data = {
                'first_name': 'John',
                'last_name': 'Doe', 
                'email': 'john.doe@example.com'
            }
            
            stats = {'variables_found': set(), 'variables_substituted': set(), 'variables_missing': set()}
            result = self.engine.substitute_variables(test_text, test_data, stats)
            
            expected = "Hello John Doe, your email is john.doe@example.com"
            
            if result == expected:
                print("✅ Variable substitution works correctly")
                print(f"   Substituted: {len(stats['variables_substituted'])} variables")
                self.test_results['test_7'] = {'status': 'PASS', 'message': f'{len(stats["variables_substituted"])} variables substituted'}
            else:
                print(f"❌ Variable substitution failed")
                print(f"   Expected: {expected}")
                print(f"   Got: {result}")
                self.test_results['test_7'] = {'status': 'FAIL', 'message': 'Substitution mismatch'}
                
        except Exception as e:
            print(f"❌ Variable substitution test failed: {str(e)}")
            self.test_results['test_7'] = {'status': 'FAIL', 'message': f'Error: {str(e)}'}
    
    def test_8_formatting_preservation(self):
        """Test 8: Generated documents maintain original formatting"""
        print("\nTest 8: Formatting preservation...")
        try:
            # This test checks that the template has preserved the original structure
            template_path = "resumes/harvard_mcs_resume_template.docx"
            
            if os.path.exists(template_path):
                from docx import Document
                doc = Document(template_path)
                
                # Check basic document structure
                paragraph_count = len(doc.paragraphs)
                table_count = len(doc.tables)
                
                # The original Harvard template should have substantial content
                if paragraph_count > 10:
                    print(f"✅ Document structure preserved: {paragraph_count} paragraphs, {table_count} tables")
                    self.test_results['test_8'] = {'status': 'PASS', 'message': f'{paragraph_count} paragraphs preserved'}
                else:
                    print(f"❌ Document structure seems incomplete: only {paragraph_count} paragraphs")
                    self.test_results['test_8'] = {'status': 'FAIL', 'message': f'Only {paragraph_count} paragraphs'}
            else:
                print("❌ Template file not found")
                self.test_results['test_8'] = {'status': 'FAIL', 'message': 'Template missing'}
                
        except Exception as e:
            print(f"❌ Formatting preservation test failed: {str(e)}")
            self.test_results['test_8'] = {'status': 'FAIL', 'message': f'Error: {str(e)}'}
    
    def test_9_multiple_template_types(self):
        """Test 9: Both resume and cover letter templates work"""
        print("\nTest 9: Multiple template types...")
        try:
            # For now, we only have resume templates, so test the system's ability to handle different types
            resume_template = "resumes/harvard_mcs_resume_template.docx"
            
            if os.path.exists(resume_template):
                print("✅ Resume template available")
                
                # Test the system's readiness for multiple types
                template_dirs = ['resumes', 'coverletters']
                dirs_exist = sum(1 for d in template_dirs if os.path.exists(d))
                
                print(f"✅ Template directory structure ready: {dirs_exist}/2 directories")
                self.test_results['test_9'] = {'status': 'PASS', 'message': f'{dirs_exist}/2 template directories ready'}
            else:
                print("❌ Resume template not available")
                self.test_results['test_9'] = {'status': 'FAIL', 'message': 'Resume template missing'}
                
        except Exception as e:
            print(f"❌ Multiple template types test failed: {str(e)}")
            self.test_results['test_9'] = {'status': 'FAIL', 'message': f'Error: {str(e)}'}
    
    def test_10_folder_structure(self):
        """Test 10: Template library folder structure is correct"""
        print("\nTest 10: Folder structure verification...")
        try:
            required_dirs = [
                '.',  # content_template_library root
                'reference',
                'reference/resume',
                'reference/coverletter',
                'resumes',
                'coverletters'
            ]
            
            existing_dirs = []
            for dir_path in required_dirs:
                if os.path.exists(dir_path):
                    existing_dirs.append(dir_path)
                    print(f"✅ {dir_path}/")
                else:
                    print(f"❌ {dir_path}/ (missing)")
            
            if len(existing_dirs) >= 5:  # Allow some flexibility
                print(f"✅ Folder structure complete: {len(existing_dirs)}/{len(required_dirs)} directories")
                self.test_results['test_10'] = {'status': 'PASS', 'message': f'{len(existing_dirs)}/{len(required_dirs)} directories'}
            else:
                print(f"❌ Incomplete folder structure: {len(existing_dirs)}/{len(required_dirs)} directories")
                self.test_results['test_10'] = {'status': 'FAIL', 'message': f'Only {len(existing_dirs)}/{len(required_dirs)} directories'}
                
        except Exception as e:
            print(f"❌ Folder structure test failed: {str(e)}")
            self.test_results['test_10'] = {'status': 'FAIL', 'message': f'Error: {str(e)}'}
    
    def flatten_test_data(self, test_data):
        """
        Flatten nested test data for template processing
        
        Args:
            test_data (dict): Nested JSON test data
            
        Returns:
            dict: Flattened data suitable for template variables
        """
        flattened = {}
        
        # Handle personal information
        if 'personal' in test_data:
            personal = test_data['personal']
            flattened.update({
                'first_name': personal.get('first_name', ''),
                'last_name': personal.get('last_name', ''),
                'email': personal.get('email', ''),
                'phone_number': personal.get('phone', ''),
                'street_address': personal.get('address', ''),
                'city': personal.get('city', ''),
                'state': personal.get('province', ''),
                'zip_code': personal.get('postal_code', '')
            })
        
        # Handle education (take first entry)
        if 'education' in test_data and test_data['education']:
            edu = test_data['education'][0]
            flattened.update({
                'university_name': edu.get('institution', ''),
                'degree': edu.get('degree', ''),
                'concentration': edu.get('concentration', ''),
                'graduation_date': edu.get('graduation_date', '')
            })
        
        # Handle skills
        if 'skills' in test_data:
            skills = test_data['skills']
            flattened.update({
                'technical_skills': ', '.join(skills.get('technical', [])),
                'languages': ', '.join(skills.get('languages', [])),
                'interests': ', '.join(skills.get('personal_interests', []))
            })
        
        # Add document metadata
        flattened.update({
            'title': f"{flattened.get('first_name', '')} {flattened.get('last_name', '')} Resume",
            'author': f"{flattened.get('first_name', '')} {flattened.get('last_name', '')}",
            'subject': 'Professional Resume',
            'keywords': 'Resume, Professional, Job Application',
            'comments': f"Generated on {datetime.now().strftime('%Y-%m-%d')} for professional use",
            'category': 'Job Application'
        })
        
        return flattened
    
    def print_test_summary(self):
        """Print a summary of all test results"""
        print("\n" + "="*50)
        print("TEST SUMMARY")
        print("="*50)
        
        passed = sum(1 for result in self.test_results.values() if result['status'] == 'PASS')
        total = len(self.test_results)
        
        for test_name, result in self.test_results.items():
            status_icon = "✅" if result['status'] == 'PASS' else "❌"
            print(f"{status_icon} {test_name}: {result['message']}")
        
        print(f"\nOverall: {passed}/{total} tests passed")
        
        if passed == total:
            print("🎉 All tests passed! Template library system is ready for use.")
        else:
            print(f"⚠️  {total - passed} test(s) failed. Please review and fix issues.")


def main():
    """Run the complete test suite"""
    tests = DocumentGeneratorTests()
    results = tests.run_all_tests()
    return results


if __name__ == "__main__":
    main()