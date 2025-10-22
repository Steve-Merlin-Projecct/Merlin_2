#!/usr/bin/env python3
"""
Deep Template Analyzer
Examines Word documents including textboxes, headers, footers, and other structures.
"""

import docx
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import zipfile
from lxml import etree


def analyze_document_xml(doc_path: str):
    """
    Analyze the raw XML structure of the Word document.
    This helps find content in textboxes and other special structures.
    """
    print(f"\n{'='*80}")
    print(f"DEEP XML ANALYSIS: {doc_path}")
    print(f"{'='*80}\n")

    # Open as ZIP (docx is a ZIP file)
    with zipfile.ZipFile(doc_path, 'r') as zip_ref:
        # List all files in the docx
        print("Files in DOCX archive:")
        for name in zip_ref.namelist():
            print(f"  {name}")

        # Read the main document XML
        if 'word/document.xml' in zip_ref.namelist():
            print("\n" + "="*80)
            print("MAIN DOCUMENT XML CONTENT")
            print("="*80 + "\n")

            xml_content = zip_ref.read('word/document.xml')
            root = etree.fromstring(xml_content)

            # Find all text elements
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'v': 'urn:schemas-microsoft-com:vml',
                'w10': 'urn:schemas-microsoft-com:office:word',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
            }

            # Extract all text from <w:t> elements
            text_elements = root.xpath('.//w:t', namespaces=namespaces)
            print(f"Found {len(text_elements)} <w:t> text elements in document.xml")

            if text_elements:
                print("\nText content from <w:t> elements:")
                for i, elem in enumerate(text_elements[:50]):  # First 50
                    text = elem.text if elem.text else ""
                    if text.strip():
                        print(f"  [{i:3d}] {text}")

            # Look for textbox content (in VML shapes)
            textboxes = root.xpath('.//v:textbox', namespaces=namespaces)
            print(f"\nFound {len(textboxes)} textboxes")

            if textboxes:
                print("\nTextbox content:")
                for i, textbox in enumerate(textboxes):
                    # Get all text within textbox
                    texts = textbox.xpath('.//w:t', namespaces=namespaces)
                    textbox_content = ' '.join([t.text for t in texts if t.text])
                    if textbox_content.strip():
                        print(f"  Textbox {i}: {textbox_content[:200]}")

        # Check for header/footer
        for part_name in ['word/header1.xml', 'word/header2.xml', 'word/footer1.xml', 'word/footer2.xml']:
            if part_name in zip_ref.namelist():
                print(f"\n{part_name} exists - checking content...")
                xml_content = zip_ref.read(part_name)
                root = etree.fromstring(xml_content)
                text_elements = root.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                for elem in text_elements:
                    if elem.text and elem.text.strip():
                        print(f"  {elem.text}")


def extract_all_text_comprehensive(doc_path: str) -> list:
    """
    Extract ALL text from document using multiple methods.
    """
    all_text = []

    doc = Document(doc_path)

    # Method 1: Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(('paragraph', para.text))

    # Method 2: Tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    if para.text.strip():
                        all_text.append(('table_cell', para.text))

    # Method 3: Headers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                if para.text.strip():
                    all_text.append(('header', para.text))

        # Footer
        if section.footer:
            for para in section.footer.paragraphs:
                if para.text.strip():
                    all_text.append(('footer', para.text))

    # Method 4: Try to access textboxes through runs
    for para in doc.paragraphs:
        for run in para.runs:
            # Check if run has any embedded objects
            if run._element.xml:
                if 'textbox' in run._element.xml.lower() or 'txbx' in run._element.xml.lower():
                    all_text.append(('potential_textbox', run._element.xml[:200]))

    return all_text


def print_comprehensive_analysis(doc_path: str):
    """Print comprehensive analysis using all methods."""
    print(f"\n{'='*80}")
    print(f"COMPREHENSIVE TEXT EXTRACTION: {doc_path}")
    print(f"{'='*80}\n")

    all_text = extract_all_text_comprehensive(doc_path)

    print(f"Total text elements found: {len(all_text)}\n")

    for i, (source, text) in enumerate(all_text):
        print(f"[{i:3d}] ({source:15s}) {text[:100]}")


if __name__ == "__main__":
    template_4_path = "/workspace/.trees/convert-raw-template-files-into-ready-for-producti/content_template_library/manual_converted/template_4_source.docx"
    template_5_path = "/workspace/.trees/convert-raw-template-files-into-ready-for-producti/content_template_library/manual_converted/template_5_source.docx"

    # Deep analysis of Template 4
    analyze_document_xml(template_4_path)
    print_comprehensive_analysis(template_4_path)

    print("\n\n" + "#"*80 + "\n\n")

    # Deep analysis of Template 5
    analyze_document_xml(template_5_path)
    print_comprehensive_analysis(template_5_path)
